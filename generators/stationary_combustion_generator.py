from __future__ import annotations

import csv
import random
import re
import zipfile
from calendar import monthrange
from datetime import date, datetime, timedelta
from decimal import Decimal, ROUND_HALF_UP
from io import BytesIO, StringIO
from typing import Any

import openpyxl
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

from components.stationary_combustion.units import default_fuel_volume_unit
from utils.bad_data import (
    FTYPE_CURRENCY_UNIT,
    FTYPE_DATE_TIME,
    FTYPE_IDENTIFIER,
    FTYPE_NUMERIC,
    FTYPE_TEXT,
    corrupt_records,
    get_bad_data_config,
)
from utils.distractor_fields import resolve_distractor_plan, resolve_tabular_value
from utils.layouts import resolve_layout_plan

# ── field-type maps for each document type ────────────────────────────────────

_FUEL_INVOICE_FIELD_TYPES: dict[str, str] = {
    "company": FTYPE_TEXT,
    "supplier": FTYPE_TEXT,
    "customer": FTYPE_TEXT,
    "site": FTYPE_TEXT,
    "country": FTYPE_TEXT,
    "equipment": FTYPE_TEXT,
    "emission_source": FTYPE_TEXT,
    "fuel": FTYPE_TEXT,
    "customer_code": FTYPE_IDENTIFIER,
    "account_number": FTYPE_IDENTIFIER,
    "invoice_no": FTYPE_IDENTIFIER,
    "invoiced_date": FTYPE_DATE_TIME,
    "received_date": FTYPE_DATE_TIME,
    "due_date": FTYPE_DATE_TIME,
    "quantity": FTYPE_NUMERIC,
    "unit_price": FTYPE_NUMERIC,
    "fuel_cost": FTYPE_NUMERIC,
    "delivery_charge": FTYPE_NUMERIC,
    "subtotal": FTYPE_NUMERIC,
    "vat": FTYPE_NUMERIC,
    "total": FTYPE_NUMERIC,
    "currency": FTYPE_CURRENCY_UNIT,
    "unit": FTYPE_CURRENCY_UNIT,
}

_DELIVERY_NOTE_FIELD_TYPES: dict[str, str] = {
    "company": FTYPE_TEXT,
    "supplier": FTYPE_TEXT,
    "customer": FTYPE_TEXT,
    "site": FTYPE_TEXT,
    "country": FTYPE_TEXT,
    "equipment": FTYPE_TEXT,
    "fuel": FTYPE_TEXT,
    "delivery_note_no": FTYPE_IDENTIFIER,
    "period_label": FTYPE_DATE_TIME,
    "quantity": FTYPE_NUMERIC,
    "currency": FTYPE_CURRENCY_UNIT,
    "unit": FTYPE_CURRENCY_UNIT,
}

_FUEL_CARD_TRANSACTION_FIELD_TYPES: dict[str, str] = {
    "site": FTYPE_TEXT,
    "country": FTYPE_TEXT,
    "equipment": FTYPE_TEXT,
    "emission_source": FTYPE_TEXT,
    "merchant": FTYPE_TEXT,
    "fuel": FTYPE_TEXT,
    "card_number": FTYPE_IDENTIFIER,
    "reference": FTYPE_IDENTIFIER,
    "quantity": FTYPE_NUMERIC,
    "unit_price": FTYPE_NUMERIC,
    "total": FTYPE_NUMERIC,
    "unit": FTYPE_CURRENCY_UNIT,
}

_GENERATOR_LOG_FIELD_TYPES: dict[str, str] = {
    "site": FTYPE_TEXT,
    "company": FTYPE_TEXT,
    "country": FTYPE_TEXT,
    "equipment": FTYPE_TEXT,
    "emission_source": FTYPE_TEXT,
    "fuel": FTYPE_TEXT,
    "period": FTYPE_DATE_TIME,
    "fuel_used": FTYPE_NUMERIC,
    "unit": FTYPE_CURRENCY_UNIT,
}

_BEMS_ASSET_FIELD_TYPES: dict[str, str] = {
    "equipment_name": FTYPE_TEXT,
    "emission_source": FTYPE_TEXT,
    "fuel": FTYPE_TEXT,
    "asset_tag": FTYPE_IDENTIFIER,
    "sensor_name": FTYPE_IDENTIFIER,
    "quantity": FTYPE_NUMERIC,
    "operating_hours": FTYPE_NUMERIC,
    "unit": FTYPE_CURRENCY_UNIT,
}

TWOPLACES = Decimal("0.01")
PAGE_W, PAGE_H = A4
_SPECIAL_CHARS_SUFFIX = ' & < " £ € \u00a0\u2014\u200f'

STATIONARY_TRANSLATIONS: dict[str, dict[str, str]] = {
    "en": {
        "fuel_invoice_title": "Fuel Invoice",
        "bill_to": "Bill To",
        "invoice_details": "Invoice Details",
        "delivery_site": "Delivery Site",
        "invoice_no": "Invoice No",
        "invoice_date": "Invoice Date",
        "account_no": "Account No",
        "vehicle": "Vehicle",
        "billing_period": "Billing Period",
        "invoiced_date": "Invoiced Date",
        "received_date": "Received Date",
        "due_date": "Due Date",
        "currency": "Currency",
        "country": "Country",
        "product": "Product",
        "quantity": "Quantity",
        "unit": "Unit",
        "unit_price": "Unit Price",
        "amount": "Amount",
        "delivery_charge": "Delivery Charge",
        "each": "Each",
        "subtotal": "Subtotal",
        "total": "Total",
        "generator_operation_log_title": "Generator Operation Log",
        "litres": "Litres",
        "received": "Received",
        "stationary_equipment": "Stationary Equipment",
        "test_run": "Test run",
        "power_outage": "Power outage",
        "maintenance_test": "Maintenance test",
        "load_bank_test": "Load bank test",
        "fuel_invoice_footer": "Generated for Scope 1 stationary combustion. Delivery-site details are illustrative.",
        "delivery_note_title": "Fuel Delivery Note",
        "supplier": "Supplier",
        "delivery_note_no": "Delivery Note No",
        "delivery_date": "Delivery Date",
        "customer": "Customer",
        "delivery_address": "Delivery Address",
        "delivery_details": "Delivery Details",
        "delivery_confirmation": "Delivery Confirmation",
        "product_delivered": "Product Delivered",
        "tank_equipment": "Tank / Equipment",
        "delivered_quantity": "Delivered Quantity",
        "driver_ref": "Driver Ref",
        "customer_signature": "Customer Signature",
        "delivery_note_footer": "Generated for Scope 1 stationary combustion delivery-note testing.",
        "fuel_card_title": "Fuel Card Statement",
        "account_name": "Account Name",
        "provider": "Provider",
        "statement_period": "Statement Period",
        "card_no": "Card No",
        "date": "Date",
        "merchant": "Merchant",
        "reference": "Reference",
        "qty": "Qty",
        "statement_total": "Statement Total",
        "net_amount": "Net Amount",
        "vat": "VAT",
        "gross_amount": "Gross Amount",
        "fuel_card_footer": "Statement generated for stationary-equipment fuel-card QA.",
        "account_details": "Account Details",
        "statement_details": "Statement Details",
        "company": "Company",
        "site": "Site",
        "equipment": "Equipment",
        "emission_source": "Emission Source",
        "generator_log_sheet_title": "Generator Log",
        "period": "Period",
        "start_time": "Start Time",
        "end_time": "End Time",
        "run_hours": "Run Hours",
        "start_fuel": "Start Fuel",
        "end_fuel": "End Fuel",
        "fuel_used": "Fuel Used",
        "notes": "Notes",
        "bems_equipment_title": "BEMS Fuel Consumption Summary",
        "assets": "Assets",
        "operating_hours": "Operating Hours",
        "top_asset": "Top Asset",
        "equipment_trend_snapshot": "Equipment Trend Snapshot",
        "equipment_tag": "Equipment Tag",
        "equipment_name": "Equipment Name",
        "fuel_type": "Fuel Type",
        "consumption": "Consumption",
        "dashboard_summary_footer": "Dashboard summary generated from BEMS trend data.",
        "bems_summary_sheet_title": "BEMS Summary",
        "reporting_period": "Reporting Period",
        "bems_time_series_title": "BEMS Time-Series Trend Export",
        "interval": "Interval",
        "rows": "Rows",
        "timestamp": "Timestamp",
        "sensor_name": "Sensor Name",
        "value": "Value",
        "time_series_footer": "Time-series export rendered as PDF preview.",
        "time_series_word_footer": "Time-series export rendered as Word preview.",
        "site_fallback": "Site",
    },
    "fr": {
        "fuel_invoice_title": "Facture de carburant",
        "bill_to": "Facturer a",
        "invoice_details": "Details de facture",
        "delivery_site": "Site de livraison",
        "invoice_no": "No de facture",
        "invoice_date": "Date de facture",
        "account_no": "No de compte",
        "vehicle": "Vehicule",
        "billing_period": "Periode de facturation",
        "invoiced_date": "Date de facturation",
        "received_date": "Date de reception",
        "due_date": "Date d'echeance",
        "currency": "Devise",
        "country": "Pays",
        "product": "Produit",
        "quantity": "Quantite",
        "unit": "Unite",
        "unit_price": "Prix unitaire",
        "amount": "Montant",
        "delivery_charge": "Frais de livraison",
        "each": "Chaque",
        "subtotal": "Sous-total",
        "total": "Total",
        "generator_operation_log_title": "Journal d'exploitation generateur",
        "litres": "Litres",
        "received": "Recu",
        "stationary_equipment": "Equipement stationnaire",
        "test_run": "Essai",
        "power_outage": "Coupure de courant",
        "maintenance_test": "Essai de maintenance",
        "load_bank_test": "Essai au banc de charge",
        "fuel_invoice_footer": "Genere pour la combustion stationnaire du scope 1. Les details du site de livraison sont indicatifs.",
        "delivery_note_title": "Bon de livraison carburant",
        "supplier": "Fournisseur",
        "delivery_note_no": "No de bon de livraison",
        "delivery_date": "Date de livraison",
        "customer": "Client",
        "delivery_address": "Adresse de livraison",
        "delivery_details": "Details de livraison",
        "delivery_confirmation": "Confirmation de livraison",
        "product_delivered": "Produit livre",
        "tank_equipment": "Cuve / Equipement",
        "delivered_quantity": "Quantite livree",
        "driver_ref": "Ref chauffeur",
        "customer_signature": "Signature client",
        "delivery_note_footer": "Genere pour les tests de bons de livraison de combustion stationnaire du scope 1.",
        "fuel_card_title": "Releve de carte carburant",
        "account_name": "Nom du compte",
        "provider": "Fournisseur",
        "statement_period": "Periode du releve",
        "card_no": "No carte",
        "date": "Date",
        "merchant": "Commercant",
        "reference": "Reference",
        "qty": "Qt",
        "statement_total": "Total du releve",
        "net_amount": "Montant net",
        "vat": "TVA",
        "gross_amount": "Montant brut",
        "fuel_card_footer": "Releve genere pour l'assurance qualite des cartes carburant d'equipements stationnaires.",
        "account_details": "Details du compte",
        "statement_details": "Details du releve",
        "company": "Entreprise",
        "site": "Site",
        "equipment": "Equipement",
        "emission_source": "Source d'emission",
        "generator_log_sheet_title": "Journal generateur",
        "period": "Periode",
        "start_time": "Heure debut",
        "end_time": "Heure fin",
        "run_hours": "Heures de marche",
        "start_fuel": "Carburant debut",
        "end_fuel": "Carburant fin",
        "fuel_used": "Carburant utilise",
        "notes": "Notes",
        "bems_equipment_title": "Resume BEMS de consommation de carburant",
        "assets": "Actifs",
        "operating_hours": "Heures de fonctionnement",
        "top_asset": "Actif principal",
        "equipment_trend_snapshot": "Apercu de tendance des equipements",
        "equipment_tag": "Identifiant equipement",
        "equipment_name": "Nom equipement",
        "fuel_type": "Type de carburant",
        "consumption": "Consommation",
        "dashboard_summary_footer": "Resume du tableau de bord genere a partir des donnees de tendance BEMS.",
        "bems_summary_sheet_title": "Resume BEMS",
        "reporting_period": "Periode de reporting",
        "bems_time_series_title": "Export BEMS de tendance chronologique",
        "interval": "Intervalle",
        "rows": "Lignes",
        "timestamp": "Horodatage",
        "sensor_name": "Nom du capteur",
        "value": "Valeur",
        "time_series_footer": "Export chronologique restitue comme apercu PDF.",
        "time_series_word_footer": "Export chronologique restitue comme apercu Word.",
        "site_fallback": "Site",
    },
    "de": {
        "fuel_invoice_title": "Kraftstoffrechnung",
        "bill_to": "Rechnung an",
        "invoice_details": "Rechnungsdetails",
        "delivery_site": "Lieferstandort",
        "invoice_no": "Rechnungsnr.",
        "invoice_date": "Rechnungsdatum",
        "account_no": "Kontonummer",
        "vehicle": "Fahrzeug",
        "billing_period": "Abrechnungszeitraum",
        "invoiced_date": "Rechnungsdatum",
        "received_date": "Eingangsdatum",
        "due_date": "Faelligkeitsdatum",
        "currency": "Wahrung",
        "country": "Land",
        "product": "Produkt",
        "quantity": "Menge",
        "unit": "Einheit",
        "unit_price": "Stuckpreis",
        "amount": "Betrag",
        "delivery_charge": "Lieferkosten",
        "each": "Je",
        "subtotal": "Zwischensumme",
        "total": "Gesamt",
        "generator_operation_log_title": "Generatorbetriebsprotokoll",
        "litres": "Liter",
        "received": "Erhalten",
        "stationary_equipment": "Stationare Anlage",
        "test_run": "Testlauf",
        "power_outage": "Stromausfall",
        "maintenance_test": "Wartungstest",
        "load_bank_test": "Lastbanktest",
        "fuel_invoice_footer": "Erstellt fur Scope-1-Standverbrennung. Lieferstandortdetails dienen nur zur Veranschaulichung.",
        "delivery_note_title": "Kraftstofflieferschein",
        "supplier": "Lieferant",
        "delivery_note_no": "Lieferscheinnr.",
        "delivery_date": "Lieferdatum",
        "customer": "Kunde",
        "delivery_address": "Lieferadresse",
        "delivery_details": "Lieferdetails",
        "delivery_confirmation": "Lieferbestatigung",
        "product_delivered": "Geliefertes Produkt",
        "tank_equipment": "Tank / Anlage",
        "delivered_quantity": "Gelieferte Menge",
        "driver_ref": "Fahrer-Ref",
        "customer_signature": "Kundenunterschrift",
        "delivery_note_footer": "Erstellt fur Tests von Lieferscheinen zur stationaren Verbrennung in Scope 1.",
        "fuel_card_title": "Tankkartenabrechnung",
        "account_name": "Kontoname",
        "provider": "Anbieter",
        "statement_period": "Abrechnungszeitraum",
        "card_no": "Kartennr.",
        "date": "Datum",
        "merchant": "Handler",
        "reference": "Referenz",
        "qty": "Menge",
        "statement_total": "Abrechnungssumme",
        "net_amount": "Nettobetrag",
        "vat": "MwSt.",
        "gross_amount": "Bruttobetrag",
        "fuel_card_footer": "Abrechnung fur die QS stationarer Tankkartenvorgange erstellt.",
        "account_details": "Kontodetails",
        "statement_details": "Abrechnungsdetails",
        "company": "Unternehmen",
        "site": "Standort",
        "equipment": "Anlage",
        "emission_source": "Emissionsquelle",
        "generator_log_sheet_title": "Generatorprotokoll",
        "period": "Zeitraum",
        "start_time": "Startzeit",
        "end_time": "Endzeit",
        "run_hours": "Betriebsstunden",
        "start_fuel": "Kraftstoff Start",
        "end_fuel": "Kraftstoff Ende",
        "fuel_used": "Verbrauchter Kraftstoff",
        "notes": "Hinweise",
        "bems_equipment_title": "BEMS-Kraftstoffverbrauchsbericht",
        "assets": "Anlagen",
        "operating_hours": "Betriebsstunden",
        "top_asset": "Top-Anlage",
        "equipment_trend_snapshot": "Anlagen-Trendubersicht",
        "equipment_tag": "Anlagenkennzeichen",
        "equipment_name": "Anlagenname",
        "fuel_type": "Kraftstoffart",
        "consumption": "Verbrauch",
        "dashboard_summary_footer": "Dashboard-Zusammenfassung aus BEMS-Trenddaten erstellt.",
        "bems_summary_sheet_title": "BEMS Ubersicht",
        "reporting_period": "Berichtszeitraum",
        "bems_time_series_title": "BEMS-Zeitreihenexport",
        "interval": "Intervall",
        "rows": "Zeilen",
        "timestamp": "Zeitstempel",
        "sensor_name": "Sensorname",
        "value": "Wert",
        "time_series_footer": "Zeitreihenexport als PDF-Vorschau dargestellt.",
        "time_series_word_footer": "Zeitreihenexport als Word-Vorschau dargestellt.",
        "site_fallback": "Standort",
    },
    "nl": {
        "fuel_invoice_title": "Brandstoffactuur",
        "bill_to": "Factureren aan",
        "invoice_details": "Factuurgegevens",
        "delivery_site": "Leveringslocatie",
        "invoice_no": "Factuurnr.",
        "invoice_date": "Factuurdatum",
        "account_no": "Rekeningnummer",
        "vehicle": "Voertuig",
        "billing_period": "Facturatieperiode",
        "invoiced_date": "Factuurdatum",
        "received_date": "Ontvangstdatum",
        "due_date": "Vervaldatum",
        "currency": "Valuta",
        "country": "Land",
        "product": "Product",
        "quantity": "Hoeveelheid",
        "unit": "Eenheid",
        "unit_price": "Eenheidsprijs",
        "amount": "Bedrag",
        "delivery_charge": "Leveringskosten",
        "each": "Per stuk",
        "subtotal": "Subtotaal",
        "total": "Totaal",
        "generator_operation_log_title": "Generatorbedrijfslog",
        "litres": "Liter",
        "received": "Ontvangen",
        "stationary_equipment": "Stationaire installatie",
        "test_run": "Testrun",
        "power_outage": "Stroomstoring",
        "maintenance_test": "Onderhoudstest",
        "load_bank_test": "Belastingbanktest",
        "fuel_invoice_footer": "Gegenereerd voor Scope 1 stationaire verbranding. Details van de leveringslocatie zijn illustratief.",
        "delivery_note_title": "Brandstofleverbon",
        "supplier": "Leverancier",
        "delivery_note_no": "Leverbonnr.",
        "delivery_date": "Leverdatum",
        "customer": "Klant",
        "delivery_address": "Leveradres",
        "delivery_details": "Leveringsdetails",
        "delivery_confirmation": "Leveringsbevestiging",
        "product_delivered": "Geleverd product",
        "tank_equipment": "Tank / Installatie",
        "delivered_quantity": "Geleverde hoeveelheid",
        "driver_ref": "Chauffeursref",
        "customer_signature": "Handtekening klant",
        "delivery_note_footer": "Gegenereerd voor het testen van leverbonnen voor stationaire verbranding in Scope 1.",
        "fuel_card_title": "Tankkaartoverzicht",
        "account_name": "Accountnaam",
        "provider": "Aanbieder",
        "statement_period": "Overzichtsperiode",
        "card_no": "Kaartnr.",
        "date": "Datum",
        "merchant": "Leverancier",
        "reference": "Referentie",
        "qty": "Aantal",
        "statement_total": "Totaal overzicht",
        "net_amount": "Nettobedrag",
        "vat": "Btw",
        "gross_amount": "Brutobedrag",
        "fuel_card_footer": "Overzicht gegenereerd voor QA van tankkaarttransacties voor stationaire apparatuur.",
        "account_details": "Accountgegevens",
        "statement_details": "Overzichtsgegevens",
        "company": "Bedrijf",
        "site": "Locatie",
        "equipment": "Installatie",
        "emission_source": "Emissiebron",
        "generator_log_sheet_title": "Generatorlog",
        "period": "Periode",
        "start_time": "Starttijd",
        "end_time": "Eindtijd",
        "run_hours": "Draaiuren",
        "start_fuel": "Brandstof start",
        "end_fuel": "Brandstof eind",
        "fuel_used": "Verbruikte brandstof",
        "notes": "Notities",
        "bems_equipment_title": "BEMS-brandstofverbruiksoverzicht",
        "assets": "Assets",
        "operating_hours": "Bedrijfsuren",
        "top_asset": "Belangrijkste asset",
        "equipment_trend_snapshot": "Momentopname apparatuurtendens",
        "equipment_tag": "Apparaatcode",
        "equipment_name": "Apparaatnaam",
        "fuel_type": "Brandstoftype",
        "consumption": "Verbruik",
        "dashboard_summary_footer": "Dashboardsamenvatting gegenereerd uit BEMS-trendgegevens.",
        "bems_summary_sheet_title": "BEMS-overzicht",
        "reporting_period": "Rapportageperiode",
        "bems_time_series_title": "BEMS-tijdreeks export",
        "interval": "Interval",
        "rows": "Rijen",
        "timestamp": "Tijdstempel",
        "sensor_name": "Sensornaam",
        "value": "Waarde",
        "time_series_footer": "Tijdreeks-export weergegeven als PDF-voorbeeld.",
        "time_series_word_footer": "Tijdreeks-export weergegeven als Word-voorbeeld.",
        "site_fallback": "Locatie",
    },
}


def _q2(value) -> Decimal:
    if not isinstance(value, Decimal):
        value = Decimal(str(value))
    return value.quantize(TWOPLACES, rounding=ROUND_HALF_UP)


def _parse_decimal(value, fallback: str = "0") -> Decimal:
    if value in (None, ""):
        return Decimal(fallback)
    if isinstance(value, Decimal):
        return value
    return Decimal(str(value))


def _parse_date(value: str) -> date:
    return datetime.strptime(value, "%Y-%m-%d").date()


def _language(raw_config: dict) -> str:
    language = str(raw_config.get("document", {}).get("language", "en")).lower()
    return language if language in STATIONARY_TRANSLATIONS else "en"


def _tr(raw_config: dict, key: str, **kwargs) -> str:
    template = STATIONARY_TRANSLATIONS[_language(raw_config)][key]
    return template.format(**kwargs) if kwargs else template


def _fmt_date(value) -> str:
    if isinstance(value, str):
        return value
    return value.strftime("%d %b %Y")


def _fmt_num(value, spec: str = ",.2f") -> str:
    """Format a numeric value; return as-is if already a corrupted string."""
    if isinstance(value, str):
        return value
    try:
        return format(value, spec)
    except Exception:
        return str(value)


def _fmt_money(value) -> str:
    if isinstance(value, str):
        return value
    try:
        return f"{_q2(value):,.2f}"
    except Exception:
        return str(value)


def _fmt_optional_number(value, suffix: str = "") -> str:
    if value in (None, ""):
        return ""
    return f"{_fmt_money(value)}{suffix}"


def _currency_symbol(currency_raw: str) -> str:
    mapping = {
        "(£)": "£",
        "(€)": "€",
        "($)": "$",
        "(¥)": "¥",
        "(kr)": "kr",
        "(Ft)": "Ft",
    }
    for token, symbol in mapping.items():
        if token in currency_raw:
            return symbol
    return ""


_FUEL_CARD_HEADER_KEYS = {
    "card_no": "card_no",
    "date": "date",
    "merchant": "merchant",
    "site": "site",
    "country": "country",
    "equipment": "equipment",
    "emission_source": "emission_source",
    "product": "product",
    "qty": "qty",
    "unit": "unit",
    "unit_price": "unit_price",
    "total": "total",
    "currency": "currency",
}

_GENERATOR_LOG_HEADER_KEYS = {
    "company": "company",
    "site": "site",
    "country": "country",
    "date": "date",
    "start_time": "start_time",
    "end_time": "end_time",
    "run_hours": "run_hours",
    "start_fuel": "start_fuel",
    "end_fuel": "end_fuel",
    "fuel_used": "fuel_used",
    "unit": "unit",
    "equipment": "equipment",
    "emission_source": "emission_source",
    "fuel_type": "fuel_type",
    "notes": "notes",
}

_BEMS_EQUIPMENT_HEADER_KEYS = {
    "equipment_tag": "equipment_tag",
    "equipment_name": "equipment_name",
    "emission_source": "emission_source",
    "fuel_type": "fuel_type",
    "consumption": "consumption",
    "unit": "unit",
    "operating_hours": "operating_hours",
}

_BEMS_TIME_SERIES_HEADER_KEYS = {
    "timestamp": "timestamp",
    "site": "site",
    "equipment_tag": "equipment_tag",
    "sensor_name": "sensor_name",
    "value": "value",
    "unit": "unit",
}


def _stationary_layout_context(raw_config: dict) -> dict:
    return {
        "language": raw_config.get("document", {}).get("language", "en"),
        "period_label": raw_config.get("financial_period", {}).get("label", ""),
        "company_labels": [company.get("label", f"Company {index + 1}") for index, company in enumerate(raw_config.get("companies", []))],
        "bems_report_type": raw_config.get("document", {}).get("bems_report_type", "equipment_trend_report"),
    }


def _stationary_layout_plan(raw_config: dict, output_format: str, *, document_type: str | None = None) -> dict:
    return resolve_layout_plan(
        raw_config.get("document", {}),
        random_seed=int(raw_config.get("random_seed", 42)),
        category="stationary_combustion",
        document_type=document_type or str(raw_config.get("document_type") or raw_config.get("document", {}).get("type") or "fuel_invoice"),
        output_format=output_format,
        context=_stationary_layout_context(raw_config),
    )


def _ordered_stationary_field_ids(plan: dict, default_ids: list[str]) -> list[str]:
    ordered = [field_id for field_id in plan.get("column_order", []) if field_id in default_ids]
    return ordered or list(default_ids)


def _stationary_header_text(
    raw_config: dict,
    plan: dict,
    field_id: str,
    default_key: str,
    distractor_fields: dict[str, Any] | None = None,
) -> str:
    if distractor_fields and field_id in distractor_fields:
        return str(distractor_fields[field_id].label)
    return plan.get("header_aliases", {}).get(field_id) or _tr(raw_config, default_key)


def _stationary_document_type(raw_config: dict, document_type: str | None = None) -> str:
    return str(document_type or raw_config.get("document_type") or raw_config.get("document", {}).get("type") or "fuel_invoice")


def _stationary_distractor_plan(
    raw_config: dict,
    output_format: str,
    *,
    document_type: str | None = None,
    artifact_key: str = "root",
):
    return resolve_distractor_plan(
        raw_config.get("document", {}),
        random_seed=int(raw_config.get("random_seed", 42)),
        category="stationary_combustion",
        document_type=_stationary_document_type(raw_config, document_type),
        output_format=output_format,
        artifact_key=artifact_key,
        context=_stationary_layout_context(raw_config),
    )


def _augment_stationary_field_ids(base_field_ids: list[str], distractor_plan) -> list[str]:
    ordered = list(base_field_ids)
    if not distractor_plan or not getattr(distractor_plan, "enabled", False):
        return ordered
    for field in distractor_plan.tabular_fields:
        insert_at = len(ordered)
        if field.anchor in ordered:
            anchor_index = ordered.index(field.anchor)
            insert_at = anchor_index + 1 if field.position == "after" else anchor_index
        ordered.insert(insert_at, field.field_id)
    return ordered


def _stationary_distractor_field_map(distractor_plan) -> dict[str, Any]:
    if not distractor_plan or not getattr(distractor_plan, "enabled", False):
        return {}
    return {field.field_id: field for field in distractor_plan.tabular_fields}


def _stationary_scope_subkey(
    field,
    row_key: str,
    *,
    statement_key: str | None = None,
    block_key: str | None = None,
) -> str:
    if getattr(field, "row_scope", "row") == "statement":
        return statement_key or block_key or row_key
    if getattr(field, "row_scope", "row") == "block":
        return block_key or statement_key or row_key
    if getattr(field, "row_scope", "row") == "file":
        return "file"
    return row_key


def _stationary_row_values(
    row_map: dict[str, Any],
    ordered_ids: list[str],
    distractor_plan,
    *,
    row_key: str,
    statement_key: str | None = None,
    block_key: str | None = None,
) -> list[Any]:
    distractor_fields = _stationary_distractor_field_map(distractor_plan)
    values: list[Any] = []
    for field_id in ordered_ids:
        if field_id in distractor_fields:
            field = distractor_fields[field_id]
            values.append(
                resolve_tabular_value(
                    distractor_plan,
                    field,
                    _stationary_scope_subkey(field, row_key, statement_key=statement_key, block_key=block_key),
                )
            )
            continue
        values.append(row_map[field_id])
    return values


def _stationary_document_lines(distractor_plan, *, placement: str | None = None) -> list[str]:
    if not distractor_plan or not getattr(distractor_plan, "enabled", False):
        return []
    return [
        f"{field.label}: {field.value}"
        for field in distractor_plan.document_fields
        if placement is None or field.placement == placement
    ]


def _stationary_document_pairs(distractor_plan, *, placement: str | None = None) -> list[tuple[str, str]]:
    if not distractor_plan or not getattr(distractor_plan, "enabled", False):
        return []
    return [
        (field.label, field.value)
        for field in distractor_plan.document_fields
        if placement is None or field.placement == placement
    ]


def _write_stationary_csv_preamble(writer: csv.writer, plan: dict) -> None:
    for row in plan.get("preamble_rows", []):
        writer.writerow(row)
    for _ in range(int(plan.get("header_row_offset", 0) or 0)):
        writer.writerow([])


def _write_stationary_xlsx_preamble(sheet, start_row: int, plan: dict) -> int:
    row_index = start_row
    for row in plan.get("preamble_rows", []):
        for column_index, value in enumerate(row, start=1):
            sheet.cell(row=row_index, column=column_index, value=value)
        row_index += 1
    row_index += int(plan.get("header_row_offset", 0) or 0)
    return row_index


def _with_special_chars(config: dict, value: str) -> str:
    if not value:
        return value
    if not config.get("document", {}).get("inject_special_chars", False):
        return value
    return value + _SPECIAL_CHARS_SUFFIX


def _financial_period(raw_config: dict) -> dict:
    fp = raw_config.get("financial_period", {})
    return {
        "label": fp.get("label", ""),
        "start_date": _parse_date(fp.get("start_date", "2026-01-01")),
        "end_date": _parse_date(fp.get("end_date", "2026-01-31")),
    }


def _months_in_range(start_date: date, end_date: date) -> list[tuple[int, int]]:
    months: list[tuple[int, int]] = []
    current = date(start_date.year, start_date.month, 1)
    while current <= end_date:
        months.append((current.year, current.month))
        if current.month == 12:
            current = date(current.year + 1, 1, 1)
        else:
            current = date(current.year, current.month + 1, 1)
    return months


def _bems_interval_minutes(raw_config: dict) -> int:
    minutes = int(raw_config.get("document", {}).get("bems_interval_minutes", 60))
    return minutes if minutes in {15, 30, 60} else 60


def _bems_report_type(raw_config: dict) -> str:
    report_type = str(raw_config.get("document", {}).get("bems_report_type", "equipment_trend_report"))
    return report_type if report_type in {"equipment_trend_report", "time_series_trend_export"} else "equipment_trend_report"


def _timestamp_range(start_date: date, end_date: date, interval_minutes: int) -> list[datetime]:
    timestamps: list[datetime] = []
    current = datetime.combine(start_date, datetime.min.time())
    end_dt = datetime.combine(end_date + timedelta(days=1), datetime.min.time())
    while current < end_dt:
        timestamps.append(current)
        current += timedelta(minutes=interval_minutes)
    return timestamps


def _iter_company_sites(raw_config: dict):
    for company_index, company in enumerate(raw_config.get("companies", []), start=1):
        for site_index, site in enumerate(company.get("sites", []), start=1):
            yield company_index, site_index, company, site


def _multi_document_count(raw_config: dict) -> int:
    """How many separate documents the user asked to generate (>= 1)."""
    try:
        count = int(raw_config.get("document", {}).get("doc_count", 1) or 1)
    except (TypeError, ValueError):
        count = 1
    return max(count, 1)


def _is_multi_document(raw_config: dict) -> bool:
    return _multi_document_count(raw_config) > 1


def _slugify_filename(value: str, fallback: str) -> str:
    slug = re.sub(r"[^A-Za-z0-9._-]+", "_", str(value or "").strip()).strip("_")
    return slug or fallback


def _build_zip_archive(documents: list[tuple[str, bytes]]) -> bytes:
    archive_buffer = BytesIO()
    with zipfile.ZipFile(archive_buffer, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for filename, content in documents:
            archive.writestr(filename, content)
    return archive_buffer.getvalue()


def _zip_documents(raw_config: dict, items: list[dict], render_one, name_for, ext: str) -> bytes:
    """Render each item as its own file and bundle them into a ZIP archive."""
    documents: list[tuple[str, bytes]] = []
    for index, item in enumerate(items):
        documents.append((name_for(item, index, ext), render_one(raw_config, [item])))
    return _build_zip_archive(documents)


def _fuel_invoice_filename(record: dict, index: int, ext: str) -> str:
    return f"{_slugify_filename(record.get('invoice_no'), f'invoice_{index + 1:04d}')}.{ext}"


def _stationary_account_number(company: dict, seed: int, company_index: int) -> str:
    explicit = str(company.get("account_number", "") or "").strip()
    if explicit:
        return explicit
    rng = random.Random(f"{seed}:stationary_account:{company_index}")
    prefix = (company.get("customer_code") or "ACC").strip() or "ACC"
    return f"{prefix}-{rng.randint(100000, 999999)}"


def _cross_scope_enabled(raw_config: dict) -> bool:
    return bool(raw_config.get("document", {}).get("cross_scope_items", False))


_CROSS_SCOPE_VEHICLE_LINES = [
    {"fuel": "Diesel", "vehicle_type": "Van (LGV)"},
    {"fuel": "Diesel", "vehicle_type": "HGV"},
    {"fuel": "Petrol", "vehicle_type": "Passenger car"},
]


def _cross_scope_vehicle_line(raw_config: dict, rng: random.Random) -> dict:
    """A road-transport fuel line mixed into a stationary invoice. Extraction
    for stationary combustion must classify it as Scope 1 mobile instead."""
    template = rng.choice(_CROSS_SCOPE_VEHICLE_LINES)
    registration = f"{rng.choice(['AB', 'BD', 'KL', 'MN'])}{rng.randint(10, 72)} {rng.choice(['CDE', 'JKL', 'RST', 'XYZ'])}"
    quantity = _q2(Decimal(rng.randrange(35, 85)))
    unit_price = _q2(Decimal(str(round(rng.uniform(1.35, 1.75), 2))))
    return {
        "category": "mobile_combustion",
        "description": f"{template['fuel']} - {_tr(raw_config, 'vehicle')} {registration}",
        "fuel": template["fuel"],
        "vehicle_reg": registration,
        "quantity": quantity,
        "unit": "Litres",
        "unit_price": unit_price,
        "amount": _q2(quantity * unit_price),
    }


def _fuel_invoice_meta_lines(raw_config: dict, record: dict, distractor_plan) -> list[str]:
    return [
        f"{_tr(raw_config, 'invoice_no')}: {record['invoice_no']}",
        f"{_tr(raw_config, 'account_no')}: {record.get('account_number', '')}",
        f"{_tr(raw_config, 'invoiced_date')}: {_fmt_date(record['invoiced_date'])}",
        f"{_tr(raw_config, 'received_date')}: {_fmt_date(record['received_date'])}",
        f"{_tr(raw_config, 'due_date')}: {_fmt_date(record['due_date'])}",
        f"{_tr(raw_config, 'currency')}: {record['currency']}",
        f"{_tr(raw_config, 'country')}: {record['country']}",
        *_stationary_document_lines(distractor_plan, placement="meta"),
    ]


def _fuel_invoice_line_rows(raw_config: dict, record: dict) -> list[tuple[str, str, str, str, str]]:
    symbol = _currency_symbol(record["currency"])
    rows = [
        (
            record["fuel"],
            _fmt_num(record["quantity"]),
            record["unit"],
            f"{symbol}{_fmt_money(record['unit_price'])}",
            f"{symbol}{_fmt_money(record['fuel_cost'])}",
        ),
        (
            _tr(raw_config, "delivery_charge"),
            "1",
            _tr(raw_config, "each"),
            f"{symbol}{_fmt_money(record['delivery_charge'])}",
            f"{symbol}{_fmt_money(record['delivery_charge'])}",
        ),
    ]
    for line in record.get("cross_scope_lines", []):
        rows.append((
            str(line.get("description", "")),
            _fmt_num(line.get("quantity")),
            str(line.get("unit", "")),
            f"{symbol}{_fmt_money(line.get('unit_price'))}",
            f"{symbol}{_fmt_money(line.get('amount'))}",
        ))
    return rows


def _site_equipment_items(raw_config: dict, site: dict, *, include_emission_source: bool) -> list[dict]:
    site_omit = site.get("_omit", {})
    equipment_omitted = bool(site_omit.get("equipment", False))
    site_emission_source_omitted = bool(site_omit.get("emission_source", False))
    raw_items = site.get("equipment_items")

    if equipment_omitted:
        raw_items = raw_items if isinstance(raw_items, list) and raw_items else []
        first_item = raw_items[0] if raw_items else {}
        if isinstance(first_item, dict):
            first_emission_source = first_item.get("emission_source", site.get("emission_source", ""))
            first_omit = first_item.get("_omit", {})
        else:
            first_emission_source = site.get("emission_source", "")
            first_omit = {}
        raw_item = dict(first_item) if isinstance(first_item, dict) else {}
        raw_item.update({
            "equipment": "",
            "emission_source": first_emission_source,
            "_omit": first_omit,
        })
        raw_items = [raw_item]

    if not isinstance(raw_items, list) or not raw_items:
        raw_items = [
            {
                "equipment": site.get("equipment", ""),
                "emission_source": site.get("emission_source", ""),
                "_omit": {"emission_source": site_emission_source_omitted},
            }
        ]

    equipment_items: list[dict] = []
    for raw_item in raw_items:
        if isinstance(raw_item, str):
            equipment = raw_item
            emission_source = site.get("emission_source", "")
            item_data: dict = {}
            item_omit: dict = {}
        else:
            equipment = raw_item.get("equipment", raw_item.get("name", ""))
            emission_source = raw_item.get("emission_source", site.get("emission_source", ""))
            item_data = raw_item
            item_omit = raw_item.get("_omit", {})

        emission_source_omitted = site_emission_source_omitted or bool(item_omit.get("emission_source", False))
        equipment_item_omitted = equipment_omitted or bool(item_omit.get("equipment", False))
        normalized_item = {
            "equipment": "" if equipment_item_omitted else _with_special_chars(raw_config, "" if equipment is None else str(equipment)),
            "emission_source": ""
            if not include_emission_source or emission_source_omitted
            else _with_special_chars(raw_config, "" if emission_source is None else str(emission_source)),
        }
        for field in [
            "fuel",
            "unit",
            "quantity",
            "unit_price",
            "delivery_charge",
            "vat_rate",
            "runs_per_month",
            "fuel_used_per_hour",
            "quantity_mode",
            "tank_capacity",
            "run_hours_min",
            "run_hours_max",
        ]:
            if field in item_data:
                normalized_item[field] = item_data.get(field)
            elif field in site:
                normalized_item[field] = site.get(field)
        normalized_item.setdefault("unit", site.get("unit", default_fuel_volume_unit(raw_config.get("document_type"))))
        normalized_item.setdefault("fuel", site.get("fuel", ""))
        equipment_items.append(normalized_item)

    return equipment_items or [{"equipment": "", "emission_source": ""}]


def _build_fuel_invoice_base_records(raw_config: dict) -> list[dict]:
    """One template record per configured company/site/equipment line. These seed
    the generated invoices; invoice numbers and dates are assigned per document."""
    records: list[dict] = []
    seed = int(raw_config.get("random_seed", 42))

    for company_index, site_index, company, site in _iter_company_sites(raw_config):
        site_omit = site.get("_omit", {})
        equipment_items = _site_equipment_items(raw_config, site, include_emission_source=True)

        for equipment_item in equipment_items:
            quantity = _q2(_parse_decimal(equipment_item.get("quantity"), "0"))
            unit_price = _q2(_parse_decimal(equipment_item.get("unit_price"), "0"))
            delivery_charge = _q2(_parse_decimal(equipment_item.get("delivery_charge"), "0"))
            vat_rate = _parse_decimal(equipment_item.get("vat_rate"), "20")
            fuel_cost = _q2(quantity * unit_price)
            subtotal = _q2(fuel_cost + delivery_charge)
            vat = _q2(subtotal * vat_rate / Decimal("100"))
            total = _q2(subtotal + vat)

            records.append({
                "company": _with_special_chars(raw_config, company.get("label", "")),
                "supplier": _with_special_chars(raw_config, company.get("supplier", "")),
                "supplier_code": company.get("supplier_code", "INV"),
                "supplier_address": [
                    _with_special_chars(raw_config, line) for line in company.get("supplier_address", [])
                ],
                "customer": _with_special_chars(raw_config, company.get("customer", "")),
                "customer_code": company.get("customer_code", ""),
                "account_number": _stationary_account_number(company, seed, company_index),
                "site": _with_special_chars(raw_config, site.get("label", "")),
                "site_address": [_with_special_chars(raw_config, line) for line in site.get("customer_address", [])],
                "country": "" if site_omit.get("country", False) else _with_special_chars(raw_config, site.get("country", "")),
                "equipment": equipment_item["equipment"],
                "emission_source": equipment_item["emission_source"],
                "fuel": _with_special_chars(raw_config, equipment_item.get("fuel", "")),
                "unit": equipment_item.get("unit", site.get("unit", _tr(raw_config, "litres"))),
                "quantity": quantity,
                "unit_price": unit_price,
                "fuel_cost": fuel_cost,
                "delivery_charge": delivery_charge,
                "subtotal": subtotal,
                "vat_rate": vat_rate,
                "vat": vat,
                "total": total,
                "currency": company.get("currency", "GBP (£)"),
            })
    return records


def _build_fuel_invoice_records(raw_config: dict) -> list[dict]:
    fp = _financial_period(raw_config)
    seed = int(raw_config.get("random_seed", 42))
    base_records = _build_fuel_invoice_base_records(raw_config)
    if not base_records:
        return []

    count = _multi_document_count(raw_config)
    days_in_period = max((fp["end_date"] - fp["start_date"]).days, 0)
    records: list[dict] = []

    for document_index in range(count):
        template = base_records[document_index % len(base_records)]
        record = dict(template)
        rng = random.Random(f"{seed}:fuel_invoice_doc:{document_index}")

        # Documents beyond the configured lines reuse the company/site identity but
        # get fresh randomized quantities and prices so every invoice is distinct.
        if document_index >= len(base_records):
            quantity = _q2(Decimal(rng.randrange(1200, 6000, 50)))
            unit_price = _q2(Decimal(str(round(rng.uniform(0.88, 1.35), 2))))
            delivery_charge = _q2(Decimal(str(round(rng.uniform(20.0, 95.0), 2))))
            vat_rate = _parse_decimal(record.get("vat_rate"), "20")
            fuel_cost = _q2(quantity * unit_price)
            subtotal = _q2(fuel_cost + delivery_charge)
            vat = _q2(subtotal * vat_rate / Decimal("100"))
            total = _q2(subtotal + vat)
            record.update({
                "quantity": quantity,
                "unit_price": unit_price,
                "delivery_charge": delivery_charge,
                "fuel_cost": fuel_cost,
                "subtotal": subtotal,
                "vat": vat,
                "total": total,
            })

        if _cross_scope_enabled(raw_config):
            cross_scope_line = _cross_scope_vehicle_line(raw_config, rng)
            vat_rate = _parse_decimal(record.get("vat_rate"), "20")
            subtotal = _q2(record["subtotal"] + cross_scope_line["amount"])
            vat = _q2(subtotal * vat_rate / Decimal("100"))
            record.update({
                "cross_scope_lines": [cross_scope_line],
                "subtotal": subtotal,
                "vat": vat,
                "total": _q2(subtotal + vat),
            })

        # One-time purchase dates rather than a billing period.
        invoiced_date = fp["start_date"] + timedelta(days=rng.randint(0, days_in_period))
        received_date = invoiced_date + timedelta(days=rng.randint(0, 4))
        due_date = invoiced_date + timedelta(days=rng.randint(14, 30))
        invoice_no = f"{record.get('supplier_code', 'INV')}-{invoiced_date.strftime('%Y%m')}-{document_index + 1:04d}"

        record.update({
            "invoice_no": invoice_no,
            "invoiced_date": invoiced_date,
            "received_date": received_date,
            "due_date": due_date,
        })
        records.append(record)

    return records


def _build_delivery_note_records(raw_config: dict) -> list[dict]:
    fp = _financial_period(raw_config)
    records: list[dict] = []
    seed = int(raw_config.get("random_seed", 42))
    days_in_period = max((fp["end_date"] - fp["start_date"]).days, 0)

    for company_index, site_index, company, site in _iter_company_sites(raw_config):
        site_omit = site.get("_omit", {})
        equipment_items = _site_equipment_items(raw_config, site, include_emission_source=False)

        for equipment_index, equipment_item in enumerate(equipment_items, start=1):
            rng = random.Random(f"{seed}:delivery_note:{company_index}:{site_index}:{equipment_index}")
            delivery_date = fp["start_date"] + timedelta(days=rng.randint(0, days_in_period))
            note_no = (
                f"DN-{delivery_date.strftime('%Y')}-"
                f"{rng.randint(10000, 99999)}"
            )

            records.append({
                "company": _with_special_chars(raw_config, company.get("label", "")),
                "supplier": _with_special_chars(raw_config, company.get("supplier", "")),
                "supplier_address": [
                    _with_special_chars(raw_config, line) for line in company.get("supplier_address", [])
                ],
                "customer": _with_special_chars(raw_config, company.get("customer", "")),
                "site": _with_special_chars(raw_config, site.get("label", "")),
                "site_address": [_with_special_chars(raw_config, line) for line in site.get("customer_address", [])],
                "country": "" if site_omit.get("country", False) else _with_special_chars(raw_config, site.get("country", "")),
                "equipment": equipment_item["equipment"],
                "fuel": _with_special_chars(raw_config, equipment_item.get("fuel", "")),
                "unit": equipment_item.get("unit", site.get("unit", _tr(raw_config, "litres"))),
                "quantity": _q2(_parse_decimal(equipment_item.get("quantity"), "0")),
                "delivery_note_no": note_no,
                "delivery_date": delivery_date,
                "driver_ref": f"TRK-{rng.randint(1, 24):02d}",
                "customer_signature": _tr(raw_config, "received"),
                "period_label": fp["label"],
                "period_start": fp["start_date"],
                "period_end": fp["end_date"],
            })

    return records


def _build_fuel_card_base_statements(raw_config: dict) -> list[dict]:
    fp = _financial_period(raw_config)
    seed = int(raw_config.get("random_seed", 42))
    days_in_period = max((fp["end_date"] - fp["start_date"]).days, 0)
    statements: list[dict] = []

    for company_index, company in enumerate(raw_config.get("companies", []), start=1):
        transactions: list[dict] = []
        for transaction_index, site in enumerate(company.get("sites", []), start=1):
            site_omit = site.get("_omit", {})
            site_value = "" if site_omit.get("label", False) else _with_special_chars(raw_config, site.get("label", ""))
            country_value = "" if site_omit.get("country", False) else _with_special_chars(raw_config, site.get("country", ""))
            equipment_items = _site_equipment_items(raw_config, site, include_emission_source=True)

            for equipment_index, equipment_item in enumerate(equipment_items, start=1):
                quantity = _q2(_parse_decimal(equipment_item.get("quantity"), "0"))
                unit_price = _q2(_parse_decimal(equipment_item.get("unit_price"), "0"))
                total = _q2(quantity * unit_price)
                rng = random.Random(f"{seed}:fuel_card:{company_index}:{transaction_index}:{equipment_index}")
                transaction_date = fp["start_date"] + timedelta(days=rng.randint(0, days_in_period))
                reference_value = site_value or equipment_item["equipment"] or _tr(raw_config, "stationary_equipment")

                transactions.append({
                    "card_number": company.get("card_number") or site.get("card_number", ""),
                    "date": transaction_date,
                    "merchant": _with_special_chars(raw_config, company.get("merchant") or site.get("merchant", "")),
                    "site": site_value,
                    "country": country_value,
                    "equipment": equipment_item["equipment"],
                    "emission_source": equipment_item["emission_source"],
                    "reference": reference_value,
                    "fuel": _with_special_chars(raw_config, equipment_item.get("fuel", "")),
                    "quantity": quantity,
                    "unit": equipment_item.get("unit", site.get("unit", default_fuel_volume_unit("fuel_card"))),
                    "unit_price": unit_price,
                    "total": total,
                })

        transactions.sort(key=lambda row: (row["date"], row["merchant"], row["card_number"]))
        statements.append({
            "company": _with_special_chars(raw_config, company.get("label", "")),
            "account_name": _with_special_chars(raw_config, company.get("customer") or company.get("label", "")),
            "provider": _with_special_chars(raw_config, company.get("supplier", "")),
            "currency": company.get("currency", "GBP (£)"),
            "period_label": fp["label"],
            "period_start": fp["start_date"],
            "period_end": fp["end_date"],
            "transactions": transactions,
            "statement_total": _q2(sum((row["total"] for row in transactions), Decimal("0"))),
        })

    return statements


def _fuel_card_vat_rate(raw_config: dict) -> Decimal:
    return _parse_decimal(raw_config.get("document", {}).get("fuel_card_vat_rate", 20), "20")


def _apply_fuel_card_totals(statement: dict, vat_rate: Decimal) -> dict:
    """Add Net / VAT / Gross totals to a statement (UK/Ireland-style breakdown)."""
    net_total = _q2(sum((row["total"] for row in statement["transactions"]), Decimal("0")))
    vat_amount = _q2(net_total * vat_rate / Decimal("100"))
    gross_total = _q2(net_total + vat_amount)
    statement["net_total"] = net_total
    statement["vat_rate"] = vat_rate
    statement["vat_amount"] = vat_amount
    statement["gross_total"] = gross_total
    statement["statement_total"] = gross_total
    return statement


def _fuel_card_summary_rows(raw_config: dict, statement: dict) -> list[tuple[str, Any, bool]]:
    """(label, amount, is_total) rows for the Net / VAT / Gross statement summary."""
    return [
        (_tr(raw_config, "net_amount"), statement["net_total"], False),
        (f"{_tr(raw_config, 'vat')} ({statement['vat_rate']}%)", statement["vat_amount"], False),
        (_tr(raw_config, "gross_amount"), statement["gross_total"], True),
    ]


def _add_fuel_card_totals_docx(document, raw_config: dict, statement: dict, currency_symbol: str) -> None:
    rows = _fuel_card_summary_rows(raw_config, statement)
    totals = document.add_table(rows=len(rows), cols=2)
    totals.style = "Table Grid"
    for row_idx, (label, amount, is_total) in enumerate(rows):
        _shade_docx_cell(totals.cell(row_idx, 0), "F5F8FB")
        _set_docx_cell_text(totals.cell(row_idx, 0), label, bold=True)
        _set_docx_cell_text(totals.cell(row_idx, 1), f"{currency_symbol}{_fmt_money(amount)}", bold=is_total)


def _expand_fuel_card_transactions(statement: dict, count: int, rng: random.Random) -> list[dict]:
    """Produce exactly `count` transaction lines for a statement. The configured
    transactions seed the first lines; extra lines reuse them with randomized
    dates, quantities, and prices, all within the statement period."""
    base_transactions = statement["transactions"]
    if not base_transactions:
        return []

    fp_start = statement["period_start"]
    days_in_period = max((statement["period_end"] - fp_start).days, 0)
    transactions: list[dict] = []
    for line_index in range(count):
        template = base_transactions[line_index % len(base_transactions)]
        transaction = dict(template)
        if line_index >= len(base_transactions):
            quantity = _q2(Decimal(rng.randrange(150, 700, 10)))
            unit_price = _q2(Decimal(str(round(rng.uniform(0.95, 1.45), 2))))
            transaction.update({
                "date": fp_start + timedelta(days=rng.randint(0, days_in_period)),
                "quantity": quantity,
                "unit_price": unit_price,
                "total": _q2(quantity * unit_price),
            })
        transactions.append(transaction)
    transactions.sort(key=lambda row: (row["date"], row["merchant"], row["card_number"]))
    return transactions


def _build_fuel_card_statements(raw_config: dict) -> list[dict]:
    seed = int(raw_config.get("random_seed", 42))
    base_statements = _build_fuel_card_base_statements(raw_config)
    if not base_statements:
        return []

    line_item_count = _multi_document_count(raw_config)
    vat_rate = _fuel_card_vat_rate(raw_config)
    statements: list[dict] = []
    for statement_index, statement in enumerate(base_statements):
        new_statement = dict(statement)
        new_statement["transactions"] = _expand_fuel_card_transactions(
            statement,
            line_item_count,
            random.Random(f"{seed}:fuel_card_lines:{statement_index}"),
        )
        statements.append(_apply_fuel_card_totals(new_statement, vat_rate))
    return statements


def _date_within_month(year: int, month: int, day: int) -> date:
    return date(year, month, min(day, monthrange(year, month)[1]))


def _fmt_percent(value: float) -> str:
    return f"{round(value):.0f}%"


def _format_time(minutes_total: int) -> str:
    hours = (minutes_total // 60) % 24
    minutes = minutes_total % 60
    return f"{hours:02d}:{minutes:02d}"


def _build_generator_log_rows(raw_config: dict) -> list[dict]:
    fp = _financial_period(raw_config)
    months = _months_in_range(fp["start_date"], fp["end_date"])
    rows: list[dict] = []
    seed = int(raw_config.get("random_seed", 42))

    for company_index, site_index, company, site in _iter_company_sites(raw_config):
        site_omit = site.get("_omit", {})
        equipment_items = _site_equipment_items(raw_config, site, include_emission_source=True)

        for equipment_index, equipment_item in enumerate(equipment_items, start=1):
            runs_per_month = max(int(equipment_item.get("runs_per_month", 3)), 1)
            tank_capacity = float(equipment_item.get("tank_capacity", 800))
            burn_rate = float(equipment_item.get("fuel_used_per_hour", 15))
            min_hours = float(equipment_item.get("run_hours_min", 0.5))
            max_hours = max(float(equipment_item.get("run_hours_max", 4.0)), min_hours)
            quantity_mode = equipment_item.get("quantity_mode", "tank_level_change")
            rng = random.Random(f"{seed}:generator_log:{company_index}:{site_index}:{equipment_index}")
            for year, month in months:
                days_in_month = monthrange(year, month)[1]
                chosen_days = sorted(rng.sample(range(1, days_in_month + 1), k=min(runs_per_month, days_in_month)))
                for day in chosen_days:
                    run_date = _date_within_month(year, month, day)
                    start_minutes = rng.choice([7 * 60, 8 * 60, 9 * 60, 13 * 60, 18 * 60])
                    run_hours = round(rng.uniform(min_hours, max_hours), 2)
                    end_minutes = start_minutes + int(run_hours * 60)

                    if quantity_mode == "explicit_fuel_used":
                        fuel_used = round(run_hours * burn_rate * rng.uniform(0.92, 1.12), 2)
                        start_pct = float(rng.randint(52, 92))
                        end_pct = max(0.0, start_pct - ((fuel_used / max(tank_capacity, 1.0)) * 100.0))
                    else:
                        start_pct = float(rng.randint(55, 95))
                        estimated_delta = max((run_hours * burn_rate / max(tank_capacity, 1.0)) * 100.0, 1.0)
                        delta_pct = min(start_pct, estimated_delta * rng.uniform(0.9, 1.15))
                        end_pct = max(0.0, start_pct - delta_pct)
                        fuel_used = round((start_pct - end_pct) / 100.0 * tank_capacity, 2)

                    rows.append({
                        "company": _with_special_chars(raw_config, company.get("label", "")),
                        "site": _with_special_chars(raw_config, site.get("label", "")),
                        "country": "" if site_omit.get("country", False) else _with_special_chars(raw_config, site.get("country", "")),
                        "equipment": equipment_item["equipment"],
                        "emission_source": equipment_item["emission_source"],
                        "fuel": _with_special_chars(raw_config, equipment_item.get("fuel", "")),
                        "period": run_date.isoformat(),
                        "date": run_date,
                        "start_time": _format_time(start_minutes),
                        "end_time": _format_time(end_minutes),
                        "run_hours": run_hours,
                        "start_fuel": _fmt_percent(start_pct),
                        "end_fuel": _fmt_percent(end_pct),
                        "fuel_used": round(fuel_used, 2),
                        "unit": equipment_item.get("unit", site.get("unit", default_fuel_volume_unit("generator_log"))),
                        "notes": _with_special_chars(
                            raw_config,
                            rng.choice(
                                [
                                    _tr(raw_config, "test_run"),
                                    _tr(raw_config, "power_outage"),
                                    _tr(raw_config, "maintenance_test"),
                                    _tr(raw_config, "load_bank_test"),
                                ]
                            ),
                        ),
                    })
    return sorted(rows, key=lambda row: (row["site"], row["equipment"], row["date"], row["start_time"]))


def _build_bems_site_blocks(raw_config: dict) -> list[dict]:
    fp = _financial_period(raw_config)
    blocks: list[dict] = []

    for _, _, company, site in _iter_company_sites(raw_config):
        site_omit = site.get("_omit", {})
        assets: list[dict] = []
        for asset in site.get("assets", []):
            asset_omit = asset.get("_omit", {})
            assets.append({
                "asset_tag": _with_special_chars(raw_config, asset.get("asset_tag", "")),
                "equipment_name": "" if asset_omit.get("equipment_name", False) else _with_special_chars(raw_config, asset.get("equipment_name", "")),
                "emission_source": "" if asset_omit.get("emission_source", False) else _with_special_chars(raw_config, asset.get("emission_source", "")),
                "fuel": "" if asset_omit.get("fuel", False) else _with_special_chars(raw_config, asset.get("fuel", "")),
                "unit": asset.get("unit", "kWh"),
                "sensor_name": "" if asset_omit.get("sensor_name", False) else _with_special_chars(raw_config, asset.get("sensor_name", "")),
                "quantity": _q2(_parse_decimal(asset.get("quantity"), "0")),
                "operating_hours": None if asset_omit.get("operating_hours", False) else _q2(_parse_decimal(asset.get("operating_hours"), "0")),
                "_omit": asset_omit,
            })
        blocks.append({
            "company": _with_special_chars(raw_config, company.get("label", "")),
            "site": _with_special_chars(raw_config, site.get("label", "")),
            "country": "" if site_omit.get("country", False) else _with_special_chars(raw_config, site.get("country", "")),
            "period_label": fp["label"],
            "period_start": fp["start_date"],
            "period_end": fp["end_date"],
            "assets": assets,
            "_omit": site_omit,
        })
    return blocks


def _bems_asset_weights(asset: dict, timestamps: list[datetime], rng: random.Random) -> list[float]:
    emission_source = asset.get("emission_source", "").lower()
    equipment = asset.get("equipment_name", "").lower()

    if "generator" in emission_source or "generator" in equipment:
        active_count = max(1, min(len(timestamps), max(4, len(timestamps) // 40)))
        active_indices = set(rng.sample(range(len(timestamps)), active_count))
        weights: list[float] = []
        for idx, ts in enumerate(timestamps):
            if idx not in active_indices:
                weights.append(0.0)
                continue
            base = 1.0 if 8 <= ts.hour <= 18 else 0.6
            weights.append(base * rng.uniform(0.9, 1.15))
        return weights

    weights = []
    for ts in timestamps:
        hour = ts.hour + ts.minute / 60
        base = 0.55
        if 5 <= hour < 8:
            base = 0.95
        elif 8 <= hour < 18:
            base = 1.25
        elif 18 <= hour < 22:
            base = 0.88
        if ts.weekday() >= 5:
            base *= 0.82
        weights.append(base * rng.uniform(0.92, 1.08))
    return weights


def _distribute_bems_series(total_quantity: Decimal, weights: list[float]) -> list[Decimal]:
    if not weights:
        return []
    weight_total = sum(weights)
    if weight_total <= 0:
        even_value = _q2(total_quantity / Decimal(str(len(weights))))
        series = [even_value for _ in weights]
        if series:
            series[-1] = _q2(total_quantity - sum(series[:-1]))
        return series

    series = [
        _q2(total_quantity * Decimal(str(weight / weight_total)))
        for weight in weights
    ]
    series[-1] = _q2(total_quantity - sum(series[:-1]))
    return series


def _build_bems_trend_exports(raw_config: dict) -> list[dict]:
    fp = _financial_period(raw_config)
    interval_minutes = _bems_interval_minutes(raw_config)
    timestamps = _timestamp_range(fp["start_date"], fp["end_date"], interval_minutes)
    seed = int(raw_config.get("random_seed", 42))
    exports: list[dict] = []

    for block_index, block in enumerate(_build_bems_site_blocks(raw_config), start=1):
        rows: list[dict] = []
        for asset_index, asset in enumerate(block["assets"], start=1):
            rng = random.Random(f"{seed}:bems:{block_index}:{asset_index}:{asset['asset_tag']}")
            weights = _bems_asset_weights(asset, timestamps, rng)
            values = _distribute_bems_series(asset["quantity"], weights)
            for timestamp, value in zip(timestamps, values):
                rows.append({
                    "timestamp": timestamp,
                    "site": block["site"],
                    "asset_tag": asset["asset_tag"],
                    "sensor_name": asset["sensor_name"],
                    "value": float(value),
                    "unit": asset["unit"],
                })
        exports.append({
            "company": block["company"],
            "site": block["site"],
            "country": block["country"],
            "period_label": block["period_label"],
            "period_start": block["period_start"],
            "period_end": block["period_end"],
            "rows": rows,
            "assets": block["assets"],
        })
    return exports


# ── corruption wrappers (call _build_* then apply bad_data) ─────────────────

def _corrupted_fuel_invoice_records(raw_config: dict) -> list[dict]:
    cfg = get_bad_data_config(raw_config)
    return corrupt_records(_build_fuel_invoice_records(raw_config), "fuel_invoice", _FUEL_INVOICE_FIELD_TYPES, cfg)


def _corrupted_delivery_note_records(raw_config: dict) -> list[dict]:
    cfg = get_bad_data_config(raw_config)
    return corrupt_records(_build_delivery_note_records(raw_config), "delivery_note", _DELIVERY_NOTE_FIELD_TYPES, cfg)


def _corrupted_fuel_card_statements(raw_config: dict) -> list[dict]:
    cfg = get_bad_data_config(raw_config)
    statements = _build_fuel_card_statements(raw_config)
    result = []
    for s_idx, statement in enumerate(statements):
        new_statement = dict(statement)
        new_statement["transactions"] = corrupt_records(
            statement["transactions"], f"fuel_card:s{s_idx}", _FUEL_CARD_TRANSACTION_FIELD_TYPES, cfg
        )
        result.append(new_statement)
    return result


def _corrupted_generator_log_rows(raw_config: dict) -> list[dict]:
    cfg = get_bad_data_config(raw_config)
    return corrupt_records(_build_generator_log_rows(raw_config), "generator_log", _GENERATOR_LOG_FIELD_TYPES, cfg)


def _corrupted_bems_site_blocks(raw_config: dict) -> list[dict]:
    """Corrupt assets inside each block; used by equipment-report generators only."""
    cfg = get_bad_data_config(raw_config)
    blocks = _build_bems_site_blocks(raw_config)
    result = []
    for b_idx, block in enumerate(blocks):
        new_block = dict(block)
        new_block["assets"] = corrupt_records(block["assets"], f"bems:b{b_idx}", _BEMS_ASSET_FIELD_TYPES, cfg)
        result.append(new_block)
    return result


def _safe_float(value, fallback: float = 0.0) -> float:
    """Return float(value) or fallback if value is a non-numeric string."""
    if isinstance(value, str):
        return fallback
    try:
        return float(value)
    except Exception:
        return fallback


def _ground_truth_entries(raw_config: dict) -> list[dict]:
    document_type = raw_config.get("document_type", "fuel_invoice")
    cfg = get_bad_data_config(raw_config)

    if document_type == "fuel_invoice":
        records = corrupt_records(
            _build_fuel_invoice_records(raw_config), "fuel_invoice", _FUEL_INVOICE_FIELD_TYPES, cfg
        )
        entries: list[dict] = []
        for record in records:
            currency = record["currency"].split()[0] if isinstance(record["currency"], str) and " " in record["currency"] else record["currency"]
            entries.append({
                "Scope": "Scope 1",
                "Category": "stationary_combustion",
                "Company": record["company"],
                "Site": record["site"],
                "Country": record["country"],
                "Invoice No": record["invoice_no"],
                "Account Number": record.get("account_number", ""),
                "Invoiced date": record["invoiced_date"].isoformat() if hasattr(record["invoiced_date"], "isoformat") else record["invoiced_date"],
                "Received date": record["received_date"].isoformat() if hasattr(record["received_date"], "isoformat") else record["received_date"],
                "Due date": record["due_date"].isoformat() if hasattr(record["due_date"], "isoformat") else record["due_date"],
                "Equipment": record["equipment"],
                "Emission source": record["emission_source"],
                "Fuel": record["fuel"],
                "Quantity": record["quantity"] if isinstance(record["quantity"], str) else float(record["quantity"]),
                "Unit": record["unit"],
                "Cost": record["fuel_cost"] if isinstance(record["fuel_cost"], str) else float(record["fuel_cost"]),
                "Currency": currency,
            })
            for line in record.get("cross_scope_lines", []):
                entries.append({
                    "Scope": "Scope 1",
                    "Category": line.get("category", "mobile_combustion"),
                    "Company": record["company"],
                    "Invoice No": record["invoice_no"],
                    "Account Number": record.get("account_number", ""),
                    "Invoiced date": record["invoiced_date"].isoformat() if hasattr(record["invoiced_date"], "isoformat") else record["invoiced_date"],
                    "Vehicle": line.get("vehicle_reg", ""),
                    "Fuel": line.get("fuel", ""),
                    "Quantity": float(line["quantity"]),
                    "Unit": line.get("unit", ""),
                    "Cost": float(line["amount"]),
                    "Currency": currency,
                })
        return entries
    if document_type == "delivery_note":
        records = corrupt_records(
            _build_delivery_note_records(raw_config), "delivery_note", _DELIVERY_NOTE_FIELD_TYPES, cfg
        )
        return [
            {
                "Scope": "Scope 1",
                "Category": "stationary_combustion",
                "Company": record["company"],
                "Site": record["site"],
                "Country": record["country"],
                "Period": record["delivery_date"].isoformat() if hasattr(record["delivery_date"], "isoformat") else record["delivery_date"],
                "Equipment": record["equipment"],
                "Fuel": record["fuel"],
                "Quantity": record["quantity"] if isinstance(record["quantity"], str) else float(record["quantity"]),
                "Unit": record["unit"],
            }
            for record in records
        ]
    if document_type == "fuel_card":
        statements = _build_fuel_card_statements(raw_config)
        corrupted_statements = []
        for s_idx, statement in enumerate(statements):
            new_statement = dict(statement)
            new_statement["transactions"] = corrupt_records(
                statement["transactions"], f"fuel_card:s{s_idx}", _FUEL_CARD_TRANSACTION_FIELD_TYPES, cfg
            )
            corrupted_statements.append(new_statement)
        return [
            {
                "Scope": "Scope 1",
                "Category": "stationary_combustion",
                "Company": statement["company"],
                "Site": transaction["site"],
                "Country": transaction["country"],
                "Period": f"{statement['period_start'].isoformat() if hasattr(statement['period_start'], 'isoformat') else statement['period_start']} to {statement['period_end'].isoformat() if hasattr(statement['period_end'], 'isoformat') else statement['period_end']}",
                "Equipment": transaction["equipment"],
                "Emission source": transaction["emission_source"],
                "Fuel": transaction["fuel"],
                "Quantity": transaction["quantity"] if isinstance(transaction["quantity"], str) else float(transaction["quantity"]),
                "Unit": transaction["unit"],
                "Cost": transaction["total"] if isinstance(transaction["total"], str) else float(transaction["total"]),
                "Currency": statement["currency"].split()[0] if isinstance(statement["currency"], str) and " " in statement["currency"] else statement["currency"],
            }
            for statement in corrupted_statements
            for transaction in statement["transactions"]
        ]
    if document_type == "bems":
        blocks = _build_bems_site_blocks(raw_config)
        corrupted_blocks = []
        for b_idx, block in enumerate(blocks):
            new_block = dict(block)
            new_block["assets"] = corrupt_records(
                block["assets"], f"bems:b{b_idx}", _BEMS_ASSET_FIELD_TYPES, cfg
            )
            corrupted_blocks.append(new_block)
        return [
            {
                "Scope": "Scope 1",
                "Category": "stationary_combustion",
                "Company": block["company"],
                "Site": block["site"],
                "Country": block["country"],
                "Period": f"{block['period_start'].isoformat() if hasattr(block['period_start'], 'isoformat') else block['period_start']} to {block['period_end'].isoformat() if hasattr(block['period_end'], 'isoformat') else block['period_end']}",
                "Equipment": asset["equipment_name"],
                "Emission source": asset["emission_source"],
                "Fuel": asset["fuel"],
                "Quantity": asset["quantity"] if isinstance(asset["quantity"], str) else float(asset["quantity"]),
                "Unit": asset["unit"],
            }
            for block in corrupted_blocks
            for asset in block["assets"]
        ]
    rows = corrupt_records(
        _build_generator_log_rows(raw_config), "generator_log", _GENERATOR_LOG_FIELD_TYPES, cfg
    )
    return [
        {
            "Scope": "Scope 1",
            "Category": "stationary_combustion",
            "Site": row["site"],
            "Country": row["country"],
            "Period": row["period"],
            "Equipment": row["equipment"],
            "Emission source": row["emission_source"],
            "Fuel": row["fuel"],
            "Quantity": row["fuel_used"] if isinstance(row["fuel_used"], str) else row["fuel_used"],
            "Unit": row["unit"],
        }
        for row in rows
    ]


def generate_ground_truth_json(raw_config: dict) -> bytes:
    import json

    return json.dumps(_ground_truth_entries(raw_config), indent=2).encode("utf-8")


def _draw_multiline(c: canvas.Canvas, x: float, y: float, lines: list[str], leading: int = 12) -> float:
    for line in lines:
        c.drawString(x, y, line)
        y -= leading
    return y


def _shade_docx_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _style_docx_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(9.5)


def _set_docx_cell_text(cell, text: str, *, bold: bool = False, color: str | None = None, size: float | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size is not None:
        run.font.size = Pt(size)


def generate_fuel_invoice_pdf(raw_config: dict) -> bytes:
    records = _corrupted_fuel_invoice_records(raw_config)
    if _is_multi_document(raw_config):
        return _zip_documents(raw_config, records, _render_fuel_invoice_pdf, _fuel_invoice_filename, "pdf")
    return _render_fuel_invoice_pdf(raw_config, records)


def _render_fuel_invoice_pdf(raw_config: dict, records: list[dict]) -> bytes:
    distractor_plan = _stationary_distractor_plan(raw_config, "PDF", document_type="fuel_invoice")
    layout_plan = _stationary_layout_plan(raw_config, "PDF")
    if layout_plan.get("enabled"):
        return _generate_fuel_invoice_pdf_variant(raw_config, records, layout_plan, distractor_plan)

    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    c.setTitle(raw_config.get("document", {}).get("title", _tr(raw_config, "fuel_invoice_title")))
    c.setSubject(raw_config.get("document", {}).get("subject", "Scope 1 stationary combustion"))

    for index, record in enumerate(records):
        if index > 0:
            c.showPage()

        accent = colors.HexColor("#1E5B88")
        c.setFillColor(accent)
        c.rect(36, PAGE_H - 72, PAGE_W - 72, 28, fill=1, stroke=0)
        c.setFillColor(colors.white)
        c.setFont("Helvetica-Bold", 14)
        c.drawString(48, PAGE_H - 62, record["supplier"])

        c.setFillColor(colors.black)
        c.setFont("Helvetica", 10)
        y = PAGE_H - 108
        y = _draw_multiline(c, 48, y, record["supplier_address"])

        c.setFont("Helvetica-Bold", 11)
        c.drawString(48, y - 10, _tr(raw_config, "bill_to"))
        c.setFont("Helvetica", 10)
        bill_to_lines = [record["customer"], *record["site_address"]]
        _draw_multiline(c, 48, y - 26, bill_to_lines)

        c.setFont("Helvetica-Bold", 11)
        c.drawString(320, PAGE_H - 108, _tr(raw_config, "invoice_details"))
        c.setFont("Helvetica", 10)
        _draw_multiline(c, 320, PAGE_H - 126, _fuel_invoice_meta_lines(raw_config, record, distractor_plan))

        c.setFont("Helvetica-Bold", 11)
        c.drawString(320, PAGE_H - 236, _tr(raw_config, "delivery_site"))
        c.setFont("Helvetica", 10)
        delivery_lines = [
            record["site"],
            record["equipment"],
            record["emission_source"],
            *_stationary_document_lines(distractor_plan, placement="summary"),
        ]
        _draw_multiline(c, 320, PAGE_H - 254, [line for line in delivery_lines if line])

        table_top = PAGE_H - 330
        table_x = 48
        table_widths = [210, 68, 58, 84, 84]
        headers = [
            _tr(raw_config, "product"),
            _tr(raw_config, "quantity"),
            _tr(raw_config, "unit"),
            _tr(raw_config, "unit_price"),
            _tr(raw_config, "amount"),
        ]
        c.setFillColor(accent)
        c.rect(table_x, table_top, sum(table_widths), 22, fill=1, stroke=0)
        c.setFillColor(colors.white)
        c.setFont("Helvetica-Bold", 9)

        x_cursor = table_x + 6
        for header, width in zip(headers, table_widths):
            c.drawString(x_cursor, table_top + 7, header)
            x_cursor += width

        rows = _fuel_invoice_line_rows(raw_config, record)

        y_row = table_top - 24
        c.setFont("Helvetica", 9)
        c.setFillColor(colors.black)
        for row in rows:
            c.rect(table_x, y_row, sum(table_widths), 20, fill=0, stroke=1)
            x_cursor = table_x + 6
            for value, width in zip(row, table_widths):
                c.drawString(x_cursor, y_row + 6, str(value))
                x_cursor += width
            y_row -= 20

        summary_y = y_row - 20
        summary_lines = [
            (_tr(raw_config, "subtotal"), record["subtotal"], False),
            (f"VAT ({record['vat_rate']}%)", record["vat"], False),
            (_tr(raw_config, "total"), record["total"], True),
        ]
        for label, value, is_total in summary_lines:
            c.setFont("Helvetica-Bold" if is_total else "Helvetica", 10)
            c.drawRightString(PAGE_W - 180, summary_y, label)
            c.drawRightString(
                PAGE_W - 48,
                summary_y,
                f"{_currency_symbol(record['currency'])}{_fmt_money(value)}",
            )
            summary_y -= 18

        c.setFont("Helvetica", 8)
        c.setFillColor(colors.grey)
        c.drawString(
            48,
            42,
            _tr(raw_config, "fuel_invoice_footer"),
        )

    c.save()
    return buffer.getvalue()


def _generate_fuel_invoice_pdf_variant(raw_config: dict, records: list[dict], layout_plan: dict, distractor_plan) -> bytes:
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    c.setTitle(raw_config.get("document", {}).get("title", _tr(raw_config, "fuel_invoice_title")))
    c.setSubject(raw_config.get("document", {}).get("subject", "Scope 1 stationary combustion"))

    for index, record in enumerate(records):
        if index > 0:
            c.showPage()

        accent = colors.HexColor("#1E5B88")
        c.setFillColor(accent)
        c.rect(36, PAGE_H - 72, PAGE_W - 72, 28, fill=1, stroke=0)
        c.setFillColor(colors.white)
        c.setFont("Helvetica-Bold", 14)
        c.drawString(48, PAGE_H - 62, record["supplier"])

        current_top = PAGE_H - 108
        section_order = [section for section in (layout_plan.get("section_order") or ["addresses", "meta", "line_items", "totals", "footer"]) if section in {"addresses", "meta", "line_items", "totals"}]
        currency_symbol = _currency_symbol(record["currency"])

        for section_name in section_order:
            if section_name == "addresses":
                c.setFillColor(colors.black)
                c.setFont("Helvetica", 10)
                current_top = _draw_multiline(c, 48, current_top, record["supplier_address"]) - 10
                c.setFont("Helvetica-Bold", 11)
                c.drawString(48, current_top, _tr(raw_config, "bill_to"))
                c.setFont("Helvetica", 10)
                current_top = _draw_multiline(c, 48, current_top - 16, [record["customer"], *record["site_address"]]) - 18
            elif section_name == "meta":
                c.setFont("Helvetica-Bold", 11)
                c.drawString(320, current_top, _tr(raw_config, "invoice_details"))
                c.setFont("Helvetica", 10)
                meta_bottom = _draw_multiline(c, 320, current_top - 18, _fuel_invoice_meta_lines(raw_config, record, distractor_plan))
                c.setFont("Helvetica-Bold", 11)
                c.drawString(48, meta_bottom - 6, _tr(raw_config, "delivery_site"))
                c.setFont("Helvetica", 10)
                current_top = _draw_multiline(
                    c,
                    48,
                    meta_bottom - 22,
                    [
                        line
                        for line in [
                            record["site"],
                            record["equipment"],
                            record["emission_source"],
                            *_stationary_document_lines(distractor_plan, placement="summary"),
                        ]
                        if line
                    ],
                ) - 18
            elif section_name == "line_items":
                table_top = current_top - 8
                table_x = 48
                table_widths = [210, 68, 58, 84, 84]
                headers = [_tr(raw_config, "product"), _tr(raw_config, "quantity"), _tr(raw_config, "unit"), _tr(raw_config, "unit_price"), _tr(raw_config, "amount")]
                c.setFillColor(accent)
                c.rect(table_x, table_top, sum(table_widths), 22, fill=1, stroke=0)
                c.setFillColor(colors.white)
                c.setFont("Helvetica-Bold", 9)
                x_cursor = table_x + 6
                for header, width in zip(headers, table_widths):
                    c.drawString(x_cursor, table_top + 7, header)
                    x_cursor += width
                rows = _fuel_invoice_line_rows(raw_config, record)
                y_row = table_top - 24
                c.setFont("Helvetica", 9)
                c.setFillColor(colors.black)
                for row in rows:
                    c.rect(table_x, y_row, sum(table_widths), 20, fill=0, stroke=1)
                    x_cursor = table_x + 6
                    for value, width in zip(row, table_widths):
                        c.drawString(x_cursor, y_row + 6, str(value))
                        x_cursor += width
                    y_row -= 20
                current_top = y_row - 14
            elif section_name == "totals":
                summary_y = current_top
                for label, value, is_total in [(_tr(raw_config, "subtotal"), record["subtotal"], False), (f"VAT ({record['vat_rate']}%)", record["vat"], False), (_tr(raw_config, "total"), record["total"], True)]:
                    c.setFont("Helvetica-Bold" if is_total else "Helvetica", 10)
                    c.drawRightString(PAGE_W - 180, summary_y, label)
                    c.drawRightString(PAGE_W - 48, summary_y, f"{currency_symbol}{_fmt_money(value)}")
                    summary_y -= 18
                current_top = summary_y - 10

        c.setFont("Helvetica", 8)
        c.setFillColor(colors.grey)
        c.drawString(48, 42, _tr(raw_config, "fuel_invoice_footer"))

    c.save()
    return buffer.getvalue()


def generate_fuel_invoice_docx(raw_config: dict) -> bytes:
    records = _corrupted_fuel_invoice_records(raw_config)
    if _is_multi_document(raw_config):
        return _zip_documents(raw_config, records, _render_fuel_invoice_docx, _fuel_invoice_filename, "docx")
    return _render_fuel_invoice_docx(raw_config, records)


def _render_fuel_invoice_docx(raw_config: dict, records: list[dict]) -> bytes:
    distractor_plan = _stationary_distractor_plan(raw_config, "DOCX", document_type="fuel_invoice")
    layout_plan = _stationary_layout_plan(raw_config, "DOCX")
    if layout_plan.get("enabled"):
        return _generate_fuel_invoice_docx_variant(raw_config, records, layout_plan, distractor_plan)

    document = Document()
    document.core_properties.title = raw_config.get("document", {}).get("title", _tr(raw_config, "fuel_invoice_title"))
    document.core_properties.subject = raw_config.get("document", {}).get("subject", "Scope 1 stationary combustion")

    for index, record in enumerate(records):
        document.add_heading(record["supplier"], level=1)

        top_table = document.add_table(rows=2, cols=2)
        top_table.style = "Table Grid"
        top_table.cell(0, 0).text = _tr(raw_config, "invoice_details")
        top_table.cell(0, 1).text = _tr(raw_config, "delivery_site")
        top_table.cell(1, 0).text = "\n".join(_fuel_invoice_meta_lines(raw_config, record, distractor_plan))
        top_table.cell(1, 1).text = "\n".join(
            [
                line
                for line in [
                    record["site"],
                    record["equipment"],
                    record["emission_source"],
                    *_stationary_document_lines(distractor_plan, placement="summary"),
                ]
                if line
            ]
        )

        bill_to_heading = document.add_paragraph()
        bill_to_heading.add_run(_tr(raw_config, "bill_to")).bold = True
        document.add_paragraph("\n".join([record["customer"], *record["site_address"]]))

        line_table = document.add_table(rows=1, cols=5)
        line_table.style = "Table Grid"
        for cell, header in zip(
            line_table.rows[0].cells,
            [
                _tr(raw_config, "product"),
                _tr(raw_config, "quantity"),
                _tr(raw_config, "unit"),
                _tr(raw_config, "unit_price"),
                _tr(raw_config, "amount"),
            ],
        ):
            cell.text = header

        line_rows = _fuel_invoice_line_rows(raw_config, record)
        for values in line_rows:
            row = line_table.add_row().cells
            for cell, value in zip(row, values):
                cell.text = str(value)

        totals = document.add_table(rows=3, cols=2)
        totals.style = "Table Grid"
        totals.cell(0, 0).text = _tr(raw_config, "subtotal")
        totals.cell(0, 1).text = f"{_currency_symbol(record['currency'])}{_fmt_money(record['subtotal'])}"
        totals.cell(1, 0).text = f"VAT ({record['vat_rate']}%)"
        totals.cell(1, 1).text = f"{_currency_symbol(record['currency'])}{_fmt_money(record['vat'])}"
        totals.cell(2, 0).text = _tr(raw_config, "total")
        totals.cell(2, 1).text = f"{_currency_symbol(record['currency'])}{_fmt_money(record['total'])}"

        if index < len(records) - 1:
            document.add_page_break()

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _generate_fuel_invoice_docx_variant(raw_config: dict, records: list[dict], layout_plan: dict, distractor_plan) -> bytes:
    document = Document()
    document.core_properties.title = raw_config.get("document", {}).get("title", _tr(raw_config, "fuel_invoice_title"))
    document.core_properties.subject = raw_config.get("document", {}).get("subject", "Scope 1 stationary combustion")

    for index, record in enumerate(records):
        document.add_heading(record["supplier"], level=1)

        def render_addresses() -> None:
            document.add_paragraph("\n".join(record["supplier_address"]))
            heading = document.add_paragraph()
            heading.add_run(_tr(raw_config, "bill_to")).bold = True
            document.add_paragraph("\n".join([record["customer"], *record["site_address"]]))

        def render_meta() -> None:
            top_table = document.add_table(rows=2, cols=2)
            top_table.style = "Table Grid"
            top_table.cell(0, 0).text = _tr(raw_config, "invoice_details")
            top_table.cell(0, 1).text = _tr(raw_config, "delivery_site")
            top_table.cell(1, 0).text = "\n".join(_fuel_invoice_meta_lines(raw_config, record, distractor_plan))
            top_table.cell(1, 1).text = "\n".join([
                line
                for line in [
                    record["site"],
                    record["equipment"],
                    record["emission_source"],
                    *_stationary_document_lines(distractor_plan, placement="summary"),
                ]
                if line
            ])

        def render_line_items() -> None:
            line_table = document.add_table(rows=1, cols=5)
            line_table.style = "Table Grid"
            for cell, header in zip(line_table.rows[0].cells, [_tr(raw_config, "product"), _tr(raw_config, "quantity"), _tr(raw_config, "unit"), _tr(raw_config, "unit_price"), _tr(raw_config, "amount")]):
                cell.text = header
            for values in _fuel_invoice_line_rows(raw_config, record):
                row = line_table.add_row().cells
                for cell, value in zip(row, values):
                    cell.text = str(value)

        def render_totals() -> None:
            totals = document.add_table(rows=3, cols=2)
            totals.style = "Table Grid"
            totals.cell(0, 0).text = _tr(raw_config, "subtotal")
            totals.cell(0, 1).text = f"{_currency_symbol(record['currency'])}{_fmt_money(record['subtotal'])}"
            totals.cell(1, 0).text = f"VAT ({record['vat_rate']}%)"
            totals.cell(1, 1).text = f"{_currency_symbol(record['currency'])}{_fmt_money(record['vat'])}"
            totals.cell(2, 0).text = _tr(raw_config, "total")
            totals.cell(2, 1).text = f"{_currency_symbol(record['currency'])}{_fmt_money(record['total'])}"

        def render_footer() -> None:
            document.add_paragraph(_tr(raw_config, "fuel_invoice_footer"))

        for section_name in (layout_plan.get("section_order") or ["addresses", "meta", "line_items", "totals", "footer"]):
            if section_name == "addresses":
                render_addresses()
            elif section_name == "meta":
                render_meta()
            elif section_name == "line_items":
                render_line_items()
            elif section_name == "totals":
                render_totals()
            elif section_name == "footer":
                render_footer()

        if index < len(records) - 1:
            document.add_page_break()

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def generate_delivery_note_pdf(raw_config: dict) -> bytes:
    records = _corrupted_delivery_note_records(raw_config)
    distractor_plan = _stationary_distractor_plan(raw_config, "PDF", document_type="delivery_note")
    layout_plan = _stationary_layout_plan(raw_config, "PDF")
    if layout_plan.get("enabled"):
        return _generate_delivery_note_pdf_variant(raw_config, records, layout_plan, distractor_plan)

    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    c.setTitle(raw_config.get("document", {}).get("title", _tr(raw_config, "delivery_note_title")))
    c.setSubject(raw_config.get("document", {}).get("subject", "Scope 1 stationary combustion fuel delivery"))

    for index, record in enumerate(records):
        if index > 0:
            c.showPage()

        accent = colors.HexColor("#1E5B88")
        c.setFillColor(accent)
        c.rect(36, PAGE_H - 72, PAGE_W - 72, 30, fill=1, stroke=0)
        c.setFillColor(colors.white)
        c.setFont("Helvetica-Bold", 15)
        c.drawString(48, PAGE_H - 62, _tr(raw_config, "delivery_note_title"))

        c.setFillColor(colors.black)
        c.setFont("Helvetica", 10)
        y = PAGE_H - 114
        meta_lines = [
            f"{_tr(raw_config, 'supplier')}: {record['supplier']}",
            f"{_tr(raw_config, 'delivery_note_no')}: {record['delivery_note_no']}",
            f"{_tr(raw_config, 'delivery_date')}: {_fmt_date(record['delivery_date'])}",
            *_stationary_document_lines(distractor_plan, placement="meta"),
        ]
        y = _draw_multiline(c, 48, y, meta_lines, leading=16)

        c.setFont("Helvetica-Bold", 11)
        c.drawString(48, y - 8, _tr(raw_config, "customer"))
        c.setFont("Helvetica", 10)
        c.drawString(48, y - 24, record["customer"])

        c.setFont("Helvetica-Bold", 11)
        c.drawString(48, y - 62, _tr(raw_config, "delivery_address"))
        c.setFont("Helvetica", 10)
        _draw_multiline(c, 48, y - 78, [record["site"], *record["site_address"]], leading=14)

        panel_top = PAGE_H - 336
        panel_height = 166
        c.setFillColor(colors.HexColor("#F5F8FB"))
        c.roundRect(36, panel_top - panel_height, PAGE_W - 72, panel_height, 8, fill=1, stroke=0)
        c.setFillColor(colors.black)
        c.setFont("Helvetica-Bold", 11)
        detail_lines = [
            (_tr(raw_config, "product_delivered"), record["fuel"]),
            (_tr(raw_config, "tank_equipment"), record["equipment"]),
            (_tr(raw_config, "delivered_quantity"), f"{_fmt_money(record['quantity'])} {record['unit']}"),
            (_tr(raw_config, "driver_ref"), record["driver_ref"]),
            (_tr(raw_config, "customer_signature"), record["customer_signature"]),
            *_stationary_document_pairs(distractor_plan, placement="summary"),
        ]
        line_y = panel_top - 28
        for label, value in detail_lines:
            if not value:
                continue
            c.drawString(52, line_y, f"{label}:")
            c.setFont("Helvetica", 10)
            c.drawString(184, line_y, value)
            c.setFont("Helvetica-Bold", 11)
            line_y -= 28

        c.setFont("Helvetica", 8)
        c.setFillColor(colors.grey)
        c.drawString(48, 42, _tr(raw_config, "delivery_note_footer"))

    c.save()
    return buffer.getvalue()


def _generate_delivery_note_pdf_variant(raw_config: dict, records: list[dict], layout_plan: dict, distractor_plan) -> bytes:
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    c.setTitle(raw_config.get("document", {}).get("title", _tr(raw_config, "delivery_note_title")))
    c.setSubject(raw_config.get("document", {}).get("subject", "Scope 1 stationary combustion fuel delivery"))

    for index, record in enumerate(records):
        if index > 0:
            c.showPage()

        accent = colors.HexColor("#1E5B88")
        c.setFillColor(accent)
        c.rect(36, PAGE_H - 72, PAGE_W - 72, 30, fill=1, stroke=0)
        c.setFillColor(colors.white)
        c.setFont("Helvetica-Bold", 15)
        c.drawString(48, PAGE_H - 62, _tr(raw_config, "delivery_note_title"))

        current_top = PAGE_H - 114
        for section_name in (layout_plan.get("section_order") or ["addresses", "delivery_details", "footer"]):
            if section_name == "addresses":
                c.setFillColor(colors.black)
                c.setFont("Helvetica", 10)
                meta_lines = [
                    f"{_tr(raw_config, 'supplier')}: {record['supplier']}",
                    f"{_tr(raw_config, 'delivery_note_no')}: {record['delivery_note_no']}",
                    f"{_tr(raw_config, 'delivery_date')}: {_fmt_date(record['delivery_date'])}",
                    *_stationary_document_lines(distractor_plan, placement="meta"),
                ]
                current_top = _draw_multiline(c, 48, current_top, meta_lines, leading=16) - 8
                c.setFont("Helvetica-Bold", 11)
                c.drawString(48, current_top, _tr(raw_config, "customer"))
                c.setFont("Helvetica", 10)
                c.drawString(48, current_top - 16, record["customer"])
                c.setFont("Helvetica-Bold", 11)
                c.drawString(48, current_top - 54, _tr(raw_config, "delivery_address"))
                c.setFont("Helvetica", 10)
                current_top = _draw_multiline(c, 48, current_top - 70, [record["site"], *record["site_address"]], leading=14) - 18
            elif section_name == "delivery_details":
                panel_top = current_top - 8
                panel_height = 166
                c.setFillColor(colors.HexColor("#F5F8FB"))
                c.roundRect(36, panel_top - panel_height, PAGE_W - 72, panel_height, 8, fill=1, stroke=0)
                c.setFillColor(colors.black)
                c.setFont("Helvetica-Bold", 11)
                detail_lines = [
                    (_tr(raw_config, "product_delivered"), record["fuel"]),
                    (_tr(raw_config, "tank_equipment"), record["equipment"]),
                    (_tr(raw_config, "delivered_quantity"), f"{_fmt_money(record['quantity'])} {record['unit']}"),
                    (_tr(raw_config, "driver_ref"), record["driver_ref"]),
                    (_tr(raw_config, "customer_signature"), record["customer_signature"]),
                    *_stationary_document_pairs(distractor_plan, placement="summary"),
                ]
                line_y = panel_top - 28
                for label, value in detail_lines:
                    if not value:
                        continue
                    c.drawString(52, line_y, f"{label}:")
                    c.setFont("Helvetica", 10)
                    c.drawString(184, line_y, value)
                    c.setFont("Helvetica-Bold", 11)
                    line_y -= 28
                current_top = panel_top - panel_height - 18

        c.setFont("Helvetica", 8)
        c.setFillColor(colors.grey)
        c.drawString(48, 42, _tr(raw_config, "delivery_note_footer"))

    c.save()
    return buffer.getvalue()


def generate_delivery_note_docx(raw_config: dict) -> bytes:
    records = _corrupted_delivery_note_records(raw_config)
    distractor_plan = _stationary_distractor_plan(raw_config, "DOCX", document_type="delivery_note")
    layout_plan = _stationary_layout_plan(raw_config, "DOCX")
    if layout_plan.get("enabled"):
        return _generate_delivery_note_docx_variant(raw_config, records, layout_plan, distractor_plan)

    document = Document()
    _style_docx_document(document)
    document.core_properties.title = raw_config.get("document", {}).get("title", _tr(raw_config, "delivery_note_title"))
    document.core_properties.subject = raw_config.get("document", {}).get("subject", "Scope 1 stationary combustion fuel delivery")

    for index, record in enumerate(records):
        banner = document.add_table(rows=1, cols=2)
        banner.style = "Table Grid"
        banner.autofit = False
        banner.columns[0].width = Inches(4.8)
        banner.columns[1].width = Inches(2.0)
        _shade_docx_cell(banner.cell(0, 0), "1E5B88")
        _shade_docx_cell(banner.cell(0, 1), "1E5B88")
        _set_docx_cell_text(banner.cell(0, 0), record["supplier"], bold=True, color="FFFFFF", size=13)
        _set_docx_cell_text(banner.cell(0, 1), _tr(raw_config, "delivery_note_title"), bold=True, color="FFFFFF", size=12)
        banner.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        meta = document.add_table(rows=2, cols=2)
        meta.style = "Table Grid"
        meta.autofit = False
        meta.columns[0].width = Inches(4.2)
        meta.columns[1].width = Inches(2.6)
        for cell, heading in zip(meta.rows[0].cells, [_tr(raw_config, "delivery_address"), _tr(raw_config, "delivery_details")]):
            _shade_docx_cell(cell, "DCEBF5")
            _set_docx_cell_text(cell, heading, bold=True)
        _set_docx_cell_text(
            meta.cell(1, 0),
            "\n".join([record["customer"], "", record["site"], *record["site_address"]]),
        )
        _set_docx_cell_text(
            meta.cell(1, 1),
            (
                f"{_tr(raw_config, 'delivery_note_no')}: {record['delivery_note_no']}\n"
                f"{_tr(raw_config, 'delivery_date')}: {_fmt_date(record['delivery_date'])}\n"
                f"{_tr(raw_config, 'country')}: {record['country']}"
                + ("\n" + "\n".join(_stationary_document_lines(distractor_plan, placement="meta")) if _stationary_document_lines(distractor_plan, placement="meta") else "")
            ),
        )

        spacer = document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(2)

        section_heading = document.add_paragraph()
        section_run = section_heading.add_run(_tr(raw_config, "delivery_confirmation"))
        section_run.bold = True
        section_run.font.size = Pt(11)
        section_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        details = document.add_table(rows=0, cols=2)
        details.style = "Table Grid"
        details.autofit = False
        details.columns[0].width = Inches(2.1)
        details.columns[1].width = Inches(4.7)
        for label, value in [
            (_tr(raw_config, "product_delivered"), record["fuel"]),
            (_tr(raw_config, "tank_equipment"), record["equipment"]),
            (_tr(raw_config, "delivered_quantity"), f"{_fmt_money(record['quantity'])} {record['unit']}"),
            (_tr(raw_config, "driver_ref"), record["driver_ref"]),
            (_tr(raw_config, "customer_signature"), record["customer_signature"]),
            *_stationary_document_pairs(distractor_plan, placement="summary"),
        ]:
            if not value:
                continue
            row = details.add_row().cells
            _shade_docx_cell(row[0], "F5F8FB")
            _set_docx_cell_text(row[0], label, bold=True)
            _set_docx_cell_text(row[1], value)

        footer = document.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_run = footer.add_run(_tr(raw_config, "delivery_note_footer"))
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor.from_string("6E7A86")

        if index < len(records) - 1:
            document.add_page_break()

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _generate_delivery_note_docx_variant(raw_config: dict, records: list[dict], layout_plan: dict, distractor_plan) -> bytes:
    document = Document()
    _style_docx_document(document)
    document.core_properties.title = raw_config.get("document", {}).get("title", _tr(raw_config, "delivery_note_title"))
    document.core_properties.subject = raw_config.get("document", {}).get("subject", "Scope 1 stationary combustion fuel delivery")

    for index, record in enumerate(records):
        banner = document.add_table(rows=1, cols=2)
        banner.style = "Table Grid"
        banner.autofit = False
        banner.columns[0].width = Inches(4.8)
        banner.columns[1].width = Inches(2.0)
        _shade_docx_cell(banner.cell(0, 0), "1E5B88")
        _shade_docx_cell(banner.cell(0, 1), "1E5B88")
        _set_docx_cell_text(banner.cell(0, 0), record["supplier"], bold=True, color="FFFFFF", size=13)
        _set_docx_cell_text(banner.cell(0, 1), _tr(raw_config, "delivery_note_title"), bold=True, color="FFFFFF", size=12)
        banner.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        def render_addresses() -> None:
            section_heading = document.add_paragraph()
            section_heading.add_run(_tr(raw_config, "delivery_address")).bold = True
            document.add_paragraph("\n".join([record["customer"], "", record["site"], *record["site_address"]]))

        def render_delivery_details() -> None:
            meta = document.add_table(rows=1, cols=2)
            meta.style = "Table Grid"
            _shade_docx_cell(meta.cell(0, 0), "DCEBF5")
            _shade_docx_cell(meta.cell(0, 1), "DCEBF5")
            _set_docx_cell_text(meta.cell(0, 0), _tr(raw_config, "delivery_details"), bold=True)
            _set_docx_cell_text(
                meta.cell(0, 1),
                f"{_tr(raw_config, 'delivery_note_no')}: {record['delivery_note_no']}\n{_tr(raw_config, 'delivery_date')}: {_fmt_date(record['delivery_date'])}\n{_tr(raw_config, 'country')}: {record['country']}"
                + ("\n" + "\n".join(_stationary_document_lines(distractor_plan, placement="meta")) if _stationary_document_lines(distractor_plan, placement="meta") else ""),
            )

            details = document.add_table(rows=0, cols=2)
            details.style = "Table Grid"
            details.autofit = False
            details.columns[0].width = Inches(2.1)
            details.columns[1].width = Inches(4.7)
            for label, value in [
                (_tr(raw_config, "product_delivered"), record["fuel"]),
                (_tr(raw_config, "tank_equipment"), record["equipment"]),
                (_tr(raw_config, "delivered_quantity"), f"{_fmt_money(record['quantity'])} {record['unit']}"),
                (_tr(raw_config, "driver_ref"), record["driver_ref"]),
                (_tr(raw_config, "customer_signature"), record["customer_signature"]),
                *_stationary_document_pairs(distractor_plan, placement="summary"),
            ]:
                if not value:
                    continue
                row = details.add_row().cells
                _shade_docx_cell(row[0], "F5F8FB")
                _set_docx_cell_text(row[0], label, bold=True)
                _set_docx_cell_text(row[1], value)

        def render_footer() -> None:
            footer = document.add_paragraph()
            footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            footer_run = footer.add_run(_tr(raw_config, "delivery_note_footer"))
            footer_run.font.size = Pt(8)
            footer_run.font.color.rgb = RGBColor.from_string("6E7A86")

        for section_name in (layout_plan.get("section_order") or ["addresses", "delivery_details", "footer"]):
            if section_name == "addresses":
                render_addresses()
            elif section_name == "delivery_details":
                render_delivery_details()
            elif section_name == "footer":
                render_footer()

        if index < len(records) - 1:
            document.add_page_break()

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def generate_fuel_card_pdf(raw_config: dict) -> bytes:
    return _render_fuel_card_pdf(raw_config, _corrupted_fuel_card_statements(raw_config))


def _render_fuel_card_pdf(raw_config: dict, statements: list[dict]) -> bytes:
    distractor_plan = _stationary_distractor_plan(raw_config, "PDF", document_type="fuel_card")
    layout_plan = _stationary_layout_plan(raw_config, "PDF")
    if layout_plan.get("enabled"):
        return _generate_fuel_card_pdf_variant(raw_config, statements, layout_plan, distractor_plan)

    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    c.setTitle(raw_config.get("document", {}).get("title", _tr(raw_config, "fuel_card_title")))
    c.setSubject(raw_config.get("document", {}).get("subject", "Scope 1 stationary combustion fuel card transactions"))

    page_size = 18
    for statement_index, statement in enumerate(statements):
        transactions = statement["transactions"]
        for page_start in range(0, max(len(transactions), 1), page_size):
            if statement_index > 0 or page_start > 0:
                c.showPage()

            accent = colors.HexColor("#1E5B88")
            c.setFillColor(accent)
            c.rect(36, PAGE_H - 74, PAGE_W - 72, 30, fill=1, stroke=0)
            c.setFillColor(colors.white)
            c.setFont("Helvetica-Bold", 15)
            c.drawString(48, PAGE_H - 63, _tr(raw_config, "fuel_card_title"))

            c.setFillColor(colors.black)
            c.setFont("Helvetica", 10)
            meta_lines = [
                f"{_tr(raw_config, 'account_name')}: {statement['account_name']}",
                f"{_tr(raw_config, 'provider')}: {statement['provider']}",
                f"{_tr(raw_config, 'statement_period')}: {_fmt_date(statement['period_start'])} - {_fmt_date(statement['period_end'])}",
                f"{_tr(raw_config, 'currency')}: {statement['currency']}",
                *_stationary_document_lines(distractor_plan),
            ]
            _draw_multiline(c, 48, PAGE_H - 108, meta_lines, leading=14)

            table_x = 36
            table_top = PAGE_H - 188
            column_widths = [62, 54, 116, 102, 66, 40, 28, 54, 54]
            headers = [
                _tr(raw_config, "card_no"),
                _tr(raw_config, "date"),
                _tr(raw_config, "merchant"),
                _tr(raw_config, "reference"),
                _tr(raw_config, "product"),
                _tr(raw_config, "qty"),
                _tr(raw_config, "unit"),
                _tr(raw_config, "unit_price"),
                _tr(raw_config, "total"),
            ]
            c.setFillColor(accent)
            c.rect(table_x, table_top, sum(column_widths), 22, fill=1, stroke=0)
            c.setFillColor(colors.white)
            c.setFont("Helvetica-Bold", 7.5)
            cursor = table_x + 4
            for header, width in zip(headers, column_widths):
                c.drawString(cursor, table_top + 7, header)
                cursor += width

            row_y = table_top - 18
            c.setFillColor(colors.black)
            c.setFont("Helvetica", 7.5)
            currency_symbol = _currency_symbol(statement["currency"])
            for transaction in transactions[page_start:page_start + page_size]:
                c.rect(table_x, row_y, sum(column_widths), 18, fill=0, stroke=1)
                cursor = table_x + 4
                row_values = [
                    transaction["card_number"],
                    transaction["date"].strftime("%d-%m-%y"),
                    transaction["merchant"],
                    transaction["reference"],
                    transaction["fuel"],
                    _fmt_money(transaction["quantity"]),
                    transaction["unit"],
                    f"{currency_symbol}{_fmt_money(transaction['unit_price'])}",
                    f"{currency_symbol}{_fmt_money(transaction['total'])}",
                ]
                for value, width in zip(row_values, column_widths):
                    c.drawString(cursor, row_y + 5, str(value))
                    cursor += width
                row_y -= 18

            if page_start + page_size >= len(transactions):
                summary_y = row_y - 18
                for label, amount, is_total in _fuel_card_summary_rows(raw_config, statement):
                    c.setFont("Helvetica-Bold" if is_total else "Helvetica", 10)
                    c.drawRightString(PAGE_W - 180, summary_y, label)
                    c.drawRightString(PAGE_W - 48, summary_y, f"{currency_symbol}{_fmt_money(amount)}")
                    summary_y -= 16

            c.setFont("Helvetica", 8)
            c.setFillColor(colors.grey)
            c.drawString(48, 42, _tr(raw_config, "fuel_card_footer"))

    c.save()
    return buffer.getvalue()


def _generate_fuel_card_pdf_variant(raw_config: dict, statements: list[dict], layout_plan: dict, distractor_plan) -> bytes:
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    c.setTitle(raw_config.get("document", {}).get("title", _tr(raw_config, "fuel_card_title")))
    c.setSubject(raw_config.get("document", {}).get("subject", "Scope 1 stationary combustion fuel card transactions"))

    page_size = 18
    for statement_index, statement in enumerate(statements):
        transactions = statement["transactions"]
        for page_start in range(0, max(len(transactions), 1), page_size):
            if statement_index > 0 or page_start > 0:
                c.showPage()

            accent = colors.HexColor("#1E5B88")
            c.setFillColor(accent)
            c.rect(36, PAGE_H - 74, PAGE_W - 72, 30, fill=1, stroke=0)
            c.setFillColor(colors.white)
            c.setFont("Helvetica-Bold", 15)
            c.drawString(48, PAGE_H - 63, _tr(raw_config, "fuel_card_title"))

            meta_lines = [
                f"{_tr(raw_config, 'account_name')}: {statement['account_name']}",
                f"{_tr(raw_config, 'provider')}: {statement['provider']}",
                f"{_tr(raw_config, 'statement_period')}: {_fmt_date(statement['period_start'])} - {_fmt_date(statement['period_end'])}",
                f"{_tr(raw_config, 'currency')}: {statement['currency']}",
                *_stationary_document_lines(distractor_plan),
            ]
            visible_transactions = transactions[page_start:page_start + page_size]
            currency_symbol = _currency_symbol(statement["currency"])
            current_top = PAGE_H - 108
            section_order = [section for section in (layout_plan.get("section_order") or ["summary", "transactions", "footer"]) if section in {"summary", "transactions"}]

            for section_name in section_order:
                if section_name == "summary":
                    c.setFillColor(colors.black)
                    c.setFont("Helvetica", 10)
                    current_top = _draw_multiline(c, 48, current_top, meta_lines, leading=14) - 18
                elif section_name == "transactions":
                    table_x = 36
                    table_top = current_top - 6
                    column_widths = [62, 54, 116, 102, 66, 40, 28, 54, 54]
                    headers = [
                        _tr(raw_config, "card_no"),
                        _tr(raw_config, "date"),
                        _tr(raw_config, "merchant"),
                        _tr(raw_config, "reference"),
                        _tr(raw_config, "product"),
                        _tr(raw_config, "qty"),
                        _tr(raw_config, "unit"),
                        _tr(raw_config, "unit_price"),
                        _tr(raw_config, "total"),
                    ]
                    c.setFillColor(accent)
                    c.rect(table_x, table_top, sum(column_widths), 22, fill=1, stroke=0)
                    c.setFillColor(colors.white)
                    c.setFont("Helvetica-Bold", 7.5)
                    cursor = table_x + 4
                    for header, width in zip(headers, column_widths):
                        c.drawString(cursor, table_top + 7, header)
                        cursor += width

                    row_y = table_top - 18
                    c.setFillColor(colors.black)
                    c.setFont("Helvetica", 7.5)
                    for transaction in visible_transactions:
                        c.rect(table_x, row_y, sum(column_widths), 18, fill=0, stroke=1)
                        cursor = table_x + 4
                        row_values = [
                            transaction["card_number"],
                            transaction["date"].strftime("%d-%m-%y"),
                            transaction["merchant"],
                            transaction["reference"],
                            transaction["fuel"],
                            _fmt_money(transaction["quantity"]),
                            transaction["unit"],
                            f"{currency_symbol}{_fmt_money(transaction['unit_price'])}",
                            f"{currency_symbol}{_fmt_money(transaction['total'])}",
                        ]
                        for value, width in zip(row_values, column_widths):
                            c.drawString(cursor, row_y + 5, str(value))
                            cursor += width
                        row_y -= 18

                    if page_start + page_size >= len(transactions):
                        summary_y = row_y - 18
                        for label, amount, is_total in _fuel_card_summary_rows(raw_config, statement):
                            c.setFont("Helvetica-Bold" if is_total else "Helvetica", 10)
                            c.drawRightString(PAGE_W - 180, summary_y, label)
                            c.drawRightString(PAGE_W - 48, summary_y, f"{currency_symbol}{_fmt_money(amount)}")
                            summary_y -= 16
                        row_y = summary_y - 16
                    current_top = row_y - 10

            c.setFont("Helvetica", 8)
            c.setFillColor(colors.grey)
            c.drawString(48, 42, _tr(raw_config, "fuel_card_footer"))

    c.save()
    return buffer.getvalue()


def generate_fuel_card_docx(raw_config: dict) -> bytes:
    return _render_fuel_card_docx(raw_config, _corrupted_fuel_card_statements(raw_config))


def _render_fuel_card_docx(raw_config: dict, statements: list[dict]) -> bytes:
    distractor_plan = _stationary_distractor_plan(raw_config, "DOCX", document_type="fuel_card")
    layout_plan = _stationary_layout_plan(raw_config, "DOCX")
    if layout_plan.get("enabled"):
        return _generate_fuel_card_docx_variant(raw_config, statements, layout_plan, distractor_plan)

    document = Document()
    _style_docx_document(document)
    document.core_properties.title = raw_config.get("document", {}).get("title", _tr(raw_config, "fuel_card_title"))
    document.core_properties.subject = raw_config.get("document", {}).get("subject", "Scope 1 stationary combustion fuel card transactions")

    for statement_index, statement in enumerate(statements):
        banner = document.add_table(rows=1, cols=2)
        banner.style = "Table Grid"
        banner.autofit = False
        banner.columns[0].width = Inches(4.8)
        banner.columns[1].width = Inches(2.0)
        _shade_docx_cell(banner.cell(0, 0), "1E5B88")
        _shade_docx_cell(banner.cell(0, 1), "1E5B88")
        _set_docx_cell_text(banner.cell(0, 0), statement["account_name"], bold=True, color="FFFFFF", size=13)
        _set_docx_cell_text(banner.cell(0, 1), _tr(raw_config, "fuel_card_title"), bold=True, color="FFFFFF", size=12)
        banner.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        meta = document.add_table(rows=2, cols=2)
        meta.style = "Table Grid"
        meta.autofit = False
        meta.columns[0].width = Inches(4.2)
        meta.columns[1].width = Inches(2.6)
        for cell, heading in zip(meta.rows[0].cells, [_tr(raw_config, "account_details"), _tr(raw_config, "statement_details")]):
            _shade_docx_cell(cell, "DCEBF5")
            _set_docx_cell_text(cell, heading, bold=True)
        _set_docx_cell_text(
            meta.cell(1, 0),
            f"{_tr(raw_config, 'account_name')}: {statement['account_name']}\n{_tr(raw_config, 'provider')}: {statement['provider']}",
        )
        _set_docx_cell_text(
            meta.cell(1, 1),
            (
                f"{_tr(raw_config, 'statement_period')}: {_fmt_date(statement['period_start'])} - {_fmt_date(statement['period_end'])}\n"
                f"{_tr(raw_config, 'currency')}: {statement['currency']}"
                + ("\n" + "\n".join(_stationary_document_lines(distractor_plan)) if _stationary_document_lines(distractor_plan) else "")
            ),
        )

        document.add_paragraph()
        table = document.add_table(rows=1, cols=9)
        table.style = "Table Grid"
        headers = [
            _tr(raw_config, "card_no"),
            _tr(raw_config, "date"),
            _tr(raw_config, "merchant"),
            _tr(raw_config, "reference"),
            _tr(raw_config, "product"),
            _tr(raw_config, "qty"),
            _tr(raw_config, "unit"),
            _tr(raw_config, "unit_price"),
            _tr(raw_config, "total"),
            _tr(raw_config, "currency"),
        ]
        for cell, header in zip(table.rows[0].cells, headers):
            _shade_docx_cell(cell, "F5F8FB")
            _set_docx_cell_text(cell, header, bold=True)

        currency_symbol = _currency_symbol(statement["currency"])
        for transaction in statement["transactions"]:
            row = table.add_row().cells
            values = [
                transaction["card_number"],
                transaction["date"].strftime("%d-%m-%y"),
                transaction["merchant"],
                transaction["reference"],
                transaction["fuel"],
                _fmt_money(transaction["quantity"]),
                transaction["unit"],
                f"{currency_symbol}{_fmt_money(transaction['unit_price'])}",
                f"{currency_symbol}{_fmt_money(transaction['total'])}",
            ]
            for cell, value in zip(row, values):
                _set_docx_cell_text(cell, str(value))

        _add_fuel_card_totals_docx(document, raw_config, statement, currency_symbol)

        if statement_index < len(statements) - 1:
            document.add_page_break()

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _generate_fuel_card_docx_variant(raw_config: dict, statements: list[dict], layout_plan: dict, distractor_plan) -> bytes:
    document = Document()
    _style_docx_document(document)
    document.core_properties.title = raw_config.get("document", {}).get("title", _tr(raw_config, "fuel_card_title"))
    document.core_properties.subject = raw_config.get("document", {}).get("subject", "Scope 1 stationary combustion fuel card transactions")

    for statement_index, statement in enumerate(statements):
        banner = document.add_table(rows=1, cols=2)
        banner.style = "Table Grid"
        banner.autofit = False
        banner.columns[0].width = Inches(4.8)
        banner.columns[1].width = Inches(2.0)
        _shade_docx_cell(banner.cell(0, 0), "1E5B88")
        _shade_docx_cell(banner.cell(0, 1), "1E5B88")
        _set_docx_cell_text(banner.cell(0, 0), statement["account_name"], bold=True, color="FFFFFF", size=13)
        _set_docx_cell_text(banner.cell(0, 1), _tr(raw_config, "fuel_card_title"), bold=True, color="FFFFFF", size=12)
        banner.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        currency_symbol = _currency_symbol(statement["currency"])

        def render_summary() -> None:
            meta = document.add_table(rows=2, cols=2)
            meta.style = "Table Grid"
            meta.autofit = False
            meta.columns[0].width = Inches(4.2)
            meta.columns[1].width = Inches(2.6)
            for cell, heading in zip(meta.rows[0].cells, [_tr(raw_config, "account_details"), _tr(raw_config, "statement_details")]):
                _shade_docx_cell(cell, "DCEBF5")
                _set_docx_cell_text(cell, heading, bold=True)
            _set_docx_cell_text(
                meta.cell(1, 0),
                f"{_tr(raw_config, 'account_name')}: {statement['account_name']}\n{_tr(raw_config, 'provider')}: {statement['provider']}",
            )
            _set_docx_cell_text(
                meta.cell(1, 1),
                (
                    f"{_tr(raw_config, 'statement_period')}: {_fmt_date(statement['period_start'])} - {_fmt_date(statement['period_end'])}\n"
                    f"{_tr(raw_config, 'currency')}: {statement['currency']}"
                    + ("\n" + "\n".join(_stationary_document_lines(distractor_plan)) if _stationary_document_lines(distractor_plan) else "")
                ),
            )

        def render_transactions() -> None:
            document.add_paragraph()
            table = document.add_table(rows=1, cols=10)
            table.style = "Table Grid"
            headers = [
                _tr(raw_config, "card_no"),
                _tr(raw_config, "date"),
                _tr(raw_config, "merchant"),
                _tr(raw_config, "reference"),
                _tr(raw_config, "product"),
                _tr(raw_config, "qty"),
                _tr(raw_config, "unit"),
                _tr(raw_config, "unit_price"),
                _tr(raw_config, "total"),
                _tr(raw_config, "currency"),
            ]
            for cell, header in zip(table.rows[0].cells, headers):
                _shade_docx_cell(cell, "F5F8FB")
                _set_docx_cell_text(cell, header, bold=True)

            for transaction in statement["transactions"]:
                row = table.add_row().cells
                values = [
                    transaction["card_number"],
                    transaction["date"].strftime("%d-%m-%y"),
                    transaction["merchant"],
                    transaction["reference"],
                    transaction["fuel"],
                    _fmt_money(transaction["quantity"]),
                    transaction["unit"],
                    f"{currency_symbol}{_fmt_money(transaction['unit_price'])}",
                    f"{currency_symbol}{_fmt_money(transaction['total'])}",
                    statement["currency"].split()[0],
                ]
                for cell, value in zip(row, values):
                    _set_docx_cell_text(cell, str(value))

            _add_fuel_card_totals_docx(document, raw_config, statement, currency_symbol)

        def render_footer() -> None:
            footer = document.add_paragraph()
            footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = footer.add_run(_tr(raw_config, "fuel_card_footer"))
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor.from_string("6E7A86")

        for section_name in (layout_plan.get("section_order") or ["summary", "transactions", "footer"]):
            if section_name == "summary":
                render_summary()
            elif section_name == "transactions":
                render_transactions()
            elif section_name == "footer":
                render_footer()

        if statement_index < len(statements) - 1:
            document.add_page_break()

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def generate_fuel_card_xlsx(raw_config: dict) -> bytes:
    return _render_fuel_card_xlsx(raw_config, _corrupted_fuel_card_statements(raw_config))


def _render_fuel_card_xlsx(raw_config: dict, statements: list[dict]) -> bytes:
    layout_plan = _stationary_layout_plan(raw_config, "XLSX")
    distractor_plan = _stationary_distractor_plan(raw_config, "XLSX", document_type="fuel_card")
    distractor_fields = _stationary_distractor_field_map(distractor_plan)
    ordered_ids = _augment_stationary_field_ids(_ordered_stationary_field_ids(layout_plan, list(_FUEL_CARD_HEADER_KEYS)), distractor_plan)
    workbook = openpyxl.Workbook()
    workbook.remove(workbook.active)

    for index, statement in enumerate(statements, start=1):
        sheet = workbook.create_sheet(title=(statement["company"] or f"Account {index}")[:31])
        sheet["A1"] = _tr(raw_config, "fuel_card_title")
        sheet["A1"].font = Font(size=14, bold=True)
        sheet["A2"] = _tr(raw_config, "account_name")
        sheet["B2"] = statement["account_name"]
        sheet["A3"] = _tr(raw_config, "statement_period")
        sheet["B3"] = f"{statement['period_start'].isoformat()} to {statement['period_end'].isoformat()}"
        sheet["A4"] = _tr(raw_config, "currency")
        sheet["B4"] = statement["currency"]

        header_row = _write_stationary_xlsx_preamble(sheet, 5, layout_plan)
        headers = [
            _stationary_header_text(
                raw_config,
                layout_plan,
                field_id,
                _FUEL_CARD_HEADER_KEYS.get(field_id, field_id),
                distractor_fields,
            )
            for field_id in ordered_ids
        ]
        header_fill = PatternFill(fill_type="solid", fgColor="1E5B88")
        for column_index, header in enumerate(headers, start=1):
            cell = sheet.cell(row=header_row, column=column_index, value=header)
            cell.font = Font(color="FFFFFF", bold=True)
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center")

        row_index = header_row + 1
        for transaction in statement["transactions"]:
            row_map = {
                "card_no": transaction["card_number"],
                "date": transaction["date"].strftime("%d-%m-%y"),
                "merchant": transaction["merchant"],
                "site": transaction["site"],
                "country": transaction["country"],
                "equipment": transaction["equipment"],
                "emission_source": transaction["emission_source"],
                "product": transaction["fuel"],
                "qty": float(transaction["quantity"]) if not isinstance(transaction["quantity"], str) else transaction["quantity"],
                "unit": transaction["unit"],
                "unit_price": float(transaction["unit_price"]) if not isinstance(transaction["unit_price"], str) else transaction["unit_price"],
                "total": float(transaction["total"]) if not isinstance(transaction["total"], str) else transaction["total"],
                "currency": statement["currency"].split()[0],
            }
            values = _stationary_row_values(
                row_map,
                ordered_ids,
                distractor_plan,
                row_key=f"{transaction['date'].isoformat()}:{transaction['card_number']}:{transaction['reference']}",
                statement_key=f"{statement['account_name']}:{statement['period_start'].isoformat()}:{statement['period_end'].isoformat()}",
            )
            for column_index, value in enumerate(values, start=1):
                sheet.cell(row=row_index, column=column_index, value=value)
            row_index += 1

        row_index += 1
        for label, amount, is_total in _fuel_card_summary_rows(raw_config, statement):
            label_cell = sheet.cell(row=row_index, column=1, value=label)
            label_cell.font = Font(bold=True)
            amount_cell = sheet.cell(
                row=row_index,
                column=2,
                value=float(amount) if not isinstance(amount, str) else amount,
            )
            amount_cell.font = Font(bold=is_total)
            row_index += 1

        width_map = {
            "card_no": 12,
            "date": 12,
            "merchant": 22,
            "site": 18,
            "country": 14,
            "equipment": 18,
            "emission_source": 18,
            "product": 14,
            "qty": 10,
            "unit": 8,
            "unit_price": 12,
            "total": 12,
            "currency": 10,
        }
        for column_index, field_id in enumerate(ordered_ids, start=1):
            width = width_map.get(field_id, 14)
            sheet.column_dimensions[get_column_letter(column_index)].width = width

    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def generate_fuel_card_csv(raw_config: dict) -> bytes:
    return _render_fuel_card_csv(raw_config, _corrupted_fuel_card_statements(raw_config))


def _render_fuel_card_csv(raw_config: dict, statements: list[dict]) -> bytes:
    layout_plan = _stationary_layout_plan(raw_config, "CSV")
    distractor_plan = _stationary_distractor_plan(raw_config, "CSV", document_type="fuel_card")
    distractor_fields = _stationary_distractor_field_map(distractor_plan)
    ordered_ids = _augment_stationary_field_ids(_ordered_stationary_field_ids(layout_plan, list(_FUEL_CARD_HEADER_KEYS)), distractor_plan)
    buffer = StringIO()
    writer = csv.writer(buffer)

    for statement_index, statement in enumerate(statements):
        if statement_index > 0:
            writer.writerow([])
        writer.writerow([_tr(raw_config, "fuel_card_title")])
        writer.writerow([_tr(raw_config, "account_name"), statement["account_name"]])
        writer.writerow([_tr(raw_config, "statement_period"), f"{statement['period_start'].isoformat()} to {statement['period_end'].isoformat()}"])
        writer.writerow([_tr(raw_config, "currency"), statement["currency"]])
        writer.writerow([])
        _write_stationary_csv_preamble(writer, layout_plan)
        writer.writerow([
            _stationary_header_text(
                raw_config,
                layout_plan,
                field_id,
                _FUEL_CARD_HEADER_KEYS.get(field_id, field_id),
                distractor_fields,
            )
            for field_id in ordered_ids
        ])
        for transaction in statement["transactions"]:
            row_map = {
                "card_no": transaction["card_number"],
                "date": transaction["date"].strftime("%d-%m-%y"),
                "merchant": transaction["merchant"],
                "site": transaction["site"],
                "country": transaction["country"],
                "equipment": transaction["equipment"],
                "emission_source": transaction["emission_source"],
                "product": transaction["fuel"],
                "qty": _fmt_num(transaction["quantity"], ".2f"),
                "unit": transaction["unit"],
                "unit_price": _fmt_num(transaction["unit_price"], ".2f"),
                "total": _fmt_num(transaction["total"], ".2f"),
                "currency": statement["currency"].split()[0] if isinstance(statement["currency"], str) and " " in statement["currency"] else statement["currency"],
            }
            writer.writerow(
                _stationary_row_values(
                    row_map,
                    ordered_ids,
                    distractor_plan,
                    row_key=f"{transaction['date'].isoformat()}:{transaction['card_number']}:{transaction['reference']}",
                    statement_key=f"{statement['account_name']}:{statement['period_start'].isoformat()}:{statement['period_end'].isoformat()}",
                )
            )

        writer.writerow([])
        for label, amount, _is_total in _fuel_card_summary_rows(raw_config, statement):
            writer.writerow([label, _fmt_num(amount, ".2f")])

    return buffer.getvalue().encode("utf-8-sig")


def _log_headers(raw_config: dict) -> list[str]:
    return [
        _tr(raw_config, "company"),
        _tr(raw_config, "site"),
        _tr(raw_config, "country"),
        _tr(raw_config, "date"),
        _tr(raw_config, "start_time"),
        _tr(raw_config, "end_time"),
        _tr(raw_config, "run_hours"),
        _tr(raw_config, "start_fuel"),
        _tr(raw_config, "end_fuel"),
        _tr(raw_config, "fuel_used"),
        _tr(raw_config, "unit"),
        _tr(raw_config, "equipment"),
        _tr(raw_config, "emission_source"),
        _tr(raw_config, "fuel_type"),
        _tr(raw_config, "notes"),
    ]


def generate_generator_log_xlsx(raw_config: dict) -> bytes:
    rows = _corrupted_generator_log_rows(raw_config)
    layout_plan = _stationary_layout_plan(raw_config, "XLSX")
    distractor_plan = _stationary_distractor_plan(raw_config, "XLSX", document_type="generator_log")
    distractor_fields = _stationary_distractor_field_map(distractor_plan)
    ordered_ids = _augment_stationary_field_ids(_ordered_stationary_field_ids(layout_plan, list(_GENERATOR_LOG_HEADER_KEYS)), distractor_plan)
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = _tr(raw_config, "generator_log_sheet_title")

    title = raw_config.get("document", {}).get("title", _tr(raw_config, "generator_operation_log_title"))
    sheet["A1"] = title
    sheet["A1"].font = Font(size=14, bold=True)
    sheet["A2"] = raw_config.get("financial_period", {}).get("label", "")
    sheet["A2"].font = Font(italic=True)

    header_row = _write_stationary_xlsx_preamble(sheet, 3, layout_plan)
    headers = [
        _stationary_header_text(
            raw_config,
            layout_plan,
            field_id,
            _GENERATOR_LOG_HEADER_KEYS.get(field_id, field_id),
            distractor_fields,
        )
        for field_id in ordered_ids
    ]
    header_fill = PatternFill(fill_type="solid", fgColor="1E5B88")
    for column_index, header in enumerate(headers, start=1):
        cell = sheet.cell(row=header_row, column=column_index, value=header)
        cell.font = Font(color="FFFFFF", bold=True)
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")

    for row_index, row in enumerate(rows, start=header_row + 1):
        row_map = {
            "company": row["company"],
            "site": row["site"],
            "country": row["country"],
            "date": row["date"].strftime("%d-%m-%y"),
            "start_time": row["start_time"],
            "end_time": row["end_time"],
            "run_hours": row["run_hours"],
            "start_fuel": row["start_fuel"],
            "end_fuel": row["end_fuel"],
            "fuel_used": row["fuel_used"],
            "unit": row["unit"],
            "equipment": row["equipment"],
            "emission_source": row["emission_source"],
            "fuel_type": row["fuel"],
            "notes": row["notes"],
        }
        values = _stationary_row_values(
            row_map,
            ordered_ids,
            distractor_plan,
            row_key=f"{row['period']}:{row['equipment']}:{row['start_time']}",
            block_key=f"{row['company']}:{row['site']}",
        )
        for column_index, value in enumerate(values, start=1):
            sheet.cell(row=row_index, column=column_index, value=value)

    width_map = {
        "company": 18,
        "site": 22,
        "country": 18,
        "date": 12,
        "start_time": 12,
        "end_time": 12,
        "run_hours": 10,
        "start_fuel": 12,
        "end_fuel": 12,
        "fuel_used": 10,
        "unit": 8,
        "equipment": 20,
        "emission_source": 22,
        "fuel_type": 18,
        "notes": 18,
    }
    for column_index, field_id in enumerate(ordered_ids, start=1):
        width = width_map.get(field_id, 14)
        sheet.column_dimensions[get_column_letter(column_index)].width = width

    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def generate_generator_log_csv(raw_config: dict) -> bytes:
    rows = _corrupted_generator_log_rows(raw_config)
    layout_plan = _stationary_layout_plan(raw_config, "CSV")
    distractor_plan = _stationary_distractor_plan(raw_config, "CSV", document_type="generator_log")
    distractor_fields = _stationary_distractor_field_map(distractor_plan)
    ordered_ids = _augment_stationary_field_ids(_ordered_stationary_field_ids(layout_plan, list(_GENERATOR_LOG_HEADER_KEYS)), distractor_plan)
    buffer = StringIO()
    writer = csv.writer(buffer)
    _write_stationary_csv_preamble(writer, layout_plan)
    writer.writerow([
        _stationary_header_text(
            raw_config,
            layout_plan,
            field_id,
            _GENERATOR_LOG_HEADER_KEYS.get(field_id, field_id),
            distractor_fields,
        )
        for field_id in ordered_ids
    ])
    for row in rows:
        row_map = {
            "company": row["company"],
            "site": row["site"],
            "country": row["country"],
            "date": row["date"].strftime("%d-%m-%y"),
            "start_time": row["start_time"],
            "end_time": row["end_time"],
            "run_hours": f"{row['run_hours']:.2f}",
            "start_fuel": row["start_fuel"],
            "end_fuel": row["end_fuel"],
            "fuel_used": _fmt_num(row["fuel_used"], ".2f"),
            "unit": row["unit"],
            "equipment": row["equipment"],
            "emission_source": row["emission_source"],
            "fuel_type": row["fuel"],
            "notes": row["notes"],
        }
        writer.writerow(
            _stationary_row_values(
                row_map,
                ordered_ids,
                distractor_plan,
                row_key=f"{row['period']}:{row['equipment']}:{row['start_time']}",
                block_key=f"{row['company']}:{row['site']}",
            )
        )
    return buffer.getvalue().encode("utf-8-sig")


def generate_bems_equipment_report_pdf(raw_config: dict) -> bytes:
    blocks = _corrupted_bems_site_blocks(raw_config)
    distractor_plan = _stationary_distractor_plan(raw_config, "PDF", document_type="bems")
    layout_plan = _stationary_layout_plan(raw_config, "PDF", document_type="bems")
    if layout_plan.get("enabled"):
        return _generate_bems_equipment_report_pdf_variant(raw_config, blocks, layout_plan, distractor_plan)

    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    c.setTitle(raw_config.get("document", {}).get("title", _tr(raw_config, "bems_equipment_title")))
    c.setSubject(raw_config.get("document", {}).get("subject", "Scope 1 stationary combustion BEMS export"))

    for block_index, block in enumerate(blocks):
        if block_index > 0:
            c.showPage()

        accent = colors.HexColor("#1E5B88")
        c.setFillColor(accent)
        c.rect(36, PAGE_H - 78, PAGE_W - 72, 34, fill=1, stroke=0)
        c.setFillColor(colors.white)
        c.setFont("Helvetica-Bold", 15)
        c.drawString(48, PAGE_H - 64, _tr(raw_config, "bems_equipment_title"))

        c.setFillColor(colors.black)
        c.setFont("Helvetica", 10)
        meta_lines = [
            f"{_tr(raw_config, 'company')}: {block['company']}",
            f"{_tr(raw_config, 'site')}: {block['site']}",
            f"{_tr(raw_config, 'country')}: {block['country']}",
            f"{_tr(raw_config, 'reporting_period')}: {block['period_label']}",
            *_stationary_document_lines(distractor_plan),
        ]
        meta_lines = [line for line in meta_lines if not line.endswith(": ")]
        _draw_multiline(c, 48, PAGE_H - 108, meta_lines, leading=14)

        total_assets = len(block["assets"])
        total_hours = sum(
            Decimal("0") if (asset["operating_hours"] is None or isinstance(asset["operating_hours"], str))
            else asset["operating_hours"]
            for asset in block["assets"]
        )
        dominant_asset = max(
            (a for a in block["assets"] if not isinstance(a["quantity"], str)),
            key=lambda asset: asset["quantity"],
            default=None,
        )
        cards = [
            (_tr(raw_config, "assets"), str(total_assets)),
            (_tr(raw_config, "operating_hours"), _fmt_optional_number(total_hours, " h") or "n/a"),
            (_tr(raw_config, "top_asset"), dominant_asset["asset_tag"] if dominant_asset else "n/a"),
        ]

        x = 48
        card_y = PAGE_H - 205
        for title, value in cards:
            c.setFillColor(colors.HexColor("#F2F6FA"))
            c.roundRect(x, card_y, 150, 46, 6, stroke=0, fill=1)
            c.setFillColor(colors.HexColor("#567389"))
            c.setFont("Helvetica", 8)
            c.drawString(x + 10, card_y + 30, title)
            c.setFillColor(colors.black)
            c.setFont("Helvetica-Bold", 12)
            c.drawString(x + 10, card_y + 13, value)
            x += 166

        table_x = 48
        table_top = PAGE_H - 290
        column_widths = [64, 136, 88, 74, 64, 48, 74]
        headers = [
            _tr(raw_config, "equipment_tag"),
            _tr(raw_config, "equipment_name"),
            _tr(raw_config, "emission_source"),
            _tr(raw_config, "fuel_type"),
            _tr(raw_config, "consumption"),
            _tr(raw_config, "unit"),
            _tr(raw_config, "operating_hours"),
        ]

        c.setFillColor(accent)
        c.rect(table_x, table_top, sum(column_widths), 24, fill=1, stroke=0)
        c.setFillColor(colors.white)
        c.setFont("Helvetica-Bold", 8)
        cursor = table_x + 6
        for header, width in zip(headers, column_widths):
            c.drawString(cursor, table_top + 8, header)
            cursor += width

        row_y = table_top - 22
        max_quantity = max(
            (_safe_float(asset["quantity"]) for asset in block["assets"]),
            default=1.0,
        ) or 1.0
        for asset in block["assets"]:
            c.setFillColor(colors.black)
            c.rect(table_x, row_y, sum(column_widths), 20, fill=0, stroke=1)
            cursor = table_x + 6
            row_values = [
                asset["asset_tag"],
                asset["equipment_name"],
                asset["emission_source"],
                asset["fuel"],
                _fmt_money(asset["quantity"]),
                asset["unit"],
                _fmt_optional_number(asset["operating_hours"]),
            ]
            c.setFont("Helvetica", 8)
            for value, width in zip(row_values, column_widths):
                c.drawString(cursor, row_y + 6, str(value))
                cursor += width
            row_y -= 20

        chart_y = row_y - 110
        c.setFont("Helvetica-Bold", 10)
        c.drawString(48, chart_y + 96, _tr(raw_config, "equipment_trend_snapshot"))
        for idx, asset in enumerate(block["assets"][:5]):
            bar_y = chart_y + 72 - (idx * 18)
            bar_width = 220 * (_safe_float(asset["quantity"]) / max_quantity if max_quantity else 0)
            c.setFont("Helvetica", 8)
            c.drawString(48, bar_y + 4, asset["asset_tag"])
            c.setFillColor(colors.HexColor("#DCEBF5"))
            c.rect(110, bar_y, 220, 10, fill=1, stroke=0)
            c.setFillColor(accent)
            c.rect(110, bar_y, bar_width, 10, fill=1, stroke=0)
            c.setFillColor(colors.black)
            c.drawString(340, bar_y + 2, f"{_fmt_money(asset['quantity'])} {asset['unit']}")

        c.setFont("Helvetica", 8)
        c.setFillColor(colors.grey)
        c.drawString(48, 42, _tr(raw_config, "dashboard_summary_footer"))

    c.save()
    return buffer.getvalue()


def _generate_bems_equipment_report_pdf_variant(raw_config: dict, blocks: list[dict], layout_plan: dict, distractor_plan) -> bytes:
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    c.setTitle(raw_config.get("document", {}).get("title", _tr(raw_config, "bems_equipment_title")))
    c.setSubject(raw_config.get("document", {}).get("subject", "Scope 1 stationary combustion BEMS export"))

    for block_index, block in enumerate(blocks):
        if block_index > 0:
            c.showPage()

        accent = colors.HexColor("#1E5B88")
        c.setFillColor(accent)
        c.rect(36, PAGE_H - 78, PAGE_W - 72, 34, fill=1, stroke=0)
        c.setFillColor(colors.white)
        c.setFont("Helvetica-Bold", 15)
        c.drawString(48, PAGE_H - 64, _tr(raw_config, "bems_equipment_title"))

        total_assets = len(block["assets"])
        total_hours = sum(
            Decimal("0") if (asset["operating_hours"] is None or isinstance(asset["operating_hours"], str)) else asset["operating_hours"]
            for asset in block["assets"]
        )
        dominant_asset = max(
            (a for a in block["assets"] if not isinstance(a["quantity"], str)),
            key=lambda asset: asset["quantity"],
            default=None,
        )

        def render_meta(current_top: float) -> float:
            c.setFillColor(colors.black)
            c.setFont("Helvetica", 10)
            meta_lines = [
                f"{_tr(raw_config, 'company')}: {block['company']}",
                f"{_tr(raw_config, 'site')}: {block['site']}",
                f"{_tr(raw_config, 'country')}: {block['country']}",
                f"{_tr(raw_config, 'reporting_period')}: {block['period_label']}",
                *_stationary_document_lines(distractor_plan),
            ]
            meta_lines = [line for line in meta_lines if not line.endswith(": ")]
            current_top = _draw_multiline(c, 48, current_top, meta_lines, leading=14) - 16

            cards = [
                (_tr(raw_config, "assets"), str(total_assets)),
                (_tr(raw_config, "operating_hours"), _fmt_optional_number(total_hours, " h") or "n/a"),
                (_tr(raw_config, "top_asset"), dominant_asset["asset_tag"] if dominant_asset else "n/a"),
            ]
            x = 48
            card_y = current_top - 46
            for title, value in cards:
                c.setFillColor(colors.HexColor("#F2F6FA"))
                c.roundRect(x, card_y, 150, 46, 6, stroke=0, fill=1)
                c.setFillColor(colors.HexColor("#567389"))
                c.setFont("Helvetica", 8)
                c.drawString(x + 10, card_y + 30, title)
                c.setFillColor(colors.black)
                c.setFont("Helvetica-Bold", 12)
                c.drawString(x + 10, card_y + 13, value)
                x += 166
            return card_y - 18

        def render_table(current_top: float) -> float:
            table_x = 48
            table_top = current_top - 8
            column_widths = [64, 136, 88, 74, 64, 48, 74]
            headers = [
                _tr(raw_config, "equipment_tag"),
                _tr(raw_config, "equipment_name"),
                _tr(raw_config, "emission_source"),
                _tr(raw_config, "fuel_type"),
                _tr(raw_config, "consumption"),
                _tr(raw_config, "unit"),
                _tr(raw_config, "operating_hours"),
            ]
            c.setFillColor(accent)
            c.rect(table_x, table_top, sum(column_widths), 24, fill=1, stroke=0)
            c.setFillColor(colors.white)
            c.setFont("Helvetica-Bold", 8)
            cursor = table_x + 6
            for header, width in zip(headers, column_widths):
                c.drawString(cursor, table_top + 8, header)
                cursor += width

            row_y = table_top - 22
            max_quantity = max((_safe_float(asset["quantity"]) for asset in block["assets"]), default=1.0) or 1.0
            for asset in block["assets"]:
                c.setFillColor(colors.black)
                c.rect(table_x, row_y, sum(column_widths), 20, fill=0, stroke=1)
                cursor = table_x + 6
                row_values = [
                    asset["asset_tag"],
                    asset["equipment_name"],
                    asset["emission_source"],
                    asset["fuel"],
                    _fmt_money(asset["quantity"]),
                    asset["unit"],
                    _fmt_optional_number(asset["operating_hours"]),
                ]
                c.setFont("Helvetica", 8)
                for value, width in zip(row_values, column_widths):
                    c.drawString(cursor, row_y + 6, str(value))
                    cursor += width
                row_y -= 20

            chart_y = row_y - 110
            c.setFont("Helvetica-Bold", 10)
            c.drawString(48, chart_y + 96, _tr(raw_config, "equipment_trend_snapshot"))
            for idx, asset in enumerate(block["assets"][:5]):
                bar_y = chart_y + 72 - (idx * 18)
                bar_width = 220 * (_safe_float(asset["quantity"]) / max_quantity if max_quantity else 0)
                c.setFont("Helvetica", 8)
                c.drawString(48, bar_y + 4, asset["asset_tag"])
                c.setFillColor(colors.HexColor("#DCEBF5"))
                c.rect(110, bar_y, 220, 10, fill=1, stroke=0)
                c.setFillColor(accent)
                c.rect(110, bar_y, bar_width, 10, fill=1, stroke=0)
                c.setFillColor(colors.black)
                c.drawString(340, bar_y + 2, f"{_fmt_money(asset['quantity'])} {asset['unit']}")
            return chart_y - 18

        current_top = PAGE_H - 108
        for section_name in (layout_plan.get("section_order") or ["meta", "table", "footer"]):
            if section_name == "meta":
                current_top = render_meta(current_top)
            elif section_name == "table":
                current_top = render_table(current_top)

        c.setFont("Helvetica", 8)
        c.setFillColor(colors.grey)
        c.drawString(48, 42, _tr(raw_config, "dashboard_summary_footer"))

    c.save()
    return buffer.getvalue()


def generate_bems_time_series_pdf(raw_config: dict) -> bytes:
    blocks = _build_bems_trend_exports(raw_config)
    distractor_plan = _stationary_distractor_plan(raw_config, "PDF", document_type="bems")
    layout_plan = _stationary_layout_plan(raw_config, "PDF", document_type="bems")
    if layout_plan.get("enabled"):
        return _generate_bems_time_series_pdf_variant(raw_config, blocks, layout_plan, distractor_plan)

    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    c.setTitle(raw_config.get("document", {}).get("title", _tr(raw_config, "bems_time_series_title")))
    c.setSubject(raw_config.get("document", {}).get("subject", "Scope 1 stationary combustion BEMS export"))

    for block_index, block in enumerate(blocks):
        rows = block["rows"]
        page_size = 28
        for page_start in range(0, max(len(rows), 1), page_size):
            if block_index > 0 or page_start > 0:
                c.showPage()

            accent = colors.HexColor("#1E5B88")
            c.setFillColor(accent)
            c.rect(36, PAGE_H - 78, PAGE_W - 72, 34, fill=1, stroke=0)
            c.setFillColor(colors.white)
            c.setFont("Helvetica-Bold", 15)
            c.drawString(48, PAGE_H - 64, _tr(raw_config, "bems_time_series_title"))

            c.setFillColor(colors.black)
            c.setFont("Helvetica", 10)
            meta_lines = [
                f"{_tr(raw_config, 'company')}: {block['company']}",
                f"{_tr(raw_config, 'site')}: {block['site']}",
                f"{_tr(raw_config, 'country')}: {block['country']}",
                f"{_tr(raw_config, 'reporting_period')}: {block['period_label']}",
                *_stationary_document_lines(distractor_plan),
            ]
            meta_lines = [line for line in meta_lines if not line.endswith(": ")]
            _draw_multiline(c, 48, PAGE_H - 108, meta_lines, leading=14)

            asset_count = len(block["assets"])
            interval_minutes = _bems_interval_minutes(raw_config)
            cards = [
                (_tr(raw_config, "assets"), str(asset_count)),
                (_tr(raw_config, "interval"), f"{interval_minutes} min"),
                (_tr(raw_config, "rows"), str(len(rows))),
            ]
            x = 48
            card_y = PAGE_H - 205
            for title, value in cards:
                c.setFillColor(colors.HexColor("#F2F6FA"))
                c.roundRect(x, card_y, 150, 46, 6, stroke=0, fill=1)
                c.setFillColor(colors.HexColor("#567389"))
                c.setFont("Helvetica", 8)
                c.drawString(x + 10, card_y + 30, title)
                c.setFillColor(colors.black)
                c.setFont("Helvetica-Bold", 12)
                c.drawString(x + 10, card_y + 13, value)
                x += 166

            table_x = 48
            table_top = PAGE_H - 292
            column_widths = [98, 66, 88, 180, 56, 40]
            headers = [
                _tr(raw_config, "timestamp"),
                _tr(raw_config, "site"),
                _tr(raw_config, "equipment_tag"),
                _tr(raw_config, "sensor_name"),
                _tr(raw_config, "value"),
                _tr(raw_config, "unit"),
            ]

            c.setFillColor(accent)
            c.rect(table_x, table_top, sum(column_widths), 24, fill=1, stroke=0)
            c.setFillColor(colors.white)
            c.setFont("Helvetica-Bold", 8)
            cursor = table_x + 4
            for header, width in zip(headers, column_widths):
                c.drawString(cursor, table_top + 8, header)
                cursor += width

            row_y = table_top - 20
            c.setFillColor(colors.black)
            c.setFont("Helvetica", 7)
            for row in rows[page_start:page_start + page_size]:
                c.rect(table_x, row_y, sum(column_widths), 18, fill=0, stroke=1)
                cursor = table_x + 4
                row_values = [
                    row["timestamp"].strftime("%Y-%m-%d %H:%M"),
                    row["site"],
                    row["asset_tag"],
                    row["sensor_name"],
                    f"{row['value']:.2f}",
                    row["unit"],
                ]
                for value, width in zip(row_values, column_widths):
                    c.drawString(cursor, row_y + 5, str(value))
                    cursor += width
                row_y -= 18

            c.setFont("Helvetica", 8)
            c.setFillColor(colors.grey)
            c.drawString(48, 42, _tr(raw_config, "time_series_footer"))

    c.save()
    return buffer.getvalue()


def _generate_bems_time_series_pdf_variant(raw_config: dict, blocks: list[dict], layout_plan: dict, distractor_plan) -> bytes:
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    c.setTitle(raw_config.get("document", {}).get("title", _tr(raw_config, "bems_time_series_title")))
    c.setSubject(raw_config.get("document", {}).get("subject", "Scope 1 stationary combustion BEMS export"))

    interval_minutes = _bems_interval_minutes(raw_config)
    for block_index, block in enumerate(blocks):
        rows = block["rows"]
        page_size = 28
        for page_start in range(0, max(len(rows), 1), page_size):
            if block_index > 0 or page_start > 0:
                c.showPage()

            accent = colors.HexColor("#1E5B88")
            c.setFillColor(accent)
            c.rect(36, PAGE_H - 78, PAGE_W - 72, 34, fill=1, stroke=0)
            c.setFillColor(colors.white)
            c.setFont("Helvetica-Bold", 15)
            c.drawString(48, PAGE_H - 64, _tr(raw_config, "bems_time_series_title"))

            visible_rows = rows[page_start:page_start + page_size]

            def render_meta(current_top: float) -> float:
                c.setFillColor(colors.black)
                c.setFont("Helvetica", 10)
                meta_lines = [
                    f"{_tr(raw_config, 'company')}: {block['company']}",
                    f"{_tr(raw_config, 'site')}: {block['site']}",
                    f"{_tr(raw_config, 'country')}: {block['country']}",
                    f"{_tr(raw_config, 'reporting_period')}: {block['period_label']}",
                    *_stationary_document_lines(distractor_plan),
                ]
                meta_lines = [line for line in meta_lines if not line.endswith(": ")]
                current_top = _draw_multiline(c, 48, current_top, meta_lines, leading=14) - 16

                cards = [
                    (_tr(raw_config, "assets"), str(len(block["assets"]))),
                    (_tr(raw_config, "interval"), f"{interval_minutes} min"),
                    (_tr(raw_config, "rows"), str(len(rows))),
                ]
                x = 48
                card_y = current_top - 46
                for title, value in cards:
                    c.setFillColor(colors.HexColor("#F2F6FA"))
                    c.roundRect(x, card_y, 150, 46, 6, stroke=0, fill=1)
                    c.setFillColor(colors.HexColor("#567389"))
                    c.setFont("Helvetica", 8)
                    c.drawString(x + 10, card_y + 30, title)
                    c.setFillColor(colors.black)
                    c.setFont("Helvetica-Bold", 12)
                    c.drawString(x + 10, card_y + 13, value)
                    x += 166
                return card_y - 18

            def render_table(current_top: float) -> float:
                table_x = 48
                table_top = current_top - 8
                column_widths = [98, 66, 88, 180, 56, 40]
                headers = [
                    _tr(raw_config, "timestamp"),
                    _tr(raw_config, "site"),
                    _tr(raw_config, "equipment_tag"),
                    _tr(raw_config, "sensor_name"),
                    _tr(raw_config, "value"),
                    _tr(raw_config, "unit"),
                ]
                c.setFillColor(accent)
                c.rect(table_x, table_top, sum(column_widths), 24, fill=1, stroke=0)
                c.setFillColor(colors.white)
                c.setFont("Helvetica-Bold", 8)
                cursor = table_x + 4
                for header, width in zip(headers, column_widths):
                    c.drawString(cursor, table_top + 8, header)
                    cursor += width

                row_y = table_top - 20
                c.setFillColor(colors.black)
                c.setFont("Helvetica", 7)
                for row in visible_rows:
                    c.rect(table_x, row_y, sum(column_widths), 18, fill=0, stroke=1)
                    cursor = table_x + 4
                    row_values = [
                        row["timestamp"].strftime("%Y-%m-%d %H:%M"),
                        row["site"],
                        row["asset_tag"],
                        row["sensor_name"],
                        f"{row['value']:.2f}",
                        row["unit"],
                    ]
                    for value, width in zip(row_values, column_widths):
                        c.drawString(cursor, row_y + 5, str(value))
                        cursor += width
                    row_y -= 18
                return row_y - 12

            current_top = PAGE_H - 108
            for section_name in (layout_plan.get("section_order") or ["meta", "table", "footer"]):
                if section_name == "meta":
                    current_top = render_meta(current_top)
                elif section_name == "table":
                    current_top = render_table(current_top)

            c.setFont("Helvetica", 8)
            c.setFillColor(colors.grey)
            c.drawString(48, 42, _tr(raw_config, "time_series_footer"))

    c.save()
    return buffer.getvalue()


def generate_bems_equipment_report_docx(raw_config: dict) -> bytes:
    blocks = _corrupted_bems_site_blocks(raw_config)
    distractor_plan = _stationary_distractor_plan(raw_config, "DOCX", document_type="bems")
    layout_plan = _stationary_layout_plan(raw_config, "DOCX", document_type="bems")
    if layout_plan.get("enabled"):
        return _generate_bems_equipment_report_docx_variant(raw_config, blocks, layout_plan, distractor_plan)

    document = Document()
    core_props = document.core_properties
    core_props.title = raw_config.get("document", {}).get("title", _tr(raw_config, "bems_equipment_title"))
    core_props.subject = raw_config.get("document", {}).get("subject", "Scope 1 stationary combustion BEMS export")

    for block_index, block in enumerate(blocks):
        document.add_heading(_tr(raw_config, "bems_equipment_title"), level=0)

        for line in [
            f"{_tr(raw_config, 'company')}: {block['company']}",
            f"{_tr(raw_config, 'site')}: {block['site']}",
            f"{_tr(raw_config, 'country')}: {block['country']}",
            f"{_tr(raw_config, 'reporting_period')}: {block['period_label']}",
            *_stationary_document_lines(distractor_plan),
        ]:
            if not line.endswith(": "):
                document.add_paragraph(line)

        total_assets = len(block["assets"])
        total_hours = sum(
            Decimal("0") if (asset["operating_hours"] is None or isinstance(asset["operating_hours"], str))
            else asset["operating_hours"]
            for asset in block["assets"]
        )
        dominant_asset = max(
            (a for a in block["assets"] if not isinstance(a["quantity"], str)),
            key=lambda asset: asset["quantity"],
            default=None,
        )
        summary_table = document.add_table(rows=2, cols=3)
        summary_table.style = "Table Grid"
        summary_headers = [_tr(raw_config, "assets"), _tr(raw_config, "operating_hours"), _tr(raw_config, "top_asset")]
        summary_values = [
            str(total_assets),
            _fmt_optional_number(total_hours, " h") or "n/a",
            dominant_asset["asset_tag"] if dominant_asset else "n/a",
        ]
        for cell, value in zip(summary_table.rows[0].cells, summary_headers):
            cell.text = value
        for cell, value in zip(summary_table.rows[1].cells, summary_values):
            cell.text = value

        document.add_paragraph(_tr(raw_config, "equipment_trend_snapshot")).runs[0].bold = True
        rank_table = document.add_table(rows=1, cols=3)
        rank_table.style = "Table Grid"
        for cell, header in zip(rank_table.rows[0].cells, [_tr(raw_config, "equipment_tag"), _tr(raw_config, "consumption"), _tr(raw_config, "unit")]):
            cell.text = header
        for asset in sorted(block["assets"], key=lambda item: item["quantity"], reverse=True)[:5]:
            row = rank_table.add_row().cells
            row[0].text = asset["asset_tag"]
            row[1].text = _fmt_money(asset["quantity"])
            row[2].text = asset["unit"]

        detail_table = document.add_table(rows=1, cols=7)
        detail_table.style = "Table Grid"
        for cell, header in zip(
            detail_table.rows[0].cells,
            [
                _tr(raw_config, "equipment_tag"),
                _tr(raw_config, "equipment_name"),
                _tr(raw_config, "emission_source"),
                _tr(raw_config, "fuel_type"),
                _tr(raw_config, "consumption"),
                _tr(raw_config, "unit"),
                _tr(raw_config, "operating_hours"),
            ],
        ):
            cell.text = header

        for asset in block["assets"]:
            row = detail_table.add_row().cells
            row[0].text = asset["asset_tag"]
            row[1].text = asset["equipment_name"]
            row[2].text = asset["emission_source"]
            row[3].text = asset["fuel"]
            row[4].text = _fmt_money(asset["quantity"])
            row[5].text = asset["unit"]
            row[6].text = _fmt_optional_number(asset["operating_hours"])

        document.add_paragraph(_tr(raw_config, "dashboard_summary_footer"))
        if block_index < len(blocks) - 1:
            document.add_page_break()

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _generate_bems_equipment_report_docx_variant(raw_config: dict, blocks: list[dict], layout_plan: dict, distractor_plan) -> bytes:
    document = Document()
    core_props = document.core_properties
    core_props.title = raw_config.get("document", {}).get("title", _tr(raw_config, "bems_equipment_title"))
    core_props.subject = raw_config.get("document", {}).get("subject", "Scope 1 stationary combustion BEMS export")

    for block_index, block in enumerate(blocks):
        document.add_heading(_tr(raw_config, "bems_equipment_title"), level=0)

        total_assets = len(block["assets"])
        total_hours = sum(
            Decimal("0") if (asset["operating_hours"] is None or isinstance(asset["operating_hours"], str)) else asset["operating_hours"]
            for asset in block["assets"]
        )
        dominant_asset = max(
            (a for a in block["assets"] if not isinstance(a["quantity"], str)),
            key=lambda asset: asset["quantity"],
            default=None,
        )

        def render_meta() -> None:
            for line in [
                f"{_tr(raw_config, 'company')}: {block['company']}",
                f"{_tr(raw_config, 'site')}: {block['site']}",
                f"{_tr(raw_config, 'country')}: {block['country']}",
                f"{_tr(raw_config, 'reporting_period')}: {block['period_label']}",
                *_stationary_document_lines(distractor_plan),
            ]:
                if not line.endswith(": "):
                    document.add_paragraph(line)

            summary_table = document.add_table(rows=2, cols=3)
            summary_table.style = "Table Grid"
            for cell, value in zip(summary_table.rows[0].cells, [_tr(raw_config, "assets"), _tr(raw_config, "operating_hours"), _tr(raw_config, "top_asset")]):
                cell.text = value
            for cell, value in zip(summary_table.rows[1].cells, [str(total_assets), _fmt_optional_number(total_hours, " h") or "n/a", dominant_asset["asset_tag"] if dominant_asset else "n/a"]):
                cell.text = value

        def render_table() -> None:
            document.add_paragraph(_tr(raw_config, "equipment_trend_snapshot")).runs[0].bold = True
            rank_table = document.add_table(rows=1, cols=3)
            rank_table.style = "Table Grid"
            for cell, header in zip(rank_table.rows[0].cells, [_tr(raw_config, "equipment_tag"), _tr(raw_config, "consumption"), _tr(raw_config, "unit")]):
                cell.text = header
            for asset in sorted(block["assets"], key=lambda item: item["quantity"], reverse=True)[:5]:
                row = rank_table.add_row().cells
                row[0].text = asset["asset_tag"]
                row[1].text = _fmt_money(asset["quantity"])
                row[2].text = asset["unit"]

            detail_table = document.add_table(rows=1, cols=7)
            detail_table.style = "Table Grid"
            for cell, header in zip(detail_table.rows[0].cells, [_tr(raw_config, "equipment_tag"), _tr(raw_config, "equipment_name"), _tr(raw_config, "emission_source"), _tr(raw_config, "fuel_type"), _tr(raw_config, "consumption"), _tr(raw_config, "unit"), _tr(raw_config, "operating_hours")]):
                cell.text = header
            for asset in block["assets"]:
                row = detail_table.add_row().cells
                row[0].text = asset["asset_tag"]
                row[1].text = asset["equipment_name"]
                row[2].text = asset["emission_source"]
                row[3].text = asset["fuel"]
                row[4].text = _fmt_money(asset["quantity"])
                row[5].text = asset["unit"]
                row[6].text = _fmt_optional_number(asset["operating_hours"])

        def render_footer() -> None:
            document.add_paragraph(_tr(raw_config, "dashboard_summary_footer"))

        for section_name in (layout_plan.get("section_order") or ["meta", "table", "footer"]):
            if section_name == "meta":
                render_meta()
            elif section_name == "table":
                render_table()
            elif section_name == "footer":
                render_footer()

        if block_index < len(blocks) - 1:
            document.add_page_break()

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def generate_bems_time_series_docx(raw_config: dict) -> bytes:
    blocks = _build_bems_trend_exports(raw_config)
    distractor_plan = _stationary_distractor_plan(raw_config, "DOCX", document_type="bems")
    layout_plan = _stationary_layout_plan(raw_config, "DOCX", document_type="bems")
    if layout_plan.get("enabled"):
        return _generate_bems_time_series_docx_variant(raw_config, blocks, layout_plan, distractor_plan)

    document = Document()
    core_props = document.core_properties
    core_props.title = raw_config.get("document", {}).get("title", _tr(raw_config, "bems_time_series_title"))
    core_props.subject = raw_config.get("document", {}).get("subject", "Scope 1 stationary combustion BEMS export")

    interval_minutes = _bems_interval_minutes(raw_config)
    for block_index, block in enumerate(blocks):
        document.add_heading(_tr(raw_config, "bems_time_series_title"), level=0)

        for line in [
            f"{_tr(raw_config, 'company')}: {block['company']}",
            f"{_tr(raw_config, 'site')}: {block['site']}",
            f"{_tr(raw_config, 'country')}: {block['country']}",
            f"{_tr(raw_config, 'reporting_period')}: {block['period_label']}",
            *_stationary_document_lines(distractor_plan),
        ]:
            if not line.endswith(": "):
                document.add_paragraph(line)

        summary_table = document.add_table(rows=2, cols=3)
        summary_table.style = "Table Grid"
        summary_headers = [_tr(raw_config, "assets"), _tr(raw_config, "interval"), _tr(raw_config, "rows")]
        summary_values = [
            str(len(block["assets"])),
            f"{interval_minutes} min",
            str(len(block["rows"])),
        ]
        for cell, value in zip(summary_table.rows[0].cells, summary_headers):
            cell.text = value
        for cell, value in zip(summary_table.rows[1].cells, summary_values):
            cell.text = value

        detail_table = document.add_table(rows=1, cols=6)
        detail_table.style = "Table Grid"
        for cell, header in zip(
            detail_table.rows[0].cells,
            [
                _tr(raw_config, "timestamp"),
                _tr(raw_config, "site"),
                _tr(raw_config, "equipment_tag"),
                _tr(raw_config, "sensor_name"),
                _tr(raw_config, "value"),
                _tr(raw_config, "unit"),
            ],
        ):
            cell.text = header

        for row_data in block["rows"]:
            row = detail_table.add_row().cells
            row[0].text = row_data["timestamp"].strftime("%Y-%m-%d %H:%M")
            row[1].text = row_data["site"]
            row[2].text = row_data["asset_tag"]
            row[3].text = row_data["sensor_name"]
            row[4].text = f"{row_data['value']:.2f}"
            row[5].text = row_data["unit"]

        document.add_paragraph(_tr(raw_config, "time_series_word_footer"))
        if block_index < len(blocks) - 1:
            document.add_page_break()

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _generate_bems_time_series_docx_variant(raw_config: dict, blocks: list[dict], layout_plan: dict, distractor_plan) -> bytes:
    document = Document()
    core_props = document.core_properties
    core_props.title = raw_config.get("document", {}).get("title", _tr(raw_config, "bems_time_series_title"))
    core_props.subject = raw_config.get("document", {}).get("subject", "Scope 1 stationary combustion BEMS export")

    interval_minutes = _bems_interval_minutes(raw_config)
    for block_index, block in enumerate(blocks):
        document.add_heading(_tr(raw_config, "bems_time_series_title"), level=0)

        def render_meta() -> None:
            for line in [
                f"{_tr(raw_config, 'company')}: {block['company']}",
                f"{_tr(raw_config, 'site')}: {block['site']}",
                f"{_tr(raw_config, 'country')}: {block['country']}",
                f"{_tr(raw_config, 'reporting_period')}: {block['period_label']}",
                *_stationary_document_lines(distractor_plan),
            ]:
                if not line.endswith(": "):
                    document.add_paragraph(line)

            summary_table = document.add_table(rows=2, cols=3)
            summary_table.style = "Table Grid"
            for cell, value in zip(summary_table.rows[0].cells, [_tr(raw_config, "assets"), _tr(raw_config, "interval"), _tr(raw_config, "rows")]):
                cell.text = value
            for cell, value in zip(summary_table.rows[1].cells, [str(len(block["assets"])), f"{interval_minutes} min", str(len(block["rows"]))]):
                cell.text = value

        def render_table() -> None:
            detail_table = document.add_table(rows=1, cols=6)
            detail_table.style = "Table Grid"
            for cell, header in zip(detail_table.rows[0].cells, [_tr(raw_config, "timestamp"), _tr(raw_config, "site"), _tr(raw_config, "equipment_tag"), _tr(raw_config, "sensor_name"), _tr(raw_config, "value"), _tr(raw_config, "unit")]):
                cell.text = header
            for row_data in block["rows"]:
                row = detail_table.add_row().cells
                row[0].text = row_data["timestamp"].strftime("%Y-%m-%d %H:%M")
                row[1].text = row_data["site"]
                row[2].text = row_data["asset_tag"]
                row[3].text = row_data["sensor_name"]
                row[4].text = f"{row_data['value']:.2f}"
                row[5].text = row_data["unit"]

        def render_footer() -> None:
            document.add_paragraph(_tr(raw_config, "time_series_word_footer"))

        for section_name in (layout_plan.get("section_order") or ["meta", "table", "footer"]):
            if section_name == "meta":
                render_meta()
            elif section_name == "table":
                render_table()
            elif section_name == "footer":
                render_footer()

        if block_index < len(blocks) - 1:
            document.add_page_break()

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def generate_bems_time_series_xlsx(raw_config: dict) -> bytes:
    blocks = _build_bems_trend_exports(raw_config)
    layout_plan = _stationary_layout_plan(raw_config, "XLSX", document_type="bems")
    distractor_plan = _stationary_distractor_plan(raw_config, "XLSX", document_type="bems")
    distractor_fields = _stationary_distractor_field_map(distractor_plan)
    ordered_ids = _augment_stationary_field_ids(_ordered_stationary_field_ids(layout_plan, list(_BEMS_TIME_SERIES_HEADER_KEYS)), distractor_plan)
    workbook = openpyxl.Workbook()
    workbook.remove(workbook.active)

    for sheet_index, block in enumerate(blocks, start=1):
        sheet_name = block["site"][:31] or f"{_tr(raw_config, 'site_fallback')} {sheet_index}"
        sheet = workbook.create_sheet(title=sheet_name)
        sheet["A1"] = _tr(raw_config, "bems_time_series_title")
        sheet["A1"].font = Font(size=14, bold=True)
        sheet["A2"] = _tr(raw_config, "company")
        sheet["B2"] = block["company"]
        sheet["A3"] = _tr(raw_config, "site")
        sheet["B3"] = block["site"]
        sheet["A4"] = _tr(raw_config, "country")
        sheet["B4"] = block["country"]
        sheet["A5"] = _tr(raw_config, "reporting_period")
        sheet["B5"] = block["period_label"]

        header_row = _write_stationary_xlsx_preamble(sheet, 6, layout_plan)
        headers = [
            _stationary_header_text(
                raw_config,
                layout_plan,
                field_id,
                _BEMS_TIME_SERIES_HEADER_KEYS.get(field_id, field_id),
                distractor_fields,
            )
            for field_id in ordered_ids
        ]
        header_fill = PatternFill(fill_type="solid", fgColor="1E5B88")
        for column_index, header in enumerate(headers, start=1):
            cell = sheet.cell(row=header_row, column=column_index, value=header)
            cell.font = Font(color="FFFFFF", bold=True)
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center")

        row_index = header_row + 1
        for row in block["rows"]:
            row_map = {
                "timestamp": row["timestamp"].strftime("%Y-%m-%d %H:%M"),
                "site": row["site"],
                "equipment_tag": row["asset_tag"],
                "sensor_name": row["sensor_name"],
                "value": row["value"],
                "unit": row["unit"],
            }
            values = _stationary_row_values(
                row_map,
                ordered_ids,
                distractor_plan,
                row_key=f"{row['timestamp'].isoformat()}:{row['asset_tag']}",
                block_key=f"{block['company']}:{block['site']}",
            )
            for column_index, value in enumerate(values, start=1):
                sheet.cell(row=row_index, column=column_index, value=value)
            row_index += 1

        width_map = {
            "timestamp": 19,
            "site": 20,
            "equipment_tag": 14,
            "sensor_name": 24,
            "value": 12,
            "unit": 8,
        }
        for column_index, field_id in enumerate(ordered_ids, start=1):
            width = width_map.get(field_id, 14)
            sheet.column_dimensions[get_column_letter(column_index)].width = width

    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def generate_bems_time_series_csv(raw_config: dict) -> bytes:
    blocks = _build_bems_trend_exports(raw_config)
    layout_plan = _stationary_layout_plan(raw_config, "CSV", document_type="bems")
    distractor_plan = _stationary_distractor_plan(raw_config, "CSV", document_type="bems")
    distractor_fields = _stationary_distractor_field_map(distractor_plan)
    ordered_ids = _augment_stationary_field_ids(_ordered_stationary_field_ids(layout_plan, list(_BEMS_TIME_SERIES_HEADER_KEYS)), distractor_plan)
    buffer = StringIO()
    writer = csv.writer(buffer)

    for block_index, block in enumerate(blocks):
        if block_index > 0:
            writer.writerow([])
        writer.writerow([_tr(raw_config, "bems_time_series_title")])
        writer.writerow([_tr(raw_config, "company"), block["company"]])
        writer.writerow([_tr(raw_config, "site"), block["site"]])
        writer.writerow([_tr(raw_config, "country"), block["country"]])
        writer.writerow([_tr(raw_config, "reporting_period"), block["period_label"]])
        writer.writerow([])
        _write_stationary_csv_preamble(writer, layout_plan)
        writer.writerow([
            _stationary_header_text(
                raw_config,
                layout_plan,
                field_id,
                _BEMS_TIME_SERIES_HEADER_KEYS.get(field_id, field_id),
                distractor_fields,
            )
            for field_id in ordered_ids
        ])
        for row in block["rows"]:
            row_map = {
                "timestamp": row["timestamp"].strftime("%Y-%m-%d %H:%M"),
                "site": row["site"],
                "equipment_tag": row["asset_tag"],
                "sensor_name": row["sensor_name"],
                "value": f"{row['value']:.2f}",
                "unit": row["unit"],
            }
            writer.writerow(
                _stationary_row_values(
                    row_map,
                    ordered_ids,
                    distractor_plan,
                    row_key=f"{row['timestamp'].isoformat()}:{row['asset_tag']}",
                    block_key=f"{block['company']}:{block['site']}",
                )
            )

    return buffer.getvalue().encode("utf-8-sig")


def generate_bems_equipment_report_xlsx(raw_config: dict) -> bytes:
    blocks = _corrupted_bems_site_blocks(raw_config)
    layout_plan = _stationary_layout_plan(raw_config, "XLSX", document_type="bems")
    distractor_plan = _stationary_distractor_plan(raw_config, "XLSX", document_type="bems")
    distractor_fields = _stationary_distractor_field_map(distractor_plan)
    ordered_ids = _ordered_stationary_field_ids(layout_plan, list(_BEMS_EQUIPMENT_HEADER_KEYS))
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = _tr(raw_config, "bems_summary_sheet_title")
    sheet["A1"] = _tr(raw_config, "bems_equipment_title")
    sheet["A1"].font = Font(size=14, bold=True)

    header_row = _write_stationary_xlsx_preamble(sheet, 2, layout_plan)
    header_keys = {
        "company": "company",
        "site": "site",
        "country": "country",
        "reporting_period": "reporting_period",
        **_BEMS_EQUIPMENT_HEADER_KEYS,
    }
    default_ids = _augment_stationary_field_ids(["company", "site", "country", "reporting_period", *ordered_ids], distractor_plan)
    headers = [
        _stationary_header_text(
            raw_config,
            layout_plan,
            field_id,
            header_keys.get(field_id, field_id),
            distractor_fields,
        )
        for field_id in default_ids
    ]
    header_fill = PatternFill(fill_type="solid", fgColor="1E5B88")
    for column_index, header in enumerate(headers, start=1):
        cell = sheet.cell(row=header_row, column=column_index, value=header)
        cell.font = Font(color="FFFFFF", bold=True)
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")

    row_index = header_row + 1
    for block in blocks:
        for asset in block["assets"]:
            row_map = {
                "company": block["company"],
                "site": block["site"],
                "country": block["country"],
                "reporting_period": block["period_label"],
                "equipment_tag": asset["asset_tag"],
                "equipment_name": asset["equipment_name"],
                "emission_source": asset["emission_source"],
                "fuel_type": asset["fuel"],
                "consumption": _safe_float(asset["quantity"]),
                "unit": asset["unit"],
                "operating_hours": None if asset["operating_hours"] is None else _safe_float(asset["operating_hours"]),
            }
            values = _stationary_row_values(
                row_map,
                default_ids,
                distractor_plan,
                row_key=str(asset["asset_tag"] or "asset-row"),
                block_key=f"{block['company']}:{block['site']}",
            )
            for column_index, value in enumerate(values, start=1):
                sheet.cell(row=row_index, column=column_index, value=value)
            row_index += 1

    width_map = {
        "company": 24,
        "site": 20,
        "country": 16,
        "reporting_period": 18,
        "equipment_tag": 14,
        "equipment_name": 22,
        "emission_source": 20,
        "fuel_type": 16,
        "consumption": 14,
        "unit": 8,
        "operating_hours": 16,
    }
    for column_index, field_id in enumerate(default_ids, start=1):
        width = width_map.get(field_id, 14)
        sheet.column_dimensions[get_column_letter(column_index)].width = width

    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def generate_bems_equipment_report_csv(raw_config: dict) -> bytes:
    blocks = _corrupted_bems_site_blocks(raw_config)
    layout_plan = _stationary_layout_plan(raw_config, "CSV", document_type="bems")
    distractor_plan = _stationary_distractor_plan(raw_config, "CSV", document_type="bems")
    distractor_fields = _stationary_distractor_field_map(distractor_plan)
    ordered_ids = _ordered_stationary_field_ids(layout_plan, list(_BEMS_EQUIPMENT_HEADER_KEYS))
    buffer = StringIO()
    writer = csv.writer(buffer)
    header_keys = {
        "company": "company",
        "site": "site",
        "country": "country",
        "reporting_period": "reporting_period",
        **_BEMS_EQUIPMENT_HEADER_KEYS,
    }
    default_ids = _augment_stationary_field_ids(["company", "site", "country", "reporting_period", *ordered_ids], distractor_plan)
    _write_stationary_csv_preamble(writer, layout_plan)
    writer.writerow([
        _stationary_header_text(
            raw_config,
            layout_plan,
            field_id,
            header_keys.get(field_id, field_id),
            distractor_fields,
        )
        for field_id in default_ids
    ])
    for block in blocks:
        for asset in block["assets"]:
            row_map = {
                "company": block["company"],
                "site": block["site"],
                "country": block["country"],
                "reporting_period": block["period_label"],
                "equipment_tag": asset["asset_tag"],
                "equipment_name": asset["equipment_name"],
                "emission_source": asset["emission_source"],
                "fuel_type": asset["fuel"],
                "consumption": _fmt_num(asset["quantity"], ".2f"),
                "unit": asset["unit"],
                "operating_hours": "" if asset["operating_hours"] is None else _fmt_num(asset["operating_hours"], ".2f"),
            }
            writer.writerow(
                _stationary_row_values(
                    row_map,
                    default_ids,
                    distractor_plan,
                    row_key=str(asset["asset_tag"] or "asset-row"),
                    block_key=f"{block['company']}:{block['site']}",
                )
            )

    return buffer.getvalue().encode("utf-8-sig")
