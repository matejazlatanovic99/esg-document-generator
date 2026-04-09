"""Electricity Consumption Statement generator (Scope 2 – Purchased Electricity).

Produces PDF, XLSX, CSV and DOCX statements.  Each site generates one statement
page (PDF) or one detail row (XLSX / CSV / DOCX).
"""
from __future__ import annotations

import csv as csv_module
import importlib.util as _ilu
import io
import math
import os
import random
from datetime import date, datetime
from decimal import Decimal, ROUND_HALF_UP

# ── shared drawing primitives from the heat PDF generator ────────────────────
_HERE = os.path.dirname(os.path.abspath(__file__))
_spec = _ilu.spec_from_file_location("_heat_gen", os.path.join(_HERE, "pdf-generator.py"))
_heat = _ilu.module_from_spec(_spec)  # type: ignore[arg-type]
_spec.loader.exec_module(_heat)  # type: ignore[union-attr]

FONT_REG  = _heat.FONT_REG
FONT_BOLD = _heat.FONT_BOLD
FONT_MONO = _heat.FONT_MONO

draw_logo            = _heat.draw_logo
draw_info_box        = _heat.draw_info_box
round_box            = _heat.round_box
build_background     = _heat.build_background
build_foreground_noise = _heat.build_foreground_noise
DEFAULT_COMPANY_STYLES = _heat.DEFAULT_COMPANY_STYLES

# Billing period helpers reused from heat generator
_bp_factor      = _heat.billing_period_factor
_bp_dates       = _heat.billing_period_dates
_bp_label       = _heat.billing_period_label
_bp_inv_suffix  = _heat.invoice_suffix
_derive_periods = _heat.derive_month_periods
_norm_periods   = _heat.normalize_billing_periods

from reportlab.lib.colors import HexColor, white  # noqa: E402
from reportlab.lib.pagesizes import A4            # noqa: E402
from reportlab.lib.utils import ImageReader       # noqa: E402
from reportlab.pdfgen import canvas               # noqa: E402

PAGE_W, PAGE_H = A4
TWOPLACES = Decimal("0.01")

# ── translations ──────────────────────────────────────────────────────────────

TRANSLATIONS: dict[str, dict[str, str]] = {
    "en": {
        "logo_subtitle":      "Electricity Supply Services",
        "doc_title_heading":  "Electricity Consumption Statement",
        "doc_subtitle":       "Scope 2 – Purchased Electricity",
        "box_supplier":       "Supplier Details",
        "box_customer":       "Customer / Site Address",
        "box_period":         "Statement Period",
        "meta_period_label":  "Period",
        "meta_period_start":  "Period Start",
        "meta_period_end":    "Period End",
        "meta_ref":           "Reference",
        "meta_currency":      "Currency",
        "tbl_meter":          "Meter & Consumption",
        "tbl_grid":           "Grid & Emissions",
        "row_site":           "Site",
        "row_city":           "City",
        "row_postcode":       "Postcode",
        "row_meter_id":       "Electricity Meter ID",
        "row_unit":           "Measurement Unit",
        "row_start_read":     "Start Meter Reading",
        "row_end_read":       "End Meter Reading",
        "row_total_qty":      "Total Consumption",
        "row_supplier_ef":    "Supplier Emission Factor (kg CO\u2082e/kWh)",
        "row_emissions_kg":   "Total Emissions (kg CO\u2082e)",
        "row_emissions_t":    "Total Emissions (tCO\u2082e)",
        "tbl_tariff":         "Tariff Breakdown",
        "col_tariff_name":    "Tariff / Rate",
        "col_tariff_qty":     "Quantity",
        "col_tariff_unit":    "Unit",
        "col_tariff_rate":    "Unit Cost",
        "col_tariff_cost":    "Cost",
        "box_total":          "Cost Summary",
        "charge_total":       "Total Electricity Cost",
        "footer_note":        "Scope 2 – Purchased Electricity (GHG Protocol, location-based method). Emission factors from national grid operator data.",
        "footer_page":        "Page {page} / {total}",
        # XLSX / DOCX labels
        "xl_meta_period":     "Financial Period",
        "xl_meta_start":      "Period Start",
        "xl_meta_end":        "Period End",
        "xl_meta_generated":  "Generated",
        "xl_sum_company":     "Company",
        "xl_sum_sites":       "Sites",
        "xl_sum_qty":         "Total Consumption",
        "xl_sum_cost":        "Total Cost",
        "xl_sum_emissions_t": "Total tCO\u2082e",
        "xl_sum_grand":       "TOTAL",
        "xl_col_ref":         "Reference",
        "xl_col_company":     "Company",
        "xl_col_site":        "Site",
        "xl_col_city":        "City",
        "xl_col_postcode":    "Postcode",
        "xl_col_meter_id":    "Meter ID",
        "xl_col_supplier_ef": "Supplier EF (kg CO\u2082e/kWh)",
        "xl_col_unit":        "Unit",
        "xl_col_start_read":  "Start Reading",
        "xl_col_end_read":    "End Reading",
        "xl_col_total_qty":   "Total Quantity",
        "xl_col_total_cost":  "Total Cost",
        "xl_col_emissions_kg": "Emissions (kg CO\u2082e)",
        "xl_col_emissions_t": "Emissions (tCO\u2082e)",
        "xl_tariff_company":  "Company",
        "xl_tariff_site":     "Site",
        "xl_tariff_meter":    "Meter ID",
        "xl_tariff_name":     "Tariff Name",
        "xl_tariff_qty":      "Quantity",
        "xl_tariff_unit":     "Unit",
        "xl_tariff_rate":     "Unit Cost",
        "xl_tariff_cost":     "Cost",
    },
    "fr": {
        "logo_subtitle":      "Services de fourniture d\u2019\u00e9lectricit\u00e9",
        "doc_title_heading":  "Relev\u00e9 de consommation \u00e9lectrique",
        "doc_subtitle":       "Scope 2 \u2013 \u00c9lectricit\u00e9 achet\u00e9e",
        "box_supplier":       "D\u00e9tails du fournisseur",
        "box_customer":       "Client / Adresse du site",
        "box_period":         "P\u00e9riode du relev\u00e9",
        "meta_period_label":  "P\u00e9riode",
        "meta_period_start":  "D\u00e9but de p\u00e9riode",
        "meta_period_end":    "Fin de p\u00e9riode",
        "meta_ref":           "R\u00e9f\u00e9rence",
        "meta_currency":      "Devise",
        "tbl_meter":          "Compteur & Consommation",
        "tbl_grid":           "R\u00e9seau & \u00c9missions",
        "row_site":           "Site",
        "row_city":           "Ville",
        "row_postcode":       "Code postal",
        "row_meter_id":       "ID du compteur \u00e9lectrique",
        "row_unit":           "Unit\u00e9 de mesure",
        "row_start_read":     "Relev\u00e9 initial",
        "row_end_read":       "Relev\u00e9 final",
        "row_total_qty":      "Consommation totale",
        "row_supplier_ef":    "Facteur d\u2019\u00e9mission fournisseur (kg CO\u2082e/kWh)",
        "row_emissions_kg":   "\u00c9missions totales (kg CO\u2082e)",
        "row_emissions_t":    "\u00c9missions totales (tCO\u2082e)",
        "tbl_tariff":         "Ventilation par tarif",
        "col_tariff_name":    "Tarif / Taux",
        "col_tariff_qty":     "Quantit\u00e9",
        "col_tariff_unit":    "Unit\u00e9",
        "col_tariff_rate":    "Co\u00fbt unitaire",
        "col_tariff_cost":    "Co\u00fbt",
        "box_total":          "R\u00e9sum\u00e9 des co\u00fbts",
        "charge_total":       "Co\u00fbt total de l\u2019\u00e9lectricit\u00e9",
        "footer_note":        "Scope 2 \u2013 \u00c9lectricit\u00e9 achet\u00e9e (protocole GHG, m\u00e9thode bas\u00e9e sur la localisation).",
        "footer_page":        "Page {page} / {total}",
        "xl_meta_period":     "P\u00e9riode financi\u00e8re",
        "xl_meta_start":      "D\u00e9but de p\u00e9riode",
        "xl_meta_end":        "Fin de p\u00e9riode",
        "xl_meta_generated":  "G\u00e9n\u00e9r\u00e9",
        "xl_sum_company":     "Entreprise",
        "xl_sum_sites":       "Sites",
        "xl_sum_qty":         "Consommation totale",
        "xl_sum_cost":        "Co\u00fbt total",
        "xl_sum_emissions_t": "Total tCO\u2082e",
        "xl_sum_grand":       "TOTAL",
        "xl_col_ref":         "R\u00e9f\u00e9rence",
        "xl_col_company":     "Entreprise",
        "xl_col_site":        "Site",
        "xl_col_city":        "Ville",
        "xl_col_postcode":    "Code postal",
        "xl_col_meter_id":    "ID compteur",
        "xl_col_supplier_ef": "FE fournisseur (kg CO\u2082e/kWh)",
        "xl_col_unit":        "Unit\u00e9",
        "xl_col_start_read":  "Relev\u00e9 initial",
        "xl_col_end_read":    "Relev\u00e9 final",
        "xl_col_total_qty":   "Quantit\u00e9 totale",
        "xl_col_total_cost":  "Co\u00fbt total",
        "xl_col_emissions_kg": "\u00c9missions (kg CO\u2082e)",
        "xl_col_emissions_t": "\u00c9missions (tCO\u2082e)",
        "xl_tariff_company":  "Entreprise",
        "xl_tariff_site":     "Site",
        "xl_tariff_meter":    "ID compteur",
        "xl_tariff_name":     "Nom du tarif",
        "xl_tariff_qty":      "Quantit\u00e9",
        "xl_tariff_unit":     "Unit\u00e9",
        "xl_tariff_rate":     "Co\u00fbt unitaire",
        "xl_tariff_cost":     "Co\u00fbt",
    },
    "de": {
        "logo_subtitle":      "Stromversorgungsdienstleistungen",
        "doc_title_heading":  "Stromverbrauchsabrechnung",
        "doc_subtitle":       "Scope 2 \u2013 Eingekaufter Strom",
        "box_supplier":       "Lieferantendetails",
        "box_customer":       "Kunde / Standortadresse",
        "box_period":         "Abrechnungszeitraum",
        "meta_period_label":  "Zeitraum",
        "meta_period_start":  "Zeitraum Beginn",
        "meta_period_end":    "Zeitraum Ende",
        "meta_ref":           "Referenz",
        "meta_currency":      "W\u00e4hrung",
        "tbl_meter":          "Z\u00e4hler & Verbrauch",
        "tbl_grid":           "Netz & Emissionen",
        "row_site":           "Standort",
        "row_city":           "Stadt",
        "row_postcode":       "Postleitzahl",
        "row_meter_id":       "Stromz\u00e4hler-ID",
        "row_unit":           "Messeinheit",
        "row_start_read":     "Anfangszählerstand",
        "row_end_read":       "Endzählerstand",
        "row_total_qty":      "Gesamtverbrauch",
        "row_supplier_ef":    "Emissionsfaktor Lieferant (kg CO\u2082e/kWh)",
        "row_emissions_kg":   "Gesamtemissionen (kg CO\u2082e)",
        "row_emissions_t":    "Gesamtemissionen (tCO\u2082e)",
        "tbl_tariff":         "Tarifaufschl\u00fcsselung",
        "col_tariff_name":    "Tarif / Satz",
        "col_tariff_qty":     "Menge",
        "col_tariff_unit":    "Einheit",
        "col_tariff_rate":    "Einheitspreis",
        "col_tariff_cost":    "Kosten",
        "box_total":          "Kostenübersicht",
        "charge_total":       "Gesamtstromkosten",
        "footer_note":        "Scope 2 \u2013 Eingekaufter Strom (GHG-Protokoll, standortbasierte Methode). Emissionsfaktoren aus nationalen Netzbetreiberdaten.",
        "footer_page":        "Seite {page} / {total}",
        "xl_meta_period":     "Finanzzeitraum",
        "xl_meta_start":      "Zeitraum Beginn",
        "xl_meta_end":        "Zeitraum Ende",
        "xl_meta_generated":  "Erstellt",
        "xl_sum_company":     "Unternehmen",
        "xl_sum_sites":       "Standorte",
        "xl_sum_qty":         "Gesamtverbrauch",
        "xl_sum_cost":        "Gesamtkosten",
        "xl_sum_emissions_t": "Gesamt tCO\u2082e",
        "xl_sum_grand":       "GESAMT",
        "xl_col_ref":         "Referenz",
        "xl_col_company":     "Unternehmen",
        "xl_col_site":        "Standort",
        "xl_col_city":        "Stadt",
        "xl_col_postcode":    "Postleitzahl",
        "xl_col_meter_id":    "Z\u00e4hler-ID",
        "xl_col_supplier_ef": "EF Lieferant (kg CO\u2082e/kWh)",
        "xl_col_unit":        "Einheit",
        "xl_col_start_read":  "Anfangszählerstand",
        "xl_col_end_read":    "Endzählerstand",
        "xl_col_total_qty":   "Gesamtmenge",
        "xl_col_total_cost":  "Gesamtkosten",
        "xl_col_emissions_kg": "Emissionen (kg CO\u2082e)",
        "xl_col_emissions_t": "Emissionen (tCO\u2082e)",
        "xl_tariff_company":  "Unternehmen",
        "xl_tariff_site":     "Standort",
        "xl_tariff_meter":    "Z\u00e4hler-ID",
        "xl_tariff_name":     "Tarifname",
        "xl_tariff_qty":      "Menge",
        "xl_tariff_unit":     "Einheit",
        "xl_tariff_rate":     "Einheitspreis",
        "xl_tariff_cost":     "Kosten",
    },
    "nl": {
        "logo_subtitle":      "Elektriciteitsleveringsdiensten",
        "doc_title_heading":  "Verklaring elektriciteitsverbruik",
        "doc_subtitle":       "Scope 2 \u2013 Ingekochte elektriciteit",
        "box_supplier":       "Leveranciersgegevens",
        "box_customer":       "Klant / Locatieadres",
        "box_period":         "Opgaveperiode",
        "meta_period_label":  "Periode",
        "meta_period_start":  "Begin periode",
        "meta_period_end":    "Einde periode",
        "meta_ref":           "Referentie",
        "meta_currency":      "Valuta",
        "tbl_meter":          "Meter & Verbruik",
        "tbl_grid":           "Net & Emissies",
        "row_site":           "Locatie",
        "row_city":           "Stad",
        "row_postcode":       "Postcode",
        "row_meter_id":       "Elektriciteitmeter-ID",
        "row_unit":           "Meeteenheid",
        "row_start_read":     "Beginmeterstand",
        "row_end_read":       "Eindmeterstand",
        "row_total_qty":      "Totaal verbruik",
        "row_supplier_ef":    "Emissiefactor leverancier (kg CO\u2082e/kWh)",
        "row_emissions_kg":   "Totale emissies (kg CO\u2082e)",
        "row_emissions_t":    "Totale emissies (tCO\u2082e)",
        "tbl_tariff":         "Tariefuitsplitsing",
        "col_tariff_name":    "Tarief",
        "col_tariff_qty":     "Hoeveelheid",
        "col_tariff_unit":    "Eenheid",
        "col_tariff_rate":    "Eenheidsprijs",
        "col_tariff_cost":    "Kosten",
        "box_total":          "Kostenoverzicht",
        "charge_total":       "Totale elektriciteitskosten",
        "footer_note":        "Scope 2 \u2013 Ingekochte elektriciteit (GHG-protocol, locatiegebaseerde methode). Emissiefactoren uit nationale netbeheerdergegevens.",
        "footer_page":        "Pagina {page} / {total}",
        "xl_meta_period":     "Financi\u00eble periode",
        "xl_meta_start":      "Begin periode",
        "xl_meta_end":        "Einde periode",
        "xl_meta_generated":  "Gegenereerd",
        "xl_sum_company":     "Bedrijf",
        "xl_sum_sites":       "Locaties",
        "xl_sum_qty":         "Totaal verbruik",
        "xl_sum_cost":        "Totale kosten",
        "xl_sum_emissions_t": "Totaal tCO\u2082e",
        "xl_sum_grand":       "TOTAAL",
        "xl_col_ref":         "Referentie",
        "xl_col_company":     "Bedrijf",
        "xl_col_site":        "Locatie",
        "xl_col_city":        "Stad",
        "xl_col_postcode":    "Postcode",
        "xl_col_meter_id":    "Meter-ID",
        "xl_col_supplier_ef": "EF leverancier (kg CO\u2082e/kWh)",
        "xl_col_unit":        "Eenheid",
        "xl_col_start_read":  "Beginmeterstand",
        "xl_col_end_read":    "Eindmeterstand",
        "xl_col_total_qty":   "Totale hoeveelheid",
        "xl_col_total_cost":  "Totale kosten",
        "xl_col_emissions_kg": "Emissies (kg CO\u2082e)",
        "xl_col_emissions_t": "Emissies (tCO\u2082e)",
        "xl_tariff_company":  "Bedrijf",
        "xl_tariff_site":     "Locatie",
        "xl_tariff_meter":    "Meter-ID",
        "xl_tariff_name":     "Tariefnaam",
        "xl_tariff_qty":      "Hoeveelheid",
        "xl_tariff_unit":     "Eenheid",
        "xl_tariff_rate":     "Eenheidsprijs",
        "xl_tariff_cost":     "Kosten",
    },
}


# ── helpers ───────────────────────────────────────────────────────────────────

def _q2(value) -> Decimal:
    if not isinstance(value, Decimal):
        value = Decimal(str(value))
    return value.quantize(TWOPLACES, rounding=ROUND_HALF_UP)


def _parse_decimal(value) -> Decimal:
    if isinstance(value, Decimal):
        return value
    return Decimal(str(value))


def _currency_symbol(currency_raw: str) -> str:
    mapping = {"(£)": "£", "(€)": "€", "($)": "$", "(¥)": "¥", "(kr)": "kr", "(Ft)": "Ft"}
    for token, sym in mapping.items():
        if token in currency_raw:
            return sym


def _split_among_tariffs(
    total_qty: Decimal, total_cost: Decimal,
    tariff_names: list[str], rng: random.Random,
) -> list[dict]:
    """Randomly split annual totals among named tariffs.

    Each tariff gets a different effective unit‐rate because quantity and cost
    are split using independent random weights.
    """
    n = len(tariff_names)
    if n == 0:
        return []
    if n == 1:
        uc = _q2(total_cost / total_qty) if total_qty > 0 else Decimal("0")
        return [{"name": tariff_names[0], "quantity": total_qty,
                 "unit_cost": uc, "cost": total_cost}]

    # Independent random weights for qty and cost so unit rates differ
    qty_weights  = [rng.random() for _ in range(n)]
    cost_weights = [w * rng.uniform(0.7, 1.3) for w in qty_weights]
    qty_total_w  = sum(qty_weights)
    cost_total_w = sum(cost_weights)

    qtys  = [_q2(total_qty  * Decimal(str(w / qty_total_w)))  for w in qty_weights]
    costs = [_q2(total_cost * Decimal(str(w / cost_total_w))) for w in cost_weights]
    # Correct last entry so sums are exact
    qtys[-1]  = _q2(total_qty  - sum(qtys[:-1]))
    costs[-1] = _q2(total_cost - sum(costs[:-1]))

    result = []
    for name, qty, cost in zip(tariff_names, qtys, costs):
        uc = _q2(cost / qty) if qty > 0 else Decimal("0")
        result.append({"name": name, "quantity": qty, "unit_cost": uc, "cost": cost})
    return result


def _distribute_annual(annual: Decimal, factors: list, rng: random.Random, jitter: float = 0.07) -> list:  # list[Decimal]
    """Distribute *annual* total proportionally across billing periods with light jitter.

    The last period is adjusted so the sum equals *annual* exactly.
    """
    total_f = sum(factors)
    if not factors or total_f == 0:
        return []
    distributed: list[Decimal] = []
    for idx, f in enumerate(factors):
        raw = annual * Decimal(str(f)) / Decimal(str(total_f))
        variation = Decimal(str(1 + rng.uniform(-jitter, jitter)))
        distributed.append(_q2(raw * variation))
    # Correct last period so totals match exactly
    distributed[-1] = _q2(annual - sum(distributed[:-1]))
    return distributed


def _generate_elec_period_records(annual_site: dict, billing_periods: list, rng: random.Random) -> list[dict]:
    """Expand an annual normalized site into per-billing-period records."""
    factors = [Decimal(str(_bp_factor(p))) for p in billing_periods]

    annual_qty  = annual_site["total_quantity"]
    annual_cost = annual_site["total_cost"]
    supplier_ef = annual_site["supplier_ef"]
    has_supplier_ef = annual_site["_has_supplier_ef"]
    unit        = annual_site["unit"]

    qtys  = _distribute_annual(annual_qty,  factors, rng)
    costs = _distribute_annual(annual_cost, factors, rng)

    # Distribute each tariff proportionally (same seasonal factors)
    annual_tariffs = annual_site["tariffs"]
    tariff_qty_by_period:  list[list[Decimal]] = []
    tariff_cost_by_period: list[list[Decimal]] = []
    for t in annual_tariffs:
        tariff_qty_by_period.append(_distribute_annual(t["quantity"], factors, rng))
        tariff_cost_by_period.append(_distribute_annual(t["cost"],    factors, rng))

    prev = annual_site["start_reading"]
    records: list[dict] = []
    for pi, period in enumerate(billing_periods):
        first, last = _bp_dates(period)
        period_label = _bp_label(period)
        qty  = qtys[pi]
        cost = costs[pi]

        curr = prev + int(qty)

        if has_supplier_ef:
            qty_kwh      = qty * Decimal("1000") if unit == "MWh" else qty
            emissions_kg = _q2(qty_kwh * supplier_ef)
            emissions_t  = (emissions_kg / Decimal("1000")).quantize(
                Decimal("1.000"), rounding=ROUND_HALF_UP
            )
        else:
            emissions_kg = Decimal("0")
            emissions_t  = Decimal("0")

        period_tariffs = [
            {
                "name":      t["name"],
                "quantity":  tariff_qty_by_period[k][pi],
                "unit":      unit,
                "unit_cost": t["unit_cost"],  # unit price stays the same
                "cost":      tariff_cost_by_period[k][pi],
            }
            for k, t in enumerate(annual_tariffs)
        ]

        suffix = _bp_inv_suffix(period, pi + 1)
        ref_no = f"{annual_site['_base_ref_no']}-{suffix}"

        records.append({
            # static site fields
            "label":            annual_site["label"],
            "customer":         annual_site["customer"],
            "customer_code":    annual_site["customer_code"],
            "customer_address": annual_site["customer_address"],
            "city":             annual_site["city"],
            "postcode":         annual_site["postcode"],
            "meter_id":         annual_site["meter_id"],
            "supplier_ef":      supplier_ef,
            "unit":             unit,
            "currency_symbol":  annual_site["currency_symbol"],
            "_omit":            annual_site["_omit"],
            "_site_uid":        annual_site["_site_uid"],
            # per-period fields
            "start_reading":       prev,
            "end_reading":         curr,
            "total_quantity":      qty,
            "total_cost":          cost,
            "emissions_kg":        emissions_kg,
            "emissions_t":         emissions_t,
            "tariffs":             period_tariffs,
            "period_start":        first,
            "period_end":          last,
            "billing_period_label": period_label,
            "ref_no":              ref_no,
        })
        prev = curr
    return records


# ── data normalisation ────────────────────────────────────────────────────────

def normalize_company(company: dict, financial_period: dict, company_index: int) -> dict:
    style = DEFAULT_COMPANY_STYLES[company_index % len(DEFAULT_COMPANY_STYLES)]
    return {
        "label": company["label"],
        "supplier": company["supplier"],
        "supplier_code": company["supplier_code"],
        "supplier_address": company["supplier_address"],
        "customer": company["customer"],
        "customer_code": company["customer_code"],
        "accent": company.get("accent", style["accent"]),
        "accent_soft": company.get("accent_soft", style["accent_soft"]),
        "skew": float(company.get("skew", style["skew"])),
        "currency": company.get("currency", "GBP (£)"),
        "sites": [normalize_site(company, site, financial_period) for site in company["sites"]],
    }


def normalize_site(company: dict, site: dict, financial_period: dict) -> dict:
    """Return a normalised site dict with *period_records* for each billing period."""
    site_omit  = site.get("_omit", {})
    total_qty  = _parse_decimal(site["total_quantity"])
    total_cost = _parse_decimal(site["total_cost"])
    unit       = site.get("unit", "kWh")

    # Determine whether a valid supplier emission factor is present and not omitted.
    supplier_ef_raw = site.get("supplier_ef")
    has_supplier_ef = (
        not site_omit.get("supplier_ef")
        and supplier_ef_raw not in (None, "", "0", 0, "0.0", 0.0)
    )
    supplier_ef = _parse_decimal(supplier_ef_raw) if has_supplier_ef else Decimal("0")

    start_reading = int(site["start_reading"])

    period_start = financial_period["start_date"]
    base_ref_no = (
        f"{company['supplier_code']}-{company['customer_code']}"
        f"-ELEC-{period_start.year}"
    )

    symbol   = _currency_symbol(company.get("currency", "GBP (£)"))
    site_uid = f"{company['label']}|{site.get('label', site.get('meter_id', '?'))}"
    rng_site = random.Random(hash(site_uid) & 0xFFFFFFFF)

    # Build resolved tariff list.
    raw_tariffs = site.get("tariffs", [])
    if raw_tariffs and "quantity" not in raw_tariffs[0]:
        tariff_names = [str(t.get("name", "")).strip() for t in raw_tariffs
                        if str(t.get("name", "")).strip()]
        tariffs = _split_among_tariffs(total_qty, _q2(total_cost), tariff_names, rng_site)
        for t in tariffs:
            t["unit"] = unit
    else:
        tariffs = []
        for t in raw_tariffs:
            if not str(t.get("name", "")).strip():
                continue
            tariffs.append({
                "name":      str(t["name"]),
                "quantity":  _parse_decimal(t.get("quantity", 0)),
                "unit":      unit,
                "unit_cost": _parse_decimal(t.get("unit_cost", 0)),
                "cost":      _q2(_parse_decimal(t.get("cost", 0))),
            })

    # Build the annual "base" dict passed to period record generator
    annual_site = {
        "label":            site.get("label", site["meter_id"]),
        "customer":         site.get("customer", company["customer"]),
        "customer_code":    site.get("customer_code", company["customer_code"]),
        "customer_address": site["customer_address"],
        "city":             site["city"],
        "postcode":         site["postcode"],
        "meter_id":         site["meter_id"],
        "supplier_ef":      supplier_ef,
        "_has_supplier_ef": has_supplier_ef,
        "unit":             unit,
        "start_reading":    start_reading,
        "total_quantity":   total_qty,
        "total_cost":       _q2(total_cost),
        "tariffs":          tariffs,
        "currency_symbol":  symbol,
        "_omit":            site_omit,
        "_base_ref_no":     base_ref_no,
        "_site_uid":        site_uid,
    }

    # Determine billing periods (custom per-site or derived from financial period)
    raw_periods = site.get("billing_periods")
    if raw_periods is not None:
        billing_periods = _norm_periods(raw_periods, financial_period["start_date"].year)
    else:
        billing_periods = _derive_periods(financial_period["start_date"], financial_period["end_date"])

    rng_site.seed(hash(site_uid) & 0xFFFFFFFF)  # re-seed so period distribution is independent of tariff split
    period_records = _generate_elec_period_records(annual_site, billing_periods, rng_site)

    return {**annual_site, "period_records": period_records}


def build_sections(config: dict) -> list[dict]:
    """Return one section per billing period per site (one page per section)."""
    sections = []
    for company in config["companies"]:
        for site in company["sites"]:
            for record in site["period_records"]:
                sections.append({"company": company, "site": record})
    return sections


# ── PDF drawing helpers ───────────────────────────────────────────────────────

def _draw_kv_table(c, x, y, w, title, rows, accent, accent_soft, row_h=17):
    """Key-value table with a coloured title bar."""
    n = len(rows)
    total_h = row_h * (n + 1)
    round_box(c, x, y, w, total_h)
    c.saveState()

    # title bar
    c.setFillColor(HexColor(accent_soft))
    c.roundRect(x, y + total_h - row_h, w, row_h, 4, stroke=0, fill=1)
    c.rect(x, y + total_h - row_h, w, row_h // 2, stroke=0, fill=1)
    c.setFillColor(HexColor(accent))
    c.setFont(FONT_BOLD, 8.1)
    c.drawString(x + 8, y + total_h - 12.2, title)

    split = x + w * 0.56
    c.setStrokeColor(HexColor("#D5DADF"))
    for i in range(n + 1):
        c.line(x, y + i * row_h, x + w, y + i * row_h)
    c.line(split, y, split, y + total_h)

    cy = y + total_h - row_h - 12.5
    for row in rows:
        field, value = row[0], row[1]
        use_mono = row[2] if len(row) > 2 else False
        c.setFillColor(HexColor("#5A6066"))
        c.setFont(FONT_REG, 7.5)
        c.drawString(x + 8, cy, field)
        c.setFillColor(HexColor("#1F2328"))
        c.setFont(FONT_MONO if use_mono else FONT_REG, 8.0)
        val = str(value)
        c.drawString(split + 8, cy, val[:45] + ("\u2026" if len(val) > 45 else ""))
        cy -= row_h
    c.restoreState()


def _draw_tariff_table(c, x, y, w, title, tariffs, accent, accent_soft, strings, symbol, row_h=17):
    """Multi-column tariff breakdown table. Returns the total height drawn."""
    n_data = len(tariffs)
    # rows: title-bar + column-header + data rows
    total_h = row_h * (n_data + 2)
    col_pct = [0.32, 0.18, 0.14, 0.18, 0.18]
    col_widths = [w * p for p in col_pct]

    round_box(c, x, y, w, total_h)
    c.saveState()

    # title bar (top row)
    title_y = y + total_h - row_h
    c.setFillColor(HexColor(accent_soft))
    c.roundRect(x, title_y, w, row_h, 4, stroke=0, fill=1)
    c.rect(x, title_y, w, row_h // 2, stroke=0, fill=1)
    c.setFillColor(HexColor(accent))
    c.setFont(FONT_BOLD, 8.1)
    c.drawString(x + 8, title_y + 4.8, title)

    # column-header row
    hdr_y = y + total_h - row_h * 2
    c.setFillColor(HexColor(accent_soft))
    c.rect(x, hdr_y, w, row_h, stroke=0, fill=1)
    hdrs = [
        strings["col_tariff_name"],
        strings["col_tariff_qty"],
        strings["col_tariff_unit"],
        strings["col_tariff_rate"],
        strings["col_tariff_cost"],
    ]
    cx = x
    for hdr, cw in zip(hdrs, col_widths):
        c.setFillColor(HexColor(accent))
        c.setFont(FONT_BOLD, 7.0)
        c.drawString(cx + 5, hdr_y + 5, hdr)
        cx += cw

    # grid lines
    c.setStrokeColor(HexColor("#D5DADF"))
    for i in range(n_data + 2):
        c.line(x, y + i * row_h, x + w, y + i * row_h)
    cx = x
    for cw in col_widths[:-1]:
        cx += cw
        c.line(cx, y, cx, y + total_h - row_h)  # skip title bar

    # data rows
    dy = hdr_y - 12.5
    for t in tariffs:
        vals = [
            t["name"],
            f"{float(t['quantity']):,.2f}",
            t["unit"],
            f"{symbol}{float(t['unit_cost']):.4f}",
            f"{symbol}{float(t['cost']):,.2f}",
        ]
        cx = x
        for val, cw in zip(vals, col_widths):
            c.setFillColor(HexColor("#1F2328"))
            c.setFont(FONT_REG, 7.5)
            c.drawString(cx + 5, dy, val)
            cx += cw
        dy -= row_h

    c.restoreState()
    return total_h


def _draw_cost_box(c, x, y, w, h, site, accent, accent_soft, strings):
    draw_info_box(c, x, y, w, h, strings["box_total"], [], accent, accent_soft)
    symbol = site.get("currency_symbol", "")
    total_str = f"{symbol}{float(site['total_cost']):,.2f}"
    c.saveState()
    sy = y + h - 30
    c.setFillColor(HexColor(accent_soft))
    c.roundRect(x + 8, sy - 8, w - 16, 24, 3, stroke=0, fill=1)
    c.setFillColor(HexColor(accent))
    c.setFont(FONT_BOLD, 10)
    c.drawString(x + 14, sy + 5, strings["charge_total"])
    c.drawRightString(x + w - 14, sy + 5, total_str)
    c.restoreState()


def _render_page(c, company, site, page_no, total_pages, bg_path, fg_path, strings, noise_level=1.0):
    accent      = company["accent"]
    accent_soft = company["accent_soft"]
    margin      = 32
    content_w   = PAGE_W - margin * 2

    c.drawImage(ImageReader(bg_path), 0, 0, width=PAGE_W, height=PAGE_H, mask="auto")
    c.saveState()
    c.translate(PAGE_W / 2, PAGE_H / 2)
    c.rotate(company["skew"] * noise_level + random.choice([-0.04, 0.03, 0.05]) * noise_level)
    c.translate(-PAGE_W / 2, -PAGE_H / 2)

    # ── header ────────────────────────────────────────────────────────────────
    draw_logo(c, margin, PAGE_H - 72, accent, company["supplier"], strings)
    c.setFillColor(HexColor("#1E2328"))
    c.setFont(FONT_BOLD, 15)
    c.drawRightString(PAGE_W - margin, PAGE_H - 50, strings["doc_title_heading"])
    c.setFont(FONT_REG, 8.2)
    c.drawRightString(PAGE_W - margin, PAGE_H - 64, strings["doc_subtitle"])
    c.setFont(FONT_REG, 7.2)
    c.drawRightString(
        PAGE_W - margin, PAGE_H - 77,
        f"{site['billing_period_label']} \u2022 {company['label']} \u2022 {site['label']}",
    )

    # ── supplier / customer boxes ─────────────────────────────────────────────
    top_y = PAGE_H - 170
    box_w = (content_w - 12) / 2
    draw_info_box(c, margin, top_y, box_w, 92,
                  strings["box_supplier"], company["supplier_address"], accent, accent_soft)
    draw_info_box(c, margin + box_w + 12, top_y, box_w, 92,
                  strings["box_customer"], site["customer_address"], accent, accent_soft)

    # ── period metadata box ───────────────────────────────────────────────────
    period_lines = [
        f"{strings['meta_period_label']}: {site['billing_period_label']}",
        f"{strings['meta_period_start']}: {site['period_start'].strftime('%d %b %Y')}",
        f"{strings['meta_period_end']}: {site['period_end'].strftime('%d %b %Y')}",
        f"{strings['meta_ref']}: {site['ref_no']}",
        f"{strings['meta_currency']}: {company['currency']}",
    ]
    draw_info_box(c, margin, top_y - 108, content_w, 84,
                  strings["box_period"], period_lines, accent, accent_soft)

    # ── meter & grid tables (side by side) ────────────────────────────────────
    half_w    = (content_w - 12) / 2
    table_top = top_y - 108 - 12  # gap below meta box

    unit = site["unit"]
    meter_rows = [
        (strings["row_site"],      site["label"]),
        (strings["row_city"],      site["city"]),
        (strings["row_postcode"],  site["postcode"]),
        (strings["row_meter_id"],  site["meter_id"], True),
        (strings["row_unit"],      unit),
        (strings["row_start_read"], f"{site['start_reading']:,}"),
        (strings["row_end_read"],  f"{site['end_reading']:,}"),
        (strings["row_total_qty"], f"{float(site['total_quantity']):,.2f} {unit}"),
    ]
    grid_rows = [
        (strings["row_supplier_ef"],  f"{float(site['supplier_ef']):.4f}"),
        (strings["row_emissions_kg"], f"{float(site['emissions_kg']):,.2f}"),
        (strings["row_emissions_t"],  f"{float(site['emissions_t']):.3f}"),
    ]

    meter_h = (len(meter_rows) + 1) * 17
    grid_h  = (len(grid_rows)  + 1) * 17
    meter_y = table_top - meter_h
    grid_y  = table_top - grid_h

    _draw_kv_table(c, margin, meter_y, half_w,
                   strings["tbl_meter"], meter_rows, accent, accent_soft)
    _draw_kv_table(c, margin + half_w + 12, grid_y, half_w,
                   strings["tbl_grid"], grid_rows, accent, accent_soft)

    # ── tariff breakdown ──────────────────────────────────────────────────────
    cursor_y = min(meter_y, grid_y) - 10  # below whichever table is taller
    symbol = site.get("currency_symbol", "")

    if site["tariffs"]:
        tariff_h = _draw_tariff_table(
            c, margin, cursor_y - (len(site["tariffs"]) + 2) * 17,
            content_w, strings["tbl_tariff"], site["tariffs"],
            accent, accent_soft, strings, symbol,
        )
        cursor_y = cursor_y - tariff_h - 10

    # ── cost summary ──────────────────────────────────────────────────────────
    cost_h = 55
    _draw_cost_box(c, margin, cursor_y - cost_h, content_w, cost_h,
                   site, accent, accent_soft, strings)

    c.restoreState()

    # ── footer ────────────────────────────────────────────────────────────────
    c.setStrokeColor(HexColor("#C9CDD2"))
    c.line(margin, 42, PAGE_W - margin, 42)
    c.setFillColor(HexColor("#5A6066"))
    c.setFont(FONT_REG, 6.0)
    c.drawString(margin, 25, strings["footer_note"])
    c.drawRightString(PAGE_W - margin, 25,
                      strings["footer_page"].format(page=page_no, total=total_pages))

    if noise_level > 0 and os.path.exists(fg_path):
        c.drawImage(ImageReader(fg_path), 0, 0, width=PAGE_W, height=PAGE_H, mask="auto")
    c.showPage()


# ── PDF public API ────────────────────────────────────────────────────────────

def render_pdf(config: dict, sections: list[dict], output_path: str, noise_level: float = 1.0):
    lang    = config["document"].get("language", "en")
    strings = TRANSLATIONS.get(lang, TRANSLATIONS["en"])
    seed    = config.get("random_seed", 42)
    bg_dir  = config["document"].get("background_dir", "/tmp")

    # Build one background per unique accent colour
    bg_paths: dict[str, str] = {}
    fg_paths: dict[str, str] = {}
    for idx, section in enumerate(sections):
        key = section["company"]["accent"]
        if key not in bg_paths:
            bp = os.path.join(bg_dir, f"elec_bg_{idx}.jpg")
            fp = os.path.join(bg_dir, f"elec_fg_{idx}.png")
            build_background(bp, accent=key, seed=seed + idx, noise_level=noise_level)
            build_foreground_noise(fp, seed=seed + idx, noise_level=noise_level)
            bg_paths[key] = bp
            fg_paths[key] = fp

    total_pages = len(sections)
    c = canvas.Canvas(output_path, pagesize=A4)
    c.setTitle(config["document"].get("title", "Electricity Consumption Statement"))
    c.setSubject(config["document"].get("subject", "Scope 2 Electricity"))
    c.setAuthor("ESG Document Generator")

    for page_no, section in enumerate(sections, start=1):
        company = section["company"]
        site    = section["site"]
        random.seed(seed + page_no)
        _render_page(
            c, company, site, page_no, total_pages,
            bg_paths[company["accent"]], fg_paths[company["accent"]],
            strings, noise_level=noise_level,
        )
    c.save()


# ── XLSX public API ───────────────────────────────────────────────────────────

def generate_xlsx(config: dict, sections: list[dict]) -> bytes:
    import openpyxl
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    lang    = config["document"].get("language", "en")
    T       = TRANSLATIONS.get(lang, TRANSLATIONS["en"])
    fp      = config["financial_period"]

    acc_hex = sections[0]["company"]["accent"] if sections else "#1E5B88"
    acc_r, acc_g, acc_b = (int(acc_hex[i:i+2], 16) for i in (1, 3, 5))

    def _hdr_fill(hex_color: str) -> PatternFill:
        r, g, b = (int(hex_color[i:i+2], 16) for i in (1, 3, 5))
        return PatternFill("solid", fgColor=f"{r:02X}{g:02X}{b:02X}")

    def _hdr_font(white_text: bool = True) -> Font:
        return Font(name="Calibri", bold=True, color="FFFFFF" if white_text else "1F2328", size=9)

    def _border() -> Border:
        s = Side(style="thin", color="D5DADF")
        return Border(left=s, right=s, top=s, bottom=s)

    wb = openpyxl.Workbook()

    # ── Summary sheet ─────────────────────────────────────────────────────────
    ws_sum = wb.active
    ws_sum.title = "Summary"

    ws_sum["A1"] = T["xl_meta_period"]
    ws_sum["B1"] = fp["label"]
    ws_sum["A2"] = T["xl_meta_start"]
    ws_sum["B2"] = fp["start_date"].strftime("%d %b %Y") if hasattr(fp["start_date"], "strftime") else str(fp["start_date"])
    ws_sum["A3"] = T["xl_meta_end"]
    ws_sum["B3"] = fp["end_date"].strftime("%d %b %Y") if hasattr(fp["end_date"], "strftime") else str(fp["end_date"])
    ws_sum["A4"] = T["xl_meta_generated"]
    ws_sum["B4"] = datetime.now().strftime("%Y-%m-%d %H:%M")
    for r in range(1, 5):
        ws_sum.cell(r, 1).font = Font(name="Calibri", bold=True, size=9)

    sum_hdrs = [T["xl_sum_company"], T["xl_sum_sites"], T["xl_sum_qty"], T["xl_sum_cost"], T["xl_sum_emissions_t"]]
    hdr_fill = _hdr_fill(acc_hex)
    hdr_font = _hdr_font()
    brd = _border()

    hdr_row = 6
    for ci, hdr in enumerate(sum_hdrs, start=1):
        cell = ws_sum.cell(hdr_row, ci, hdr)
        cell.fill = hdr_fill
        cell.font = hdr_font
        cell.border = brd
        cell.alignment = Alignment(horizontal="center")

    # Group by company
    from collections import defaultdict
    by_company: dict[str, list] = defaultdict(list)
    for sec in sections:
        by_company[sec["company"]["label"]].append(sec)

    grand_qty = Decimal("0")
    grand_cost = Decimal("0")
    grand_t = Decimal("0")

    dr = hdr_row + 1
    for co_label, secs in by_company.items():
        n_sites    = len({s["site"]["_site_uid"] for s in secs})
        total_qty  = sum(s["site"]["total_quantity"] for s in secs)
        total_cost = sum(s["site"]["total_cost"] for s in secs)
        total_t    = sum(s["site"]["emissions_t"] for s in secs)
        symbol     = secs[0]["site"].get("currency_symbol", "")

        row_data = [co_label, n_sites,
                    float(total_qty), float(total_cost), float(total_t)]
        for ci, val in enumerate(row_data, start=1):
            cell = ws_sum.cell(dr, ci, val)
            cell.border = brd
            cell.font = Font(name="Calibri", size=9)
            if ci >= 3:
                cell.number_format = "#,##0.00"
        dr += 1
        grand_qty  += total_qty
        grand_cost += total_cost
        grand_t    += total_t

    # Grand total row
    grand_fill = PatternFill("solid", fgColor=f"{acc_r:02X}{acc_g:02X}{acc_b:02X}")
    grand_data = [T["xl_sum_grand"], len({s["site"]["_site_uid"] for s in sections}), float(grand_qty), float(grand_cost), float(grand_t)]
    for ci, val in enumerate(grand_data, start=1):
        cell = ws_sum.cell(dr, ci, val)
        cell.fill = grand_fill
        cell.font = Font(name="Calibri", bold=True, color="FFFFFF", size=9)
        cell.border = brd
        if ci >= 3:
            cell.number_format = "#,##0.00"

    for ci, w in enumerate([30, 8, 18, 18, 14], start=1):
        ws_sum.column_dimensions[get_column_letter(ci)].width = w

    # ── Detail sheet ──────────────────────────────────────────────────────────
    ws_det = wb.create_sheet("Detail")

    # Determine the max number of tariffs across all sections for dynamic columns
    max_tariffs = max((len(s["site"].get("tariffs", [])) for s in sections), default=0)

    det_hdrs = [
        T["xl_col_ref"],    T["xl_col_company"], T["xl_col_site"],
        T.get("xl_col_period", "Billing Period"),
        T["xl_col_city"],   T["xl_col_postcode"], T["xl_col_meter_id"],
        T["xl_col_supplier_ef"], T["xl_col_unit"],
        T["xl_col_start_read"], T["xl_col_end_read"],
        T["xl_col_total_qty"], T["xl_col_total_cost"],
        T["xl_col_emissions_kg"], T["xl_col_emissions_t"],
    ]
    for k in range(max_tariffs):
        prefix = f"Tariff {k + 1}"
        det_hdrs += [
            f"{prefix}: {T['xl_tariff_name']}",
            f"{prefix}: {T['xl_tariff_qty']}",
            f"{prefix}: {T['xl_tariff_rate']}",
            f"{prefix}: {T['xl_tariff_cost']}",
        ]

    for ci, hdr in enumerate(det_hdrs, start=1):
        cell = ws_det.cell(1, ci, hdr)
        cell.fill = hdr_fill
        cell.font = hdr_font
        cell.border = brd
        cell.alignment = Alignment(horizontal="center")

    for dr2, sec in enumerate(sections, start=2):
        co   = sec["company"]
        site = sec["site"]
        row_vals = [
            site["ref_no"], co["label"], site["label"],
            site["billing_period_label"],
            site["city"], site["postcode"], site["meter_id"],
            float(site["supplier_ef"]), site["unit"],
            site["start_reading"], site["end_reading"],
            float(site["total_quantity"]), float(site["total_cost"]),
            float(site["emissions_kg"]), float(site["emissions_t"]),
        ]
        tariffs_here = site.get("tariffs", [])
        for k in range(max_tariffs):
            if k < len(tariffs_here):
                t = tariffs_here[k]
                row_vals += [
                    t["name"], float(t["quantity"]),
                    float(t["unit_cost"]), float(t["cost"]),
                ]
            else:
                row_vals += ["", "", "", ""]

        # 1-based numeric columns: supplier_ef=8, qty=12, cost=13, emis×2=14-15
        # tariff groups start at col 16: qty=base+1, rate=base+2, cost=base+3
        num_cols = {8, 12, 13, 14, 15}
        rate_cols = set()
        for k in range(max_tariffs):
            base = 16 + k * 4
            num_cols |= {base + 1, base + 2, base + 3}
            rate_cols.add(base + 2)

        for ci, val in enumerate(row_vals, start=1):
            cell = ws_det.cell(dr2, ci, val)
            cell.border = brd
            cell.font = Font(name="Calibri", size=9)
            if ci in num_cols and isinstance(val, float):
                cell.number_format = "#,##0.0000" if ci in rate_cols else "#,##0.00"

    base_widths = [30, 28, 22, 18, 16, 10, 24, 16, 8, 14, 14, 14, 14, 16, 16]
    det_widths = base_widths + [26, 12, 12, 12] * max_tariffs
    for ci, w in enumerate(det_widths, start=1):
        ws_det.column_dimensions[get_column_letter(ci)].width = w

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ── CSV public API ────────────────────────────────────────────────────────────

def generate_csv(config: dict, sections: list[dict]) -> bytes:
    lang = config["document"].get("language", "en")
    T    = TRANSLATIONS.get(lang, TRANSLATIONS["en"])

    buf = io.StringIO()
    writer = csv_module.writer(buf)

    # Determine max tariffs across all sections for dynamic columns (mirrors XLSX Detail)
    max_tariffs = max((len(s["site"].get("tariffs", [])) for s in sections), default=0)

    # Header — matches XLSX Detail sheet order
    hdrs = [
        T["xl_col_ref"],       T["xl_col_company"],   T["xl_col_site"],
        T.get("xl_col_period", "Billing Period"),
        T["meta_period_start"], T["meta_period_end"],
        T["xl_col_city"],      T["xl_col_postcode"],   T["xl_col_meter_id"],
        T["xl_col_supplier_ef"],  T["xl_col_unit"],
        T["xl_col_start_read"], T["xl_col_end_read"],
        T["xl_col_total_qty"], T["xl_col_total_cost"],
        T["xl_col_emissions_kg"], T["xl_col_emissions_t"],
    ]
    for k in range(max_tariffs):
        prefix = f"Tariff {k + 1}"
        hdrs += [
            f"{prefix}: {T['xl_tariff_name']}",
            f"{prefix}: {T['xl_tariff_qty']}",
            f"{prefix}: {T['xl_tariff_rate']}",
            f"{prefix}: {T['xl_tariff_cost']}",
        ]
    writer.writerow(hdrs)

    for sec in sections:
        co   = sec["company"]
        site = sec["site"]
        ps = site["period_start"]
        pe = site["period_end"]
        row = [
            site["ref_no"],          co["label"],              site["label"],
            site["billing_period_label"],
            ps.strftime("%Y-%m-%d") if hasattr(ps, "strftime") else str(ps),
            pe.strftime("%Y-%m-%d") if hasattr(pe, "strftime") else str(pe),
            site["city"],            site["postcode"],          site["meter_id"],
            f"{float(site['supplier_ef']):.4f}", site["unit"],
            site["start_reading"],   site["end_reading"],
            f"{float(site['total_quantity']):.2f}", f"{float(site['total_cost']):.2f}",
            f"{float(site['emissions_kg']):.2f}",   f"{float(site['emissions_t']):.3f}",
        ]
        tariffs_here = site.get("tariffs", [])
        for k in range(max_tariffs):
            if k < len(tariffs_here):
                t = tariffs_here[k]
                row += [
                    t["name"],
                    f"{float(t['quantity']):.2f}",
                    f"{float(t['unit_cost']):.4f}",
                    f"{float(t['cost']):.2f}",
                ]
            else:
                row += ["", "", "", ""]
        writer.writerow(row)

    return buf.getvalue().encode("utf-8-sig")


# ── DOCX public API ───────────────────────────────────────────────────────────

def generate_docx(config: dict, sections: list[dict]) -> bytes:
    import docx
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    lang    = config["document"].get("language", "en")
    T       = TRANSLATIONS.get(lang, TRANSLATIONS["en"])
    fp      = config["financial_period"]

    document = docx.Document()
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)

    # Narrow margins
    for section in document.sections:
        section.top_margin    = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin   = Inches(0.8)
        section.right_margin  = Inches(0.8)

    # Cover heading
    h = document.add_heading(T["doc_title_heading"], level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.runs[0].font.size = Pt(16)

    sub = document.add_paragraph(T["doc_subtitle"])
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(0x5A, 0x60, 0x66)

    start_str = fp["start_date"].strftime("%d %b %Y") if hasattr(fp["start_date"], "strftime") else str(fp["start_date"])
    end_str   = fp["end_date"].strftime("%d %b %Y")   if hasattr(fp["end_date"],   "strftime") else str(fp["end_date"])
    meta_str  = f"{fp['label']}  |  {start_str} \u2013 {end_str}"
    mp = document.add_paragraph(meta_str)
    mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()

    for sec in sections:
        co   = sec["company"]
        site = sec["site"]
        acc_hex = co["accent"]
        acc_r, acc_g, acc_b = (int(acc_hex[i:i+2], 16) for i in (1, 3, 5))

        # Section heading: company / site / billing period
        h2 = document.add_heading(f"{co['label']}  \u2014  {site['label']}", level=2)
        h2.runs[0].font.color.rgb = RGBColor(acc_r, acc_g, acc_b)

        period_start_str = site["period_start"].strftime("%d %b %Y") if hasattr(site["period_start"], "strftime") else str(site["period_start"])
        period_end_str   = site["period_end"].strftime("%d %b %Y")   if hasattr(site["period_end"],   "strftime") else str(site["period_end"])
        period_p = document.add_paragraph(
            f"{site['billing_period_label']}  \u2022  "
            f"{T['meta_period_start']}: {period_start_str}  \u2013  "
            f"{T['meta_period_end']}: {period_end_str}"
        )
        period_p.runs[0].font.color.rgb = RGBColor(0x5A, 0x60, 0x66)
        period_p.runs[0].font.size = Pt(9)

        # ── Meter & Consumption table ─────────────────────────────────────────
        def _add_section_table(title, rows):
            document.add_paragraph(title).runs[0].font.bold = True
            tbl = document.add_table(rows=len(rows) + 1, cols=2)
            tbl.style = "Table Grid"
            hc0 = tbl.cell(0, 0)
            hc1 = tbl.cell(0, 1)
            for hc, txt in [(hc0, "Field"), (hc1, "Value")]:
                hc.paragraphs[0].text = txt
                hc.paragraphs[0].runs[0].font.bold = True
                hc.paragraphs[0].runs[0].font.color.rgb = RGBColor(acc_r, acc_g, acc_b)
            for ri, (k, v) in enumerate(rows, start=1):
                tbl.cell(ri, 0).paragraphs[0].text = k
                tbl.cell(ri, 1).paragraphs[0].text = str(v)
            document.add_paragraph()

        unit   = site["unit"]
        symbol = site.get("currency_symbol", "")

        _add_section_table(T["tbl_meter"], [
            (T["row_site"],       site["label"]),
            (T["row_city"],       site["city"]),
            (T["row_postcode"],   site["postcode"]),
            (T["row_meter_id"],   site["meter_id"]),
            (T["row_unit"],       unit),
            (T["row_start_read"], f"{site['start_reading']:,}"),
            (T["row_end_read"],   f"{site['end_reading']:,}"),
            (T["row_total_qty"],  f"{float(site['total_quantity']):,.2f} {unit}"),
        ])

        _add_section_table(T["tbl_grid"], [
            (T["row_supplier_ef"],  f"{float(site['supplier_ef']):.4f}"),
            (T["row_emissions_kg"], f"{float(site['emissions_kg']):,.2f}"),
            (T["row_emissions_t"],  f"{float(site['emissions_t']):.3f}"),
        ])

        # ── Tariff breakdown ──────────────────────────────────────────────────
        if site["tariffs"]:
            document.add_paragraph(T["tbl_tariff"]).runs[0].font.bold = True
            tar_tbl = document.add_table(rows=len(site["tariffs"]) + 1, cols=5)
            tar_tbl.style = "Table Grid"
            for ci, hdr in enumerate([
                T["col_tariff_name"], T["col_tariff_qty"], T["col_tariff_unit"],
                T["col_tariff_rate"], T["col_tariff_cost"],
            ]):
                cell = tar_tbl.cell(0, ci)
                cell.paragraphs[0].text = hdr
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(acc_r, acc_g, acc_b)
            for ri, t in enumerate(site["tariffs"], start=1):
                vals = [
                    t["name"],
                    f"{float(t['quantity']):,.2f}",
                    t["unit"],
                    f"{symbol}{float(t['unit_cost']):.4f}",
                    f"{symbol}{float(t['cost']):,.2f}",
                ]
                for ci, val in enumerate(vals):
                    tar_tbl.cell(ri, ci).paragraphs[0].text = val
            document.add_paragraph()

        # ── Cost summary ──────────────────────────────────────────────────────
        cost_p = document.add_paragraph()
        run = cost_p.add_run(
            f"{T['charge_total']}: {symbol}{float(site['total_cost']):,.2f}"
        )
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(acc_r, acc_g, acc_b)

        document.add_paragraph()
        document.add_paragraph("\u2014" * 40)  # section separator
        document.add_paragraph()

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()
