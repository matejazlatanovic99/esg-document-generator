from __future__ import annotations

from decimal import Decimal, ROUND_HALF_UP
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ── translations ───────────────────────────────────────────────────────────────

TRANSLATIONS: dict[str, dict[str, str]] = {
    "en": {
        "logo_subtitle":     "Metered District Heating Services",
        "doc_title_heading": "Heating Billing Document",
        "box_supplier":      "Supplier Details",
        "box_customer":      "Customer / Service Address",
        "box_invoice":       "Invoice Metadata",
        "meta_invoice_no":   "Invoice Number",
        "meta_issue_date":   "Issue Date",
        "meta_due_date":     "Payment Due Date",
        "meta_currency":     "Currency",
        "tbl_header_field":  "Billing Field",
        "tbl_header_value":  "Recorded Value",
        "row_supplier":      "Supplier",
        "row_customer":      "Customer Name",
        "row_site":          "Site",
        "row_city":          "City",
        "row_postcode":      "Postcode",
        "row_period_start":  "Billing Period Start",
        "row_period_end":    "Billing Period End",
        "row_meter_id":      "Heat Meter ID",
        "row_prev_read":     "Previous Meter Reading (kWh)",
        "row_curr_read":     "Current Meter Reading (kWh)",
        "row_consumption":   "Heat Consumption (kWh)",
        "row_unit_price":    "Heat Unit Price (£/kWh)",
        "row_capacity":      "Contracted Capacity (kW)",
        "row_capacity_rate": "Capacity Charge (£/kW/month)",
        "box_charges":       "Charges & VAT Summary",
        "charge_heat":       "Heat Consumption Cost",
        "charge_capacity":   "Capacity Charge",
        "charge_subtotal":   "Subtotal",
        "charge_vat":        "VAT (5%)",
        "charge_total":      "Total Amount Due",
        "footer_vat":        "VAT applied at 5%. Payment terms: 14 days from issue date unless otherwise specified in the supply agreement.",
        "footer_page":       "Page {page} / {total}",
    },
    "fr": {
        "logo_subtitle":     "Services de chauffage urbain mesurés",
        "doc_title_heading": "Document de facturation thermique",
        "box_supplier":      "Détails du fournisseur",
        "box_customer":      "Client / Adresse du service",
        "box_invoice":       "Métadonnées de la facture",
        "meta_invoice_no":   "Numéro de facture",
        "meta_issue_date":   "Date d'émission",
        "meta_due_date":     "Date d'échéance",
        "meta_currency":     "Devise",
        "tbl_header_field":  "Champ de facturation",
        "tbl_header_value":  "Valeur enregistrée",
        "row_supplier":      "Fournisseur",
        "row_customer":      "Nom du client",
        "row_site":          "Site",
        "row_city":          "Ville",
        "row_postcode":      "Code postal",
        "row_period_start":  "Début de la période",
        "row_period_end":    "Fin de la période",
        "row_meter_id":      "Identifiant du compteur",
        "row_prev_read":     "Relevé précédent (kWh)",
        "row_curr_read":     "Relevé actuel (kWh)",
        "row_consumption":   "Consommation thermique (kWh)",
        "row_unit_price":    "Prix unitaire (£/kWh)",
        "row_capacity":      "Capacité contractée (kW)",
        "row_capacity_rate": "Frais de capacité (£/kW/mois)",
        "box_charges":       "Résumé des charges et TVA",
        "charge_heat":       "Coût de consommation thermique",
        "charge_capacity":   "Frais de capacité",
        "charge_subtotal":   "Sous-total",
        "charge_vat":        "TVA (5%)",
        "charge_total":      "Montant total dû",
        "footer_vat":        "TVA appliquée à 5 %. Conditions de paiement : 14 jours à compter de la date d'émission, sauf accord contraire.",
        "footer_page":       "Page {page} / {total}",
    },
    "de": {
        "logo_subtitle":     "Gemessene Fernwärmedienste",
        "doc_title_heading": "Fernwärme-Abrechnungsdokument",
        "box_supplier":      "Lieferantendetails",
        "box_customer":      "Kunde / Serviceadresse",
        "box_invoice":       "Rechnungsmetadaten",
        "meta_invoice_no":   "Rechnungsnummer",
        "meta_issue_date":   "Ausstellungsdatum",
        "meta_due_date":     "Zahlungsfälligkeitsdatum",
        "meta_currency":     "Währung",
        "tbl_header_field":  "Abrechnungsfeld",
        "tbl_header_value":  "Erfasster Wert",
        "row_supplier":      "Lieferant",
        "row_customer":      "Kundenname",
        "row_site":          "Standort",
        "row_city":          "Stadt",
        "row_postcode":      "Postleitzahl",
        "row_period_start":  "Abrechnungszeitraum Beginn",
        "row_period_end":    "Abrechnungszeitraum Ende",
        "row_meter_id":      "Wärmezähler-ID",
        "row_prev_read":     "Vorheriger Zählerstand (kWh)",
        "row_curr_read":     "Aktueller Zählerstand (kWh)",
        "row_consumption":   "Wärmeverbrauch (kWh)",
        "row_unit_price":    "Einheitspreis (£/kWh)",
        "row_capacity":      "Vertragsleistung (kW)",
        "row_capacity_rate": "Leistungsgebühr (£/kW/Monat)",
        "box_charges":       "Kosten- und MwSt.-Übersicht",
        "charge_heat":       "Wärmeverbrauchskosten",
        "charge_capacity":   "Leistungsgebühr",
        "charge_subtotal":   "Zwischensumme",
        "charge_vat":        "MwSt. (5%)",
        "charge_total":      "Gesamtbetrag fällig",
        "footer_vat":        "MwSt. zu 5 % angewendet. Zahlungsbedingungen: 14 Tage ab Ausstellungsdatum.",
        "footer_page":       "Seite {page} / {total}",
    },
    "nl": {
        "logo_subtitle":     "Gemeten stadsverwarmingsdiensten",
        "doc_title_heading": "Factuur stadsverwarming",
        "box_supplier":      "Leveranciersgegevens",
        "box_customer":      "Klant / Serviceadres",
        "box_invoice":       "Factuurmetadata",
        "meta_invoice_no":   "Factuurnummer",
        "meta_issue_date":   "Uitgiftedatum",
        "meta_due_date":     "Vervaldatum",
        "meta_currency":     "Valuta",
        "tbl_header_field":  "Factuurveld",
        "tbl_header_value":  "Geregistreerde waarde",
        "row_supplier":      "Leverancier",
        "row_customer":      "Klantnaam",
        "row_site":          "Locatie",
        "row_city":          "Stad",
        "row_postcode":      "Postcode",
        "row_period_start":  "Begin facturatieperiode",
        "row_period_end":    "Einde facturatieperiode",
        "row_meter_id":      "Warmtemeter-ID",
        "row_prev_read":     "Vorige meterstand (kWh)",
        "row_curr_read":     "Huidige meterstand (kWh)",
        "row_consumption":   "Warmteverbruik (kWh)",
        "row_unit_price":    "Eenheidsprijs (£/kWh)",
        "row_capacity":      "Gecontracteerd vermogen (kW)",
        "row_capacity_rate": "Vermogenstoeslag (£/kW/maand)",
        "box_charges":       "Kosten- en BTW-overzicht",
        "charge_heat":       "Warmteverbruikskosten",
        "charge_capacity":   "Vermogenstoeslag",
        "charge_subtotal":   "Subtotaal",
        "charge_vat":        "BTW (5%)",
        "charge_total":      "Totaal verschuldigd bedrag",
        "footer_vat":        "BTW 5 % toegepast. Betalingsvoorwaarden: 14 dagen na factuurdatum.",
        "footer_page":       "Pagina {page} / {total}",
    },
}

# ── layout constants ───────────────────────────────────────────────────────────

_USABLE_W   = Cm(18.0)   # 21cm page - 1.5cm * 2 margins
_HALF_W     = Cm(9.0)
_LABEL_W    = Cm(10.2)   # 56 % for billing-field label column
_VALUE_W    = Cm(7.8)    # 44 % for billing-field value column
_CHARGE_L   = Cm(12.0)
_CHARGE_V   = Cm(6.0)

_WHITE      = (255, 255, 255)
_DARK       = (31,  35,  40)
_GRAY       = (90,  96,  102)
_BORDER_HEX = "C9CDD2"

# ── XML / styling helpers ──────────────────────────────────────────────────────

def _hex_rgb(hex_color: str) -> tuple[int, int, int]:
    h = hex_color.lstrip("#")
    return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)


def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#").upper())
    tcPr.append(shd)


def _apply_table_borders(table, color: str = _BORDER_HEX, none: bool = False) -> None:
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none" if none else "single")
        el.set(qn("w:sz"), "0" if none else "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto" if none else color.lstrip("#").upper())
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _set_col_widths(table, *widths_cm: float) -> None:
    for i, w in enumerate(widths_cm):
        for row in table.rows:
            row.cells[i].width = Cm(w)


def _p(cell, text: str, bold: bool = False, size: float = 9.0,
       rgb: tuple | None = None, align=WD_ALIGN_PARAGRAPH.LEFT,
       mono: bool = False, first: bool = True) -> None:
    """Write text into a cell paragraph (first=True uses existing, else adds new)."""
    para = cell.paragraphs[0] if first else cell.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(1)
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if rgb:
        run.font.color.rgb = RGBColor(*rgb)
    if mono:
        run.font.name = "Courier New"


def _spacer(doc, pts: float = 4.0) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before  = Pt(0)
    p.paragraph_format.space_after   = Pt(0)
    p.paragraph_format.line_spacing  = Pt(pts)


# ── formatting helpers ─────────────────────────────────────────────────────────

def _q2(v) -> Decimal:
    if not isinstance(v, Decimal):
        v = Decimal(str(v))
    return v.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)


def _fmt_money(v) -> str:
    return f"£{_q2(v):,.2f}"


def _fmt_rate(v, places: int = 3) -> str:
    if not isinstance(v, Decimal):
        v = Decimal(str(v))
    q = Decimal("1." + "0" * places)
    return str(v.quantize(q, rounding=ROUND_HALF_UP))


# ── per-invoice page builder ───────────────────────────────────────────────────

def _render_invoice(
    doc: Document,
    company: dict,
    site: dict,
    rec: dict,
    page_no: int,
    total_pages: int,
    fp_label: str,
    strings: dict,
    omit: set[str],
) -> None:
    accent     = company["accent"]
    accent_soft = company["accent_soft"]
    accent_rgb = _hex_rgb(accent)
    soft_rgb   = _hex_rgb(accent_soft)

    # ── 1. Header ──────────────────────────────────────────────────────────────
    hdr = doc.add_table(rows=2, cols=2)
    _apply_table_borders(hdr, none=True)
    _set_col_widths(hdr, 11.0, 7.0)

    _p(hdr.cell(0, 0), company["supplier"], bold=True, size=14, rgb=accent_rgb)
    _p(hdr.cell(0, 1), strings["doc_title_heading"], bold=True, size=12,
       rgb=_DARK, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _p(hdr.cell(1, 0), strings["logo_subtitle"], size=8, rgb=_GRAY)
    _p(hdr.cell(1, 1),
       f"{fp_label}  •  {company['label']}  •  {rec['billing_period_label']}",
       size=7.5, rgb=_GRAY, align=WD_ALIGN_PARAGRAPH.RIGHT)

    _spacer(doc, 5)

    # ── 2. Address boxes ───────────────────────────────────────────────────────
    sup_lines  = list(company.get("supplier_address", []))
    cust_lines = list(site.get("customer_address", []))
    n_addr = max(len(sup_lines), len(cust_lines), 1)

    addr = doc.add_table(rows=n_addr + 1, cols=2)
    _apply_table_borders(addr)
    _set_col_widths(addr, 9.0, 9.0)

    _set_cell_bg(addr.cell(0, 0), accent)
    _set_cell_bg(addr.cell(0, 1), accent)
    _p(addr.cell(0, 0), strings["box_supplier"], bold=True, size=9, rgb=_WHITE)
    _p(addr.cell(0, 1), strings["box_customer"], bold=True, size=9, rgb=_WHITE)

    for i in range(n_addr):
        _p(addr.cell(i + 1, 0), sup_lines[i]  if i < len(sup_lines)  else "", size=8.5, rgb=_DARK)
        _p(addr.cell(i + 1, 1), cust_lines[i] if i < len(cust_lines) else "", size=8.5, rgb=_DARK)

    _spacer(doc, 5)

    # ── 3. Invoice metadata ────────────────────────────────────────────────────
    meta = doc.add_table(rows=3, cols=2)
    _apply_table_borders(meta)
    _set_col_widths(meta, 9.0, 9.0)

    hdr_cell = meta.cell(0, 0).merge(meta.cell(0, 1))
    _set_cell_bg(hdr_cell, accent)
    _p(hdr_cell, strings["box_invoice"], bold=True, size=9, rgb=_WHITE)

    _p(meta.cell(1, 0), f"{strings['meta_invoice_no']}:  {rec['invoice_no']}", size=8.5, rgb=_DARK)
    _p(meta.cell(1, 1), f"{strings['meta_issue_date']}:  {rec['issue_date'].strftime('%d %b %Y')}", size=8.5, rgb=_DARK)
    _p(meta.cell(2, 0), f"{strings['meta_due_date']}:  {rec['due_date'].strftime('%d %b %Y')}", size=8.5, rgb=_DARK)
    _p(meta.cell(2, 1), f"{strings['meta_currency']}:  {company['currency']}", size=8.5, rgb=_DARK)

    _spacer(doc, 5)

    # ── 4. Billing fields table ────────────────────────────────────────────────
    billing_rows: list[tuple[str, str, bool]] = [
        # (label, value, is_mono)
        (strings["row_supplier"],      rec["supplier"],                                  False),
        (strings["row_customer"],      rec["customer"],                                  False),
        (strings["row_site"],          rec["site_label"],                                False),
        (strings["row_city"],          rec["city"],                                      False),
        (strings["row_postcode"],      rec["postcode"],                                  False),
        (strings["row_period_start"],  rec["period_start"].strftime("%d %b %Y"),         False),
        (strings["row_period_end"],    rec["period_end"].strftime("%d %b %Y"),           False),
        (strings["row_meter_id"],      rec["meter_id"],                                  True),
        (strings["row_prev_read"],     "" if "prev_read"   in omit else f"{rec['prev_read']:,}",    False),
        (strings["row_curr_read"],     "" if "curr_read"   in omit else f"{rec['curr_read']:,}",    False),
        (strings["row_consumption"],   "" if "consumption" in omit else f"{rec['consumption']:,}",  False),
        (strings["row_unit_price"],    "" if "unit_price"  in omit else _fmt_rate(rec["unit_price"], 3), False),
        (strings["row_capacity"],      "" if "capacity_kw" in omit else str(rec["capacity_kw"]),    False),
        (strings["row_capacity_rate"], "" if "capacity_rate" in omit else _fmt_rate(rec["capacity_rate"], 2), False),
    ]

    btbl = doc.add_table(rows=len(billing_rows) + 1, cols=2)
    _apply_table_borders(btbl)
    _set_col_widths(btbl, 10.2, 7.8)

    _set_cell_bg(btbl.cell(0, 0), accent)
    _set_cell_bg(btbl.cell(0, 1), accent)
    _p(btbl.cell(0, 0), strings["tbl_header_field"], bold=True, size=9, rgb=_WHITE)
    _p(btbl.cell(0, 1), strings["tbl_header_value"], bold=True, size=9, rgb=_WHITE)

    for i, (label, value, mono) in enumerate(billing_rows):
        row_idx = i + 1
        bg = accent_soft if i % 2 == 1 else None
        lc = btbl.cell(row_idx, 0)
        vc = btbl.cell(row_idx, 1)
        if bg:
            _set_cell_bg(lc, accent_soft)
            _set_cell_bg(vc, accent_soft)
        _p(lc, label,  size=8.5, rgb=_GRAY)
        _p(vc, value,  size=8.5, rgb=_DARK, mono=mono)

    _spacer(doc, 5)

    # ── 5. Charges summary ─────────────────────────────────────────────────────
    charge_rows: list[tuple[str, str, bool]] = [
        (strings["charge_heat"],     _fmt_money(rec["heat_cost"]),       False),
        (strings["charge_capacity"], _fmt_money(rec["capacity_charge"]), False),
        (strings["charge_subtotal"], _fmt_money(rec["subtotal"]),        False),
        (strings["charge_vat"],      _fmt_money(rec["vat"]),             False),
        (strings["charge_total"],    _fmt_money(rec["total"]),           True),   # highlighted
    ]

    ctbl = doc.add_table(rows=len(charge_rows) + 1, cols=2)
    _apply_table_borders(ctbl)
    _set_col_widths(ctbl, 12.0, 6.0)

    hdr_c = ctbl.cell(0, 0).merge(ctbl.cell(0, 1))
    _set_cell_bg(hdr_c, accent)
    _p(hdr_c, strings["box_charges"], bold=True, size=9, rgb=_WHITE)

    for i, (label, value, is_total) in enumerate(charge_rows):
        lc = ctbl.cell(i + 1, 0)
        vc = ctbl.cell(i + 1, 1)
        if is_total:
            _set_cell_bg(lc, accent_soft)
            _set_cell_bg(vc, accent_soft)
            _p(lc, label, bold=True, size=9, rgb=accent_rgb)
            _p(vc, value, bold=True, size=9, rgb=accent_rgb,
               align=WD_ALIGN_PARAGRAPH.RIGHT)
        else:
            _p(lc, label, size=8.5, rgb=_GRAY)
            _p(vc, value, size=8.5, rgb=_DARK, align=WD_ALIGN_PARAGRAPH.RIGHT)

    _spacer(doc, 5)

    # ── 6. Footer ──────────────────────────────────────────────────────────────
    ftbl = doc.add_table(rows=1, cols=2)
    _apply_table_borders(ftbl, none=True)
    _set_col_widths(ftbl, 13.0, 5.0)

    _p(ftbl.cell(0, 0), strings["footer_vat"], size=7, rgb=_GRAY)
    _p(ftbl.cell(0, 1),
       strings["footer_page"].format(page=page_no, total=total_pages),
       size=7, rgb=_GRAY, align=WD_ALIGN_PARAGRAPH.RIGHT)


# ── public API ─────────────────────────────────────────────────────────────────

def generate_docx(
    config: dict,
    sections: list[dict],
    blank_fields: set[str] | None = None,
) -> bytes:
    """Build a DOCX billing document from sections and return bytes."""
    lang    = config["document"].get("language", "en")
    strings = TRANSLATIONS.get(lang, TRANSLATIONS["en"])
    omit    = blank_fields or set()

    doc = Document()

    # Page setup: A4, narrow margins
    page_sec = doc.sections[0]
    page_sec.page_height    = Cm(29.7)
    page_sec.page_width     = Cm(21.0)
    page_sec.left_margin    = Cm(1.5)
    page_sec.right_margin   = Cm(1.5)
    page_sec.top_margin     = Cm(1.5)
    page_sec.bottom_margin  = Cm(1.5)

    doc.core_properties.title   = config["document"].get("title", "")
    doc.core_properties.subject = config["document"].get("subject", "")

    # Remove the default empty paragraph python-docx adds on creation (if any)
    if doc.paragraphs:
        default_para = doc.paragraphs[0]._element
        default_para.getparent().remove(default_para)

    fp_label    = config["financial_period"]["label"]
    total_pages = sum(len(s["records"]) for s in sections)
    page_no     = 1
    first       = True

    for sec in sections:
        company = sec["company"]
        site    = sec["site"]
        for rec in sec["records"]:
            if not first:
                doc.add_page_break()
            first = False
            _render_invoice(
                doc, company, site, rec,
                page_no, total_pages, fp_label, strings, omit,
            )
            page_no += 1

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
