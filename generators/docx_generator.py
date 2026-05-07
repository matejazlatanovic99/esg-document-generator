from __future__ import annotations

from decimal import Decimal, ROUND_HALF_UP
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from utils.currency import format_money, replace_pound_labels

# ── translations ───────────────────────────────────────────────────────────────

TRANSLATIONS: dict[str, dict[str, str]] = {
    "en": {
        "logo_subtitle":     "Metered District Heating Services",
        "doc_title_heading": "Heating Billing Document",
        "box_supplier":      "Supplier Details",
        "box_customer":      "Customer / Service Address",
        "box_invoice":       "Invoice Metadata",
        "meta_invoice_no":   "Invoice Number",
        "meta_issue_date":   "Issue Date",
        "meta_due_date":     "Payment Due Date",
        "meta_currency":     "Currency",
        "tbl_header_field":  "Billing Field",
        "tbl_header_value":  "Recorded Value",
        "row_supplier":      "Supplier",
        "row_customer":      "Customer Name",
        "row_site":          "Site",
        "row_city":          "City",
        "row_postcode":      "Postcode",
        "row_period_start":  "Billing Period Start",
        "row_period_end":    "Billing Period End",
        "row_meter_id":      "Heat Meter ID",
        "row_prev_read":     "Previous Meter Reading (kWh)",
        "row_curr_read":     "Current Meter Reading (kWh)",
        "row_consumption":   "Heat Consumption (kWh)",
        "row_unit_price":    "Heat Unit Price (£/kWh)",
        "row_capacity":      "Contracted Capacity (kW)",
        "row_capacity_rate": "Capacity Charge (£/kW/month)",
        "row_supplier_ef":   "Supplier Emission Factor (kg CO\u2082e/kWh)",
        "box_charges":       "Charges & VAT Summary",
        "charge_heat":       "Heat Consumption Cost",
        "charge_capacity":   "Capacity Charge",
        "charge_subtotal":   "Subtotal",
        "charge_vat":        "VAT (5%)",
        "charge_total":      "Total Amount Due",
        "footer_vat":        "VAT applied at 5%. Payment terms: 14 days from issue date unless otherwise specified in the supply agreement.",
        "footer_page":       "Page {page} / {total}",
    },
    "fr": {
        "logo_subtitle":     "Services de chauffage urbain mesurés",
        "doc_title_heading": "Document de facturation thermique",
        "box_supplier":      "Détails du fournisseur",
        "box_customer":      "Client / Adresse du service",
        "box_invoice":       "Métadonnées de la facture",
        "meta_invoice_no":   "Numéro de facture",
        "meta_issue_date":   "Date d'émission",
        "meta_due_date":     "Date d'échéance",
        "meta_currency":     "Devise",
        "tbl_header_field":  "Champ de facturation",
        "tbl_header_value":  "Valeur enregistrée",
        "row_supplier":      "Fournisseur",
        "row_customer":      "Nom du client",
        "row_site":          "Site",
        "row_city":          "Ville",
        "row_postcode":      "Code postal",
        "row_period_start":  "Début de la période",
        "row_period_end":    "Fin de la période",
        "row_meter_id":      "Identifiant du compteur",
        "row_prev_read":     "Relevé précédent (kWh)",
        "row_curr_read":     "Relevé actuel (kWh)",
        "row_consumption":   "Consommation thermique (kWh)",
        "row_unit_price":    "Prix unitaire (£/kWh)",
        "row_capacity":      "Capacité contractée (kW)",
        "row_capacity_rate": "Frais de capacité (£/kW/mois)",
        "row_supplier_ef":   "Facteur d'émission fournisseur (kg CO\u2082e/kWh)",
        "box_charges":       "Résumé des charges et TVA",
        "charge_heat":       "Coût de consommation thermique",
        "charge_capacity":   "Frais de capacité",
        "charge_subtotal":   "Sous-total",
        "charge_vat":        "TVA (5%)",
        "charge_total":      "Montant total dû",
        "footer_vat":        "TVA appliquée à 5 %. Conditions de paiement : 14 jours à compter de la date d'émission, sauf accord contraire.",
        "footer_page":       "Page {page} / {total}",
    },
    "de": {
        "logo_subtitle":     "Gemessene Fernwärmedienste",
        "doc_title_heading": "Fernwärme-Abrechnungsdokument",
        "box_supplier":      "Lieferantendetails",
        "box_customer":      "Kunde / Serviceadresse",
        "box_invoice":       "Rechnungsmetadaten",
        "meta_invoice_no":   "Rechnungsnummer",
        "meta_issue_date":   "Ausstellungsdatum",
        "meta_due_date":     "Zahlungsfälligkeitsdatum",
        "meta_currency":     "Währung",
        "tbl_header_field":  "Abrechnungsfeld",
        "tbl_header_value":  "Erfasster Wert",
        "row_supplier":      "Lieferant",
        "row_customer":      "Kundenname",
        "row_site":          "Standort",
        "row_city":          "Stadt",
        "row_postcode":      "Postleitzahl",
        "row_period_start":  "Abrechnungszeitraum Beginn",
        "row_period_end":    "Abrechnungszeitraum Ende",
        "row_meter_id":      "Wärmezähler-ID",
        "row_prev_read":     "Vorheriger Zählerstand (kWh)",
        "row_curr_read":     "Aktueller Zählerstand (kWh)",
        "row_consumption":   "Wärmeverbrauch (kWh)",
        "row_unit_price":    "Einheitspreis (£/kWh)",
        "row_capacity":      "Vertragsleistung (kW)",
        "row_capacity_rate": "Leistungsgebühr (£/kW/Monat)",
        "row_supplier_ef":   "Emissionsfaktor Lieferant (kg CO\u2082e/kWh)",
        "box_charges":       "Kosten- und MwSt.-Übersicht",
        "charge_heat":       "Wärmeverbrauchskosten",
        "charge_capacity":   "Leistungsgebühr",
        "charge_subtotal":   "Zwischensumme",
        "charge_vat":        "MwSt. (5%)",
        "charge_total":      "Gesamtbetrag fällig",
        "footer_vat":        "MwSt. zu 5 % angewendet. Zahlungsbedingungen: 14 Tage ab Ausstellungsdatum.",
        "footer_page":       "Seite {page} / {total}",
    },
    "nl": {
        "logo_subtitle":     "Gemeten stadsverwarmingsdiensten",
        "doc_title_heading": "Factuur stadsverwarming",
        "box_supplier":      "Leveranciersgegevens",
        "box_customer":      "Klant / Serviceadres",
        "box_invoice":       "Factuurmetadata",
        "meta_invoice_no":   "Factuurnummer",
        "meta_issue_date":   "Uitgiftedatum",
        "meta_due_date":     "Vervaldatum",
        "meta_currency":     "Valuta",
        "tbl_header_field":  "Factuurveld",
        "tbl_header_value":  "Geregistreerde waarde",
        "row_supplier":      "Leverancier",
        "row_customer":      "Klantnaam",
        "row_site":          "Locatie",
        "row_city":          "Stad",
        "row_postcode":      "Postcode",
        "row_period_start":  "Begin facturatieperiode",
        "row_period_end":    "Einde facturatieperiode",
        "row_meter_id":      "Warmtemeter-ID",
        "row_prev_read":     "Vorige meterstand (kWh)",
        "row_curr_read":     "Huidige meterstand (kWh)",
        "row_consumption":   "Warmteverbruik (kWh)",
        "row_unit_price":    "Eenheidsprijs (£/kWh)",
        "row_capacity":      "Gecontracteerd vermogen (kW)",
        "row_capacity_rate": "Vermogenstoeslag (£/kW/maand)",
        "row_supplier_ef":   "Emissiefactor leverancier (kg CO\u2082e/kWh)",
        "box_charges":       "Kosten- en BTW-overzicht",
        "charge_heat":       "Warmteverbruikskosten",
        "charge_capacity":   "Vermogenstoeslag",
        "charge_subtotal":   "Subtotaal",
        "charge_vat":        "BTW (5%)",
        "charge_total":      "Totaal verschuldigd bedrag",
        "footer_vat":        "BTW 5 % toegepast. Betalingsvoorwaarden: 14 dagen na factuurdatum.",
        "footer_page":       "Pagina {page} / {total}",
    },
}

ELECTRICITY_TRANSLATIONS: dict[str, dict[str, str]] = {
    "en": {
        "doc_title_heading": "Electricity Consumption Statement",
        "doc_subtitle": "Scope 2 - Purchased Electricity",
        "box_supplier": "Supplier Details",
        "box_customer": "Customer / Site Address",
        "box_period": "Statement Period",
        "meta_period_label": "Period",
        "meta_period_start": "Period Start",
        "meta_period_end": "Period End",
        "meta_ref": "Reference",
        "meta_currency": "Currency",
        "tbl_meter": "Meter & Consumption",
        "tbl_grid": "Grid & Emissions",
        "row_site": "Site",
        "row_city": "City",
        "row_postcode": "Postcode",
        "row_meter_id": "Electricity Meter ID",
        "row_unit": "Measurement Unit",
        "row_start_read": "Start Meter Reading",
        "row_end_read": "End Meter Reading",
        "row_total_qty": "Total Consumption",
        "row_supplier_ef": "Supplier Emission Factor (kg CO\u2082e/kWh)",
        "row_emissions_kg": "Total Emissions (kg CO\u2082e)",
        "row_emissions_t": "Total Emissions (tCO\u2082e)",
        "tbl_tariff": "Tariff Breakdown",
        "col_tariff_name": "Tariff / Rate",
        "col_tariff_qty": "Quantity",
        "col_tariff_unit": "Unit",
        "col_tariff_rate": "Unit Cost",
        "col_tariff_cost": "Cost",
        "charge_total": "Total Electricity Cost",
    },
    "fr": {
        "doc_title_heading": "Relev\u00e9 de consommation \u00e9lectrique",
        "doc_subtitle": "Scope 2 - \u00c9lectricit\u00e9 achet\u00e9e",
        "box_supplier": "D\u00e9tails du fournisseur",
        "box_customer": "Client / Adresse du site",
        "box_period": "P\u00e9riode du relev\u00e9",
        "meta_period_label": "P\u00e9riode",
        "meta_period_start": "D\u00e9but de p\u00e9riode",
        "meta_period_end": "Fin de p\u00e9riode",
        "meta_ref": "R\u00e9f\u00e9rence",
        "meta_currency": "Devise",
        "tbl_meter": "Compteur & Consommation",
        "tbl_grid": "R\u00e9seau & \u00c9missions",
        "row_site": "Site",
        "row_city": "Ville",
        "row_postcode": "Code postal",
        "row_meter_id": "ID du compteur \u00e9lectrique",
        "row_unit": "Unit\u00e9 de mesure",
        "row_start_read": "Relev\u00e9 initial",
        "row_end_read": "Relev\u00e9 final",
        "row_total_qty": "Consommation totale",
        "row_supplier_ef": "Facteur d\u2019\u00e9mission fournisseur (kg CO\u2082e/kWh)",
        "row_emissions_kg": "\u00c9missions totales (kg CO\u2082e)",
        "row_emissions_t": "\u00c9missions totales (tCO\u2082e)",
        "tbl_tariff": "Ventilation par tarif",
        "col_tariff_name": "Tarif / Taux",
        "col_tariff_qty": "Quantit\u00e9",
        "col_tariff_unit": "Unit\u00e9",
        "col_tariff_rate": "Co\u00fbt unitaire",
        "col_tariff_cost": "Co\u00fbt",
        "charge_total": "Co\u00fbt total de l\u2019\u00e9lectricit\u00e9",
    },
    "de": {
        "doc_title_heading": "Stromverbrauchsabrechnung",
        "doc_subtitle": "Scope 2 - Eingekaufter Strom",
        "box_supplier": "Lieferantendetails",
        "box_customer": "Kunde / Standortadresse",
        "box_period": "Abrechnungszeitraum",
        "meta_period_label": "Zeitraum",
        "meta_period_start": "Zeitraum Beginn",
        "meta_period_end": "Zeitraum Ende",
        "meta_ref": "Referenz",
        "meta_currency": "W\u00e4hrung",
        "tbl_meter": "Z\u00e4hler & Verbrauch",
        "tbl_grid": "Netz & Emissionen",
        "row_site": "Standort",
        "row_city": "Stadt",
        "row_postcode": "Postleitzahl",
        "row_meter_id": "Stromz\u00e4hler-ID",
        "row_unit": "Messeinheit",
        "row_start_read": "Anfangsz\u00e4hlerstand",
        "row_end_read": "Endz\u00e4hlerstand",
        "row_total_qty": "Gesamtverbrauch",
        "row_supplier_ef": "Emissionsfaktor Lieferant (kg CO\u2082e/kWh)",
        "row_emissions_kg": "Gesamtemissionen (kg CO\u2082e)",
        "row_emissions_t": "Gesamtemissionen (tCO\u2082e)",
        "tbl_tariff": "Tarifaufschl\u00fcsselung",
        "col_tariff_name": "Tarif / Satz",
        "col_tariff_qty": "Menge",
        "col_tariff_unit": "Einheit",
        "col_tariff_rate": "Einheitspreis",
        "col_tariff_cost": "Kosten",
        "charge_total": "Gesamtstromkosten",
    },
    "nl": {
        "doc_title_heading": "Verklaring elektriciteitsverbruik",
        "doc_subtitle": "Scope 2 - Ingekochte elektriciteit",
        "box_supplier": "Leveranciersgegevens",
        "box_customer": "Klant / Locatieadres",
        "box_period": "Overzichtsperiode",
        "meta_period_label": "Periode",
        "meta_period_start": "Begin periode",
        "meta_period_end": "Einde periode",
        "meta_ref": "Referentie",
        "meta_currency": "Valuta",
        "tbl_meter": "Meter & Verbruik",
        "tbl_grid": "Net & Emissies",
        "row_site": "Locatie",
        "row_city": "Stad",
        "row_postcode": "Postcode",
        "row_meter_id": "Elektriciteitmeter-ID",
        "row_unit": "Meeteenheid",
        "row_start_read": "Beginmeterstand",
        "row_end_read": "Eindmeterstand",
        "row_total_qty": "Totaal verbruik",
        "row_supplier_ef": "Emissiefactor leverancier (kg CO\u2082e/kWh)",
        "row_emissions_kg": "Totale emissies (kg CO\u2082e)",
        "row_emissions_t": "Totale emissies (tCO\u2082e)",
        "tbl_tariff": "Tariefuitsplitsing",
        "col_tariff_name": "Tarief",
        "col_tariff_qty": "Hoeveelheid",
        "col_tariff_unit": "Eenheid",
        "col_tariff_rate": "Eenheidsprijs",
        "col_tariff_cost": "Kosten",
        "charge_total": "Totale elektriciteitskosten",
    },
}

# ── layout constants ───────────────────────────────────────────────────────────

_USABLE_W   = Cm(18.0)   # 21cm page - 1.5cm * 2 margins
_HALF_W     = Cm(9.0)
_LABEL_W    = Cm(10.2)   # 56 % for billing-field label column
_VALUE_W    = Cm(7.8)    # 44 % for billing-field value column
_CHARGE_L   = Cm(12.0)
_CHARGE_V   = Cm(6.0)

_WHITE      = (255, 255, 255)
_DARK       = (31,  35,  40)
_GRAY       = (90,  96,  102)
_BORDER_HEX = "C9CDD2"

# ── XML / styling helpers ──────────────────────────────────────────────────────


def _document_distractor_lines(config: dict) -> list[str]:
    plan = config.get("document", {}).get("distractor_plan")
    if not plan or not getattr(plan, "enabled", False):
        return []
    return [f"{field.label}:  {field.value}" for field in plan.document_fields]


def _paired_lines(lines: list[str]) -> list[tuple[str, str]]:
    pairs: list[tuple[str, str]] = []
    for index in range(0, len(lines), 2):
        pairs.append((lines[index], lines[index + 1] if index + 1 < len(lines) else ""))
    return pairs


def _has_visible_value(value) -> bool:
    if value is None:
        return False
    if isinstance(value, str):
        return bool(value.strip())
    return True


def _join_display_parts(*parts: str) -> str:
    return "  •  ".join(str(part).strip() for part in parts if str(part).strip())

def _hex_rgb(hex_color: str) -> tuple[int, int, int]:
    h = hex_color.lstrip("#")
    return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)


def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#").upper())
    tcPr.append(shd)


def _apply_table_borders(table, color: str = _BORDER_HEX, none: bool = False) -> None:
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none" if none else "single")
        el.set(qn("w:sz"), "0" if none else "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto" if none else color.lstrip("#").upper())
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _set_col_widths(table, *widths_cm: float) -> None:
    for i, w in enumerate(widths_cm):
        for row in table.rows:
            row.cells[i].width = Cm(w)


def _p(cell, text: str, bold: bool = False, size: float = 9.0,
       rgb: tuple | None = None, align=WD_ALIGN_PARAGRAPH.LEFT,
       mono: bool = False, first: bool = True) -> None:
    """Write text into a cell paragraph (first=True uses existing, else adds new)."""
    para = cell.paragraphs[0] if first else cell.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(1)
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if rgb:
        run.font.color.rgb = RGBColor(*rgb)
    if mono:
        run.font.name = "Courier New"


def _spacer(doc, pts: float = 4.0) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before  = Pt(0)
    p.paragraph_format.space_after   = Pt(0)
    p.paragraph_format.line_spacing  = Pt(pts)


def _render_docx_transposed_rows(
    doc: Document,
    rows: list[dict],
    accent: str,
    accent_soft: str,
    *,
    chunk_size: int = 5,
    value_align=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    accent_rgb = _hex_rgb(accent)
    for start_idx in range(0, len(rows), chunk_size):
        chunk = rows[start_idx:start_idx + chunk_size]
        table = doc.add_table(rows=2, cols=len(chunk))
        _apply_table_borders(table)
        _set_col_widths(table, *([18.0 / len(chunk)] * len(chunk)))

        for col_idx, row in enumerate(chunk):
            label_cell = table.cell(0, col_idx)
            value_cell = table.cell(1, col_idx)
            _set_cell_bg(label_cell, accent)
            _p(label_cell, row["label"], bold=True, size=8.2, rgb=_WHITE)

            if row.get("emphasize"):
                _set_cell_bg(value_cell, accent_soft)
            _p(
                value_cell,
                row["value"],
                bold=bool(row.get("emphasize")),
                size=8.3,
                rgb=accent_rgb if row.get("emphasize") else _DARK,
                mono=bool(row.get("mono")),
                align=value_align,
            )

        _spacer(doc, 5)


# ── formatting helpers ─────────────────────────────────────────────────────────

def _q2(v) -> Decimal:
    if not isinstance(v, Decimal):
        v = Decimal(str(v))
    return v.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)


def _fmt_money(v, currency: str | None = "GBP (£)") -> str:
    if isinstance(v, str):
        return v
    try:
        return format_money(_q2(v), currency)
    except Exception:
        return str(v)


def _fmt_rate(v, places: int = 3) -> str:
    if isinstance(v, str):
        return v
    try:
        if not isinstance(v, Decimal):
            v = Decimal(str(v))
        q = Decimal("1." + "0" * places)
        return str(v.quantize(q, rounding=ROUND_HALF_UP))
    except Exception:
        return str(v)


# ── per-invoice page builder ───────────────────────────────────────────────────

def _render_invoice(
    doc: Document,
    company: dict,
    site: dict,
    rec: dict,
    page_no: int,
    total_pages: int,
    fp_label: str,
    strings: dict,
    omit: set[str],
    distractor_lines: list[str],
) -> None:
    strings = replace_pound_labels(strings, company.get("currency", "GBP (£)"))
    accent     = company["accent"]
    accent_soft = company["accent_soft"]
    accent_rgb = _hex_rgb(accent)
    soft_rgb   = _hex_rgb(accent_soft)

    # ── 1. Header ──────────────────────────────────────────────────────────────
    hdr = doc.add_table(rows=2, cols=2)
    _apply_table_borders(hdr, none=True)
    _set_col_widths(hdr, 11.0, 7.0)

    _p(hdr.cell(0, 0), company["supplier"], bold=True, size=14, rgb=accent_rgb)
    _p(hdr.cell(0, 1), strings["doc_title_heading"], bold=True, size=12,
       rgb=_DARK, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _p(hdr.cell(1, 0), strings["logo_subtitle"], size=8, rgb=_GRAY)
    _p(hdr.cell(1, 1),
       f"{fp_label}  •  {company['label']}  •  {rec['billing_period_label']}",
       size=7.5, rgb=_GRAY, align=WD_ALIGN_PARAGRAPH.RIGHT)

    _spacer(doc, 5)

    # ── 2. Address boxes ───────────────────────────────────────────────────────
    sup_lines  = list(company.get("supplier_address", []))
    cust_lines = list(site.get("customer_address", []))
    n_addr = max(len(sup_lines), len(cust_lines), 1)

    addr = doc.add_table(rows=n_addr + 1, cols=2)
    _apply_table_borders(addr)
    _set_col_widths(addr, 9.0, 9.0)

    _set_cell_bg(addr.cell(0, 0), accent)
    _set_cell_bg(addr.cell(0, 1), accent)
    _p(addr.cell(0, 0), strings["box_supplier"], bold=True, size=9, rgb=_WHITE)
    _p(addr.cell(0, 1), strings["box_customer"], bold=True, size=9, rgb=_WHITE)

    for i in range(n_addr):
        _p(addr.cell(i + 1, 0), sup_lines[i]  if i < len(sup_lines)  else "", size=8.5, rgb=_DARK)
        _p(addr.cell(i + 1, 1), cust_lines[i] if i < len(cust_lines) else "", size=8.5, rgb=_DARK)

    _spacer(doc, 5)

    # ── 3. Invoice metadata ────────────────────────────────────────────────────
    meta_lines = [
        f"{strings['meta_invoice_no']}:  {rec['invoice_no']}",
        f"{strings['meta_issue_date']}:  {rec['issue_date'].strftime('%d %b %Y') if hasattr(rec['issue_date'], 'strftime') else rec['issue_date']}",
        f"{strings['meta_due_date']}:  {rec['due_date'].strftime('%d %b %Y') if hasattr(rec['due_date'], 'strftime') else rec['due_date']}",
        f"{strings['meta_currency']}:  {company['currency']}",
        *distractor_lines,
    ]
    meta_rows = _paired_lines(meta_lines)
    meta = doc.add_table(rows=1 + len(meta_rows), cols=2)
    _apply_table_borders(meta)
    _set_col_widths(meta, 9.0, 9.0)

    hdr_cell = meta.cell(0, 0).merge(meta.cell(0, 1))
    _set_cell_bg(hdr_cell, accent)
    _p(hdr_cell, strings["box_invoice"], bold=True, size=9, rgb=_WHITE)

    for row_index, (left, right) in enumerate(meta_rows, start=1):
        _p(meta.cell(row_index, 0), left, size=8.5, rgb=_DARK)
        _p(meta.cell(row_index, 1), right, size=8.5, rgb=_DARK)

    _spacer(doc, 5)

    # ── 4. Billing fields table ────────────────────────────────────────────────
    billing_rows: list[tuple[str, str, bool]] = [
        # (label, value, is_mono)
        (strings["row_supplier"],      rec["supplier"],                                  False),
        (strings["row_customer"],      rec["customer"],                                  False),
        (strings["row_site"],          rec["site_label"],                                False),
        (strings["row_city"],          rec["city"],                                      False),
        (strings["row_postcode"],      rec["postcode"],                                  False),
        (strings["row_period_start"],  rec["period_start"].strftime("%d %b %Y") if hasattr(rec["period_start"], "strftime") else str(rec["period_start"]),         False),
        (strings["row_period_end"],    rec["period_end"].strftime("%d %b %Y") if hasattr(rec["period_end"], "strftime") else str(rec["period_end"]),           False),
        (strings["row_meter_id"],      rec["meter_id"],                                  True),
        (strings["row_prev_read"],     "" if "prev_read"   in omit else (f"{rec['prev_read']:,}" if isinstance(rec["prev_read"], int) else str(rec["prev_read"])),    False),
        (strings["row_curr_read"],     "" if "curr_read"   in omit else (f"{rec['curr_read']:,}" if isinstance(rec["curr_read"], int) else str(rec["curr_read"])),    False),
        (strings["row_consumption"],   "" if "consumption" in omit else (f"{rec['consumption']:,}" if isinstance(rec["consumption"], int) else str(rec["consumption"])),  False),
        (strings["row_unit_price"],    "" if "unit_price"  in omit else _fmt_rate(rec["unit_price"], 3), False),
        (strings["row_capacity"],      "" if "capacity_kw" in omit else str(rec["capacity_kw"]),    False),
        (strings["row_capacity_rate"], "" if "capacity_rate" in omit else _fmt_rate(rec["capacity_rate"], 2), False),
        (strings["row_supplier_ef"],   "" if "supplier_ef"   in omit else _fmt_rate(rec["supplier_ef"],   4), False),
    ]
    billing_rows = [row for row in billing_rows if row[0] != strings["row_site"] or _has_visible_value(row[1])]

    btbl = doc.add_table(rows=len(billing_rows) + 1, cols=2)
    _apply_table_borders(btbl)
    _set_col_widths(btbl, 10.2, 7.8)

    _set_cell_bg(btbl.cell(0, 0), accent)
    _set_cell_bg(btbl.cell(0, 1), accent)
    _p(btbl.cell(0, 0), strings["tbl_header_field"], bold=True, size=9, rgb=_WHITE)
    _p(btbl.cell(0, 1), strings["tbl_header_value"], bold=True, size=9, rgb=_WHITE)

    for i, (label, value, mono) in enumerate(billing_rows):
        row_idx = i + 1
        bg = accent_soft if i % 2 == 1 else None
        lc = btbl.cell(row_idx, 0)
        vc = btbl.cell(row_idx, 1)
        if bg:
            _set_cell_bg(lc, accent_soft)
            _set_cell_bg(vc, accent_soft)
        _p(lc, label,  size=8.5, rgb=_GRAY)
        _p(vc, value,  size=8.5, rgb=_DARK, mono=mono)

    _spacer(doc, 5)

    # ── 5. Charges summary ─────────────────────────────────────────────────────
    charge_rows: list[tuple[str, str, bool]] = [
        (strings["charge_heat"],     _fmt_money(rec["heat_cost"], company.get("currency")),       False),
        (strings["charge_capacity"], _fmt_money(rec["capacity_charge"], company.get("currency")), False),
        (strings["charge_subtotal"], _fmt_money(rec["subtotal"], company.get("currency")),        False),
        (strings["charge_vat"],      _fmt_money(rec["vat"], company.get("currency")),             False),
        (strings["charge_total"],    _fmt_money(rec["total"], company.get("currency")),           True),   # highlighted
    ]

    ctbl = doc.add_table(rows=len(charge_rows) + 1, cols=2)
    _apply_table_borders(ctbl)
    _set_col_widths(ctbl, 12.0, 6.0)

    hdr_c = ctbl.cell(0, 0).merge(ctbl.cell(0, 1))
    _set_cell_bg(hdr_c, accent)
    _p(hdr_c, strings["box_charges"], bold=True, size=9, rgb=_WHITE)

    for i, (label, value, is_total) in enumerate(charge_rows):
        lc = ctbl.cell(i + 1, 0)
        vc = ctbl.cell(i + 1, 1)
        if is_total:
            _set_cell_bg(lc, accent_soft)
            _set_cell_bg(vc, accent_soft)
            _p(lc, label, bold=True, size=9, rgb=accent_rgb)
            _p(vc, value, bold=True, size=9, rgb=accent_rgb,
               align=WD_ALIGN_PARAGRAPH.RIGHT)
        else:
            _p(lc, label, size=8.5, rgb=_GRAY)
            _p(vc, value, size=8.5, rgb=_DARK, align=WD_ALIGN_PARAGRAPH.RIGHT)

    _spacer(doc, 5)

    # ── 6. Footer ──────────────────────────────────────────────────────────────
    ftbl = doc.add_table(rows=1, cols=2)
    _apply_table_borders(ftbl, none=True)
    _set_col_widths(ftbl, 13.0, 5.0)

    _p(ftbl.cell(0, 0), strings["footer_vat"], size=7, rgb=_GRAY)
    _p(ftbl.cell(0, 1),
       strings["footer_page"].format(page=page_no, total=total_pages),
       size=7, rgb=_GRAY, align=WD_ALIGN_PARAGRAPH.RIGHT)


def _render_invoice_variant(
    doc: Document,
    company: dict,
    site: dict,
    rec: dict,
    page_no: int,
    total_pages: int,
    fp_label: str,
    strings: dict,
    omit: set[str],
    layout_plan: dict,
    distractor_lines: list[str],
) -> None:
    strings = replace_pound_labels(strings, company.get("currency", "GBP (£)"))
    accent = company["accent"]
    accent_soft = company["accent_soft"]
    accent_rgb = _hex_rgb(accent)

    hdr = doc.add_table(rows=2, cols=2)
    _apply_table_borders(hdr, none=True)
    _set_col_widths(hdr, 11.0, 7.0)

    _p(hdr.cell(0, 0), company["supplier"], bold=True, size=14, rgb=accent_rgb)
    _p(hdr.cell(0, 1), strings["doc_title_heading"], bold=True, size=12, rgb=_DARK, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _p(hdr.cell(1, 0), strings["logo_subtitle"], size=8, rgb=_GRAY)
    _p(
        hdr.cell(1, 1),
        f"{fp_label}  •  {company['label']}  •  {rec['billing_period_label']}",
        size=7.5,
        rgb=_GRAY,
        align=WD_ALIGN_PARAGRAPH.RIGHT,
    )
    _spacer(doc, 5)

    billing_rows = [
        {"label": strings["row_supplier"], "value": rec["supplier"]},
        {"label": strings["row_customer"], "value": rec["customer"]},
        {"label": strings["row_site"], "value": rec["site_label"]},
        {"label": strings["row_city"], "value": rec["city"]},
        {"label": strings["row_postcode"], "value": rec["postcode"]},
        {"label": strings["row_period_start"], "value": rec["period_start"].strftime("%d %b %Y") if hasattr(rec["period_start"], "strftime") else str(rec["period_start"])},
        {"label": strings["row_period_end"], "value": rec["period_end"].strftime("%d %b %Y") if hasattr(rec["period_end"], "strftime") else str(rec["period_end"])},
        {"label": strings["row_meter_id"], "value": rec["meter_id"], "mono": True},
        {"label": strings["row_prev_read"], "value": "" if "prev_read" in omit else (f"{rec['prev_read']:,}" if isinstance(rec["prev_read"], int) else str(rec["prev_read"]))},
        {"label": strings["row_curr_read"], "value": "" if "curr_read" in omit else (f"{rec['curr_read']:,}" if isinstance(rec["curr_read"], int) else str(rec["curr_read"]))},
        {"label": strings["row_consumption"], "value": "" if "consumption" in omit else (f"{rec['consumption']:,}" if isinstance(rec["consumption"], int) else str(rec["consumption"]))},
        {"label": strings["row_unit_price"], "value": "" if "unit_price" in omit else _fmt_rate(rec["unit_price"], 3)},
        {"label": strings["row_capacity"], "value": "" if "capacity_kw" in omit else str(rec["capacity_kw"])},
        {"label": strings["row_capacity_rate"], "value": "" if "capacity_rate" in omit else _fmt_rate(rec["capacity_rate"], 2)},
        {"label": strings["row_supplier_ef"], "value": "" if "supplier_ef" in omit else _fmt_rate(rec["supplier_ef"], 4)},
    ]
    billing_rows = [row for row in billing_rows if row["label"] != strings["row_site"] or _has_visible_value(row["value"])]
    charge_rows = [
        {"label": strings["charge_heat"], "value": _fmt_money(rec["heat_cost"], company.get("currency"))},
        {"label": strings["charge_capacity"], "value": _fmt_money(rec["capacity_charge"], company.get("currency"))},
        {"label": strings["charge_subtotal"], "value": _fmt_money(rec["subtotal"], company.get("currency"))},
        {"label": strings["charge_vat"], "value": _fmt_money(rec["vat"], company.get("currency"))},
        {"label": strings["charge_total"], "value": _fmt_money(rec["total"], company.get("currency")), "emphasize": True},
    ]

    def render_addresses() -> None:
        sup_lines = list(company.get("supplier_address", []))
        cust_lines = list(site.get("customer_address", []))
        n_addr = max(len(sup_lines), len(cust_lines), 1)
        addr = doc.add_table(rows=n_addr + 1, cols=2)
        _apply_table_borders(addr)
        _set_col_widths(addr, 9.0, 9.0)
        _set_cell_bg(addr.cell(0, 0), accent)
        _set_cell_bg(addr.cell(0, 1), accent)
        _p(addr.cell(0, 0), strings["box_supplier"], bold=True, size=9, rgb=_WHITE)
        _p(addr.cell(0, 1), strings["box_customer"], bold=True, size=9, rgb=_WHITE)
        for idx in range(n_addr):
            _p(addr.cell(idx + 1, 0), sup_lines[idx] if idx < len(sup_lines) else "", size=8.5, rgb=_DARK)
            _p(addr.cell(idx + 1, 1), cust_lines[idx] if idx < len(cust_lines) else "", size=8.5, rgb=_DARK)
        _spacer(doc, 5)

    def render_meta() -> None:
        meta_lines = [
            f"{strings['meta_invoice_no']}:  {rec['invoice_no']}",
            f"{strings['meta_issue_date']}:  {rec['issue_date'].strftime('%d %b %Y') if hasattr(rec['issue_date'], 'strftime') else rec['issue_date']}",
            f"{strings['meta_due_date']}:  {rec['due_date'].strftime('%d %b %Y') if hasattr(rec['due_date'], 'strftime') else rec['due_date']}",
            f"{strings['meta_currency']}:  {company['currency']}",
            *distractor_lines,
        ]
        meta_rows = _paired_lines(meta_lines)
        meta = doc.add_table(rows=1 + len(meta_rows), cols=2)
        _apply_table_borders(meta)
        _set_col_widths(meta, 9.0, 9.0)
        hdr_cell = meta.cell(0, 0).merge(meta.cell(0, 1))
        _set_cell_bg(hdr_cell, accent)
        _p(hdr_cell, strings["box_invoice"], bold=True, size=9, rgb=_WHITE)
        for row_index, (left, right) in enumerate(meta_rows, start=1):
            _p(meta.cell(row_index, 0), left, size=8.5, rgb=_DARK)
            _p(meta.cell(row_index, 1), right, size=8.5, rgb=_DARK)
        _spacer(doc, 5)

    def render_billing() -> None:
        if layout_plan.get("table_transforms", {}).get("billing_fields") == "transposed":
            _render_docx_transposed_rows(doc, billing_rows, accent, accent_soft, chunk_size=5)
            return

        table = doc.add_table(rows=len(billing_rows) + 1, cols=2)
        _apply_table_borders(table)
        _set_col_widths(table, 10.2, 7.8)
        _set_cell_bg(table.cell(0, 0), accent)
        _set_cell_bg(table.cell(0, 1), accent)
        _p(table.cell(0, 0), strings["tbl_header_field"], bold=True, size=9, rgb=_WHITE)
        _p(table.cell(0, 1), strings["tbl_header_value"], bold=True, size=9, rgb=_WHITE)
        for idx, row in enumerate(billing_rows, start=1):
            if idx % 2 == 0:
                _set_cell_bg(table.cell(idx, 0), accent_soft)
                _set_cell_bg(table.cell(idx, 1), accent_soft)
            _p(table.cell(idx, 0), row["label"], size=8.5, rgb=_GRAY)
            _p(table.cell(idx, 1), row["value"], size=8.5, rgb=_DARK, mono=bool(row.get("mono")))
        _spacer(doc, 5)

    def render_charges() -> None:
        if layout_plan.get("table_transforms", {}).get("charges") == "transposed":
            _render_docx_transposed_rows(doc, charge_rows, accent, accent_soft, chunk_size=5, value_align=WD_ALIGN_PARAGRAPH.RIGHT)
            return

        table = doc.add_table(rows=len(charge_rows) + 1, cols=2)
        _apply_table_borders(table)
        _set_col_widths(table, 12.0, 6.0)
        hdr_cell = table.cell(0, 0).merge(table.cell(0, 1))
        _set_cell_bg(hdr_cell, accent)
        _p(hdr_cell, strings["box_charges"], bold=True, size=9, rgb=_WHITE)
        for idx, row in enumerate(charge_rows, start=1):
            if row.get("emphasize"):
                _set_cell_bg(table.cell(idx, 0), accent_soft)
                _set_cell_bg(table.cell(idx, 1), accent_soft)
                _p(table.cell(idx, 0), row["label"], bold=True, size=9, rgb=accent_rgb)
                _p(table.cell(idx, 1), row["value"], bold=True, size=9, rgb=accent_rgb, align=WD_ALIGN_PARAGRAPH.RIGHT)
            else:
                _p(table.cell(idx, 0), row["label"], size=8.5, rgb=_GRAY)
                _p(table.cell(idx, 1), row["value"], size=8.5, rgb=_DARK, align=WD_ALIGN_PARAGRAPH.RIGHT)
        _spacer(doc, 5)

    def render_footer() -> None:
        ftbl = doc.add_table(rows=1, cols=2)
        _apply_table_borders(ftbl, none=True)
        _set_col_widths(ftbl, 13.0, 5.0)
        _p(ftbl.cell(0, 0), strings["footer_vat"], size=7, rgb=_GRAY)
        _p(ftbl.cell(0, 1), strings["footer_page"].format(page=page_no, total=total_pages), size=7, rgb=_GRAY, align=WD_ALIGN_PARAGRAPH.RIGHT)

    renderers = {
        "addresses": render_addresses,
        "meta": render_meta,
        "billing_fields": render_billing,
        "charges": render_charges,
    }
    section_order = [section for section in (layout_plan.get("section_order") or ["addresses", "meta", "billing_fields", "charges", "footer"]) if section in renderers]
    for section in section_order:
        renderers[section]()
    render_footer()


# ── public API ─────────────────────────────────────────────────────────────────

def _generate_heat_docx(
    config: dict,
    sections: list[dict],
    blank_fields: set[str] | None = None,
) -> bytes:
    """Build a DOCX billing document from sections and return bytes."""
    lang    = config["document"].get("language", "en")
    strings = TRANSLATIONS.get(lang, TRANSLATIONS["en"])
    omit    = blank_fields or set()

    doc = Document()

    # Page setup: A4, narrow margins
    page_sec = doc.sections[0]
    page_sec.page_height    = Cm(29.7)
    page_sec.page_width     = Cm(21.0)
    page_sec.left_margin    = Cm(1.5)
    page_sec.right_margin   = Cm(1.5)
    page_sec.top_margin     = Cm(1.5)
    page_sec.bottom_margin  = Cm(1.5)

    doc.core_properties.title   = config["document"].get("title", "")
    doc.core_properties.subject = config["document"].get("subject", "")

    # Remove the default empty paragraph python-docx adds on creation (if any)
    if doc.paragraphs:
        default_para = doc.paragraphs[0]._element
        default_para.getparent().remove(default_para)

    fp_label    = config["financial_period"]["label"]
    total_pages = sum(len(s["records"]) for s in sections)
    page_no     = 1
    first       = True
    layout_plan = config["document"].get("layout_plan", {})
    distractor_lines = _document_distractor_lines(config)

    for sec in sections:
        company = sec["company"]
        site    = sec["site"]
        for rec in sec["records"]:
            if not first:
                doc.add_page_break()
            first = False
            if layout_plan.get("enabled"):
                _render_invoice_variant(
                    doc, company, site, rec,
                    page_no, total_pages, fp_label, strings, omit, layout_plan, distractor_lines,
                )
            else:
                _render_invoice(
                    doc, company, site, rec,
                    page_no, total_pages, fp_label, strings, omit, distractor_lines,
                )
            page_no += 1

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _generate_electricity_docx(config: dict, sections: list[dict]) -> bytes:
    layout_plan = config["document"].get("layout_plan", {})
    if layout_plan.get("enabled"):
        return _generate_electricity_docx_variant(config, sections, layout_plan)

    lang = config["document"].get("language", "en")
    strings = ELECTRICITY_TRANSLATIONS.get(lang, ELECTRICITY_TRANSLATIONS["en"])
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Aptos"
    style.font.size = Pt(9)

    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    doc.core_properties.title = config["document"].get("title", "")
    doc.core_properties.subject = config["document"].get("subject", "")
    distractor_lines = _document_distractor_lines(config)

    if doc.paragraphs:
        default_para = doc.paragraphs[0]._element
        default_para.getparent().remove(default_para)

    def add_kv_table(title: str, rows: list[tuple[str, str]], accent: str, accent_rgb: tuple[int, int, int]) -> None:
        table = doc.add_table(rows=len(rows) + 1, cols=2)
        _apply_table_borders(table)
        _set_col_widths(table, 10.0, 8.0)

        _set_cell_bg(table.cell(0, 0), accent)
        _set_cell_bg(table.cell(0, 1), accent)
        _p(table.cell(0, 0), title, bold=True, size=9.2, rgb=_WHITE)
        _p(table.cell(0, 1), "Value", bold=True, size=9.2, rgb=_WHITE)

        for idx, (label, value) in enumerate(rows, start=1):
            if idx % 2 == 0:
                _set_cell_bg(table.cell(idx, 0), "F7F9FB")
                _set_cell_bg(table.cell(idx, 1), "F7F9FB")
            _p(table.cell(idx, 0), label, size=8.6, rgb=_GRAY)
            _p(table.cell(idx, 1), value, size=8.6, rgb=_DARK)

        _spacer(doc, 5)

    for page_no, section in enumerate(sections, start=1):
        if page_no > 1:
            doc.add_page_break()

        company = section["company"]
        site = section["site"]
        accent = company["accent"]
        accent_rgb = _hex_rgb(accent)
        site_omit = site.get("_omit", {})
        show_readings = not site_omit.get("start_reading")
        show_costs = not site_omit.get("total_cost")
        show_emissions = not site_omit.get("supplier_ef")
        unit = site["unit"]
        symbol = site.get("currency_symbol", "")

        header = doc.add_table(rows=2, cols=2)
        _apply_table_borders(header, none=True)
        _set_col_widths(header, 11.2, 6.8)
        _p(header.cell(0, 0), company["supplier"], bold=True, size=16, rgb=accent_rgb)
        _p(header.cell(0, 1), strings["doc_title_heading"], bold=True, size=13.5, rgb=_DARK, align=WD_ALIGN_PARAGRAPH.RIGHT)
        _p(header.cell(1, 0), strings["doc_subtitle"], size=8.3, rgb=_GRAY)
        _p(
            header.cell(1, 1),
            _join_display_parts(site["billing_period_label"], company["label"], site["label"]),
            size=8.0,
            rgb=_GRAY,
            align=WD_ALIGN_PARAGRAPH.RIGHT,
        )
        _spacer(doc, 5)

        supplier_lines = list(company.get("supplier_address", []))
        customer_lines = list(site.get("customer_address", []))
        info_rows = max(len(supplier_lines), len(customer_lines), 1)
        info = doc.add_table(rows=info_rows + 1, cols=2)
        _apply_table_borders(info)
        _set_col_widths(info, 9.0, 9.0)
        _set_cell_bg(info.cell(0, 0), accent)
        _set_cell_bg(info.cell(0, 1), accent)
        _p(info.cell(0, 0), strings["box_supplier"], bold=True, size=9, rgb=_WHITE)
        _p(info.cell(0, 1), strings["box_customer"], bold=True, size=9, rgb=_WHITE)
        for row_idx in range(info_rows):
            _p(info.cell(row_idx + 1, 0), supplier_lines[row_idx] if row_idx < len(supplier_lines) else "", size=8.4, rgb=_DARK)
            _p(info.cell(row_idx + 1, 1), customer_lines[row_idx] if row_idx < len(customer_lines) else "", size=8.4, rgb=_DARK)
        _spacer(doc, 5)

        meta_lines = [
            f"{strings.get('meta_period_label', 'Period')}:  {site['billing_period_label']}",
            f"{strings.get('meta_ref', 'Reference')}:  {site['ref_no']}",
            f"{strings['meta_period_start']}:  {site['period_start'].strftime('%d %b %Y') if hasattr(site['period_start'], 'strftime') else site['period_start']}",
            f"{strings['meta_period_end']}:  {site['period_end'].strftime('%d %b %Y') if hasattr(site['period_end'], 'strftime') else site['period_end']}",
            f"{strings['meta_currency']}: {company['currency']}",
            *distractor_lines,
        ]
        meta_rows = _paired_lines(meta_lines)
        meta = doc.add_table(rows=1 + len(meta_rows), cols=2)
        _apply_table_borders(meta)
        _set_col_widths(meta, 9.0, 9.0)
        meta_hdr = meta.cell(0, 0).merge(meta.cell(0, 1))
        _set_cell_bg(meta_hdr, accent)
        _p(meta_hdr, strings.get("box_period", "Statement Period"), bold=True, size=9, rgb=_WHITE)
        for row_index, (left, right) in enumerate(meta_rows, start=1):
            _p(meta.cell(row_index, 0), left, size=8.4, rgb=_DARK)
            _p(meta.cell(row_index, 1), right, size=8.4, rgb=_DARK)
        _spacer(doc, 5)

        meter_rows = [
            (strings["row_site"], site["label"]),
            (strings["row_city"], site["city"]),
            (strings["row_postcode"], site["postcode"]),
            (strings["row_meter_id"], site["meter_id"]),
            (strings["row_unit"], unit),
            (strings["row_total_qty"], f"{float(site['total_quantity']):,.2f} {unit}" if not isinstance(site["total_quantity"], str) else f"{site['total_quantity']} {unit}"),
        ]
        meter_rows = [row for row in meter_rows if row[0] != strings["row_site"] or _has_visible_value(row[1])]
        if show_readings:
            meter_rows.insert(5, (strings["row_start_read"], f"{site['start_reading']:,}" if isinstance(site["start_reading"], int) else str(site["start_reading"])))
            meter_rows.insert(6, (strings["row_end_read"], f"{site['end_reading']:,}" if isinstance(site["end_reading"], int) else str(site["end_reading"])))
        add_kv_table(strings["tbl_meter"], meter_rows, accent, accent_rgb)

        if show_emissions:
            add_kv_table(strings["tbl_grid"], [
                (strings["row_supplier_ef"], f"{float(site['supplier_ef']):.4f}" if not isinstance(site["supplier_ef"], str) else site["supplier_ef"]),
                (strings["row_emissions_kg"], f"{float(site['emissions_kg']):,.2f}" if not isinstance(site.get("emissions_kg"), str) else site["emissions_kg"]),
                (strings["row_emissions_t"], f"{float(site['emissions_t']):.3f}" if not isinstance(site.get("emissions_t"), str) else site["emissions_t"]),
            ], accent, accent_rgb)

        if site["tariffs"]:
            tariff_headers = [
                strings["col_tariff_name"],
                strings["col_tariff_qty"],
                strings["col_tariff_unit"],
            ]
            if show_costs:
                tariff_headers.extend([strings["col_tariff_rate"], strings["col_tariff_cost"]])

            tariff_table = doc.add_table(rows=len(site["tariffs"]) + 1, cols=len(tariff_headers))
            _apply_table_borders(tariff_table)
            for col_idx, header_text in enumerate(tariff_headers):
                _set_cell_bg(tariff_table.cell(0, col_idx), accent)
                _p(tariff_table.cell(0, col_idx), header_text, bold=True, size=8.6, rgb=_WHITE)

            for row_idx, tariff in enumerate(site["tariffs"], start=1):
                if row_idx % 2 == 0:
                    for col_idx in range(len(tariff_headers)):
                        _set_cell_bg(tariff_table.cell(row_idx, col_idx), "F7F9FB")
                values = [
                    tariff["name"],
                    f"{float(tariff['quantity']):,.2f}",
                    tariff["unit"],
                ]
                if show_costs:
                    values.extend([
                        f"{symbol}{float(tariff['unit_cost']):.4f}",
                        f"{symbol}{float(tariff['cost']):,.2f}",
                    ])
                for col_idx, value in enumerate(values):
                    _p(tariff_table.cell(row_idx, col_idx), str(value), size=8.4, rgb=_DARK)
            _spacer(doc, 5)

        if show_costs:
            total_table = doc.add_table(rows=1, cols=2)
            _apply_table_borders(total_table)
            _set_col_widths(total_table, 12.5, 5.5)
            _set_cell_bg(total_table.cell(0, 0), company["accent_soft"])
            _set_cell_bg(total_table.cell(0, 1), company["accent_soft"])
            _p(total_table.cell(0, 0), strings["charge_total"], bold=True, size=10.2, rgb=accent_rgb)
            _p(
                total_table.cell(0, 1),
                str(site['total_cost']) if isinstance(site['total_cost'], str) else f"{symbol}{float(site['total_cost']):,.2f}",
                bold=True,
                size=10.2,
                rgb=accent_rgb,
                align=WD_ALIGN_PARAGRAPH.RIGHT,
            )

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _generate_electricity_docx_variant(config: dict, sections: list[dict], layout_plan: dict) -> bytes:
    lang = config["document"].get("language", "en")
    strings = ELECTRICITY_TRANSLATIONS.get(lang, ELECTRICITY_TRANSLATIONS["en"])
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Aptos"
    style.font.size = Pt(9)

    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    doc.core_properties.title = config["document"].get("title", "")
    doc.core_properties.subject = config["document"].get("subject", "")
    distractor_lines = _document_distractor_lines(config)

    if doc.paragraphs:
        default_para = doc.paragraphs[0]._element
        default_para.getparent().remove(default_para)

    total_pages = len(sections)

    def add_vertical_table(title: str, rows: list[dict], accent: str) -> None:
        table = doc.add_table(rows=len(rows) + 1, cols=2)
        _apply_table_borders(table)
        _set_col_widths(table, 10.0, 8.0)

        _set_cell_bg(table.cell(0, 0), accent)
        _set_cell_bg(table.cell(0, 1), accent)
        _p(table.cell(0, 0), title, bold=True, size=9.2, rgb=_WHITE)
        _p(table.cell(0, 1), "Value", bold=True, size=9.2, rgb=_WHITE)

        for idx, row in enumerate(rows, start=1):
            if idx % 2 == 0:
                _set_cell_bg(table.cell(idx, 0), "F7F9FB")
                _set_cell_bg(table.cell(idx, 1), "F7F9FB")
            _p(table.cell(idx, 0), row["label"], size=8.6, rgb=_GRAY)
            _p(table.cell(idx, 1), row["value"], size=8.6, rgb=_DARK, mono=bool(row.get("mono")))

        _spacer(doc, 5)

    for page_no, section in enumerate(sections, start=1):
        if page_no > 1:
            doc.add_page_break()

        company = section["company"]
        site = section["site"]
        accent = company["accent"]
        accent_soft = company["accent_soft"]
        accent_rgb = _hex_rgb(accent)
        site_omit = site.get("_omit", {})
        show_readings = not site_omit.get("start_reading")
        show_costs = not site_omit.get("total_cost")
        show_emissions = not site_omit.get("supplier_ef")
        unit = site["unit"]
        symbol = site.get("currency_symbol", "")

        header = doc.add_table(rows=2, cols=2)
        _apply_table_borders(header, none=True)
        _set_col_widths(header, 11.2, 6.8)
        _p(header.cell(0, 0), company["supplier"], bold=True, size=16, rgb=accent_rgb)
        _p(header.cell(0, 1), strings["doc_title_heading"], bold=True, size=13.5, rgb=_DARK, align=WD_ALIGN_PARAGRAPH.RIGHT)
        _p(header.cell(1, 0), strings["doc_subtitle"], size=8.3, rgb=_GRAY)
        _p(
            header.cell(1, 1),
            _join_display_parts(site["billing_period_label"], company["label"], site["label"]),
            size=8.0,
            rgb=_GRAY,
            align=WD_ALIGN_PARAGRAPH.RIGHT,
        )
        _spacer(doc, 5)

        meter_rows = [
            {"label": strings["row_site"], "value": site["label"]},
            {"label": strings["row_city"], "value": site["city"]},
            {"label": strings["row_postcode"], "value": site["postcode"]},
            {"label": strings["row_meter_id"], "value": site["meter_id"], "mono": True},
            {"label": strings["row_unit"], "value": unit},
            {
                "label": strings["row_total_qty"],
                "value": f"{float(site['total_quantity']):,.2f} {unit}" if not isinstance(site["total_quantity"], str) else f"{site['total_quantity']} {unit}",
            },
        ]
        meter_rows = [row for row in meter_rows if row["label"] != strings["row_site"] or _has_visible_value(row["value"])]
        if show_readings:
            meter_rows.insert(5, {"label": strings["row_start_read"], "value": f"{site['start_reading']:,}" if isinstance(site["start_reading"], int) else str(site["start_reading"])})
            meter_rows.insert(6, {"label": strings["row_end_read"], "value": f"{site['end_reading']:,}" if isinstance(site["end_reading"], int) else str(site["end_reading"])})

        grid_rows = []
        if show_emissions:
            grid_rows = [
                {"label": strings["row_supplier_ef"], "value": f"{float(site['supplier_ef']):.4f}" if not isinstance(site["supplier_ef"], str) else site["supplier_ef"]},
                {"label": strings["row_emissions_kg"], "value": f"{float(site['emissions_kg']):,.2f}" if not isinstance(site.get("emissions_kg"), str) else site["emissions_kg"]},
                {"label": strings["row_emissions_t"], "value": f"{float(site['emissions_t']):.3f}" if not isinstance(site.get("emissions_t"), str) else site["emissions_t"]},
            ]

        def render_addresses() -> None:
            supplier_lines = list(company.get("supplier_address", []))
            customer_lines = list(site.get("customer_address", []))
            info_rows = max(len(supplier_lines), len(customer_lines), 1)
            info = doc.add_table(rows=info_rows + 1, cols=2)
            _apply_table_borders(info)
            _set_col_widths(info, 9.0, 9.0)
            _set_cell_bg(info.cell(0, 0), accent)
            _set_cell_bg(info.cell(0, 1), accent)
            _p(info.cell(0, 0), strings["box_supplier"], bold=True, size=9, rgb=_WHITE)
            _p(info.cell(0, 1), strings["box_customer"], bold=True, size=9, rgb=_WHITE)
            for row_idx in range(info_rows):
                _p(info.cell(row_idx + 1, 0), supplier_lines[row_idx] if row_idx < len(supplier_lines) else "", size=8.4, rgb=_DARK)
                _p(info.cell(row_idx + 1, 1), customer_lines[row_idx] if row_idx < len(customer_lines) else "", size=8.4, rgb=_DARK)
            _spacer(doc, 5)

        def render_period_meta() -> None:
            meta_lines = [
                f"{strings.get('meta_period_label', 'Period')}:  {site['billing_period_label']}",
                f"{strings.get('meta_ref', 'Reference')}:  {site['ref_no']}",
                f"{strings['meta_period_start']}:  {site['period_start'].strftime('%d %b %Y') if hasattr(site['period_start'], 'strftime') else site['period_start']}",
                f"{strings['meta_period_end']}:  {site['period_end'].strftime('%d %b %Y') if hasattr(site['period_end'], 'strftime') else site['period_end']}",
                f"{strings['meta_currency']}: {company['currency']}",
                *distractor_lines,
            ]
            meta_rows = _paired_lines(meta_lines)
            meta = doc.add_table(rows=1 + len(meta_rows), cols=2)
            _apply_table_borders(meta)
            _set_col_widths(meta, 9.0, 9.0)
            meta_hdr = meta.cell(0, 0).merge(meta.cell(0, 1))
            _set_cell_bg(meta_hdr, accent)
            _p(meta_hdr, strings.get("box_period", "Statement Period"), bold=True, size=9, rgb=_WHITE)
            for row_index, (left, right) in enumerate(meta_rows, start=1):
                _p(meta.cell(row_index, 0), left, size=8.4, rgb=_DARK)
                _p(meta.cell(row_index, 1), right, size=8.4, rgb=_DARK)
            _spacer(doc, 5)

        def render_meter_table() -> None:
            if layout_plan.get("table_transforms", {}).get("meter_table") == "transposed":
                _render_docx_transposed_rows(doc, meter_rows, accent, accent_soft, chunk_size=4)
                return
            add_vertical_table(strings["tbl_meter"], meter_rows, accent)

        def render_grid_table() -> None:
            if not grid_rows:
                return
            if layout_plan.get("table_transforms", {}).get("grid_table") == "transposed":
                _render_docx_transposed_rows(doc, grid_rows, accent, accent_soft, chunk_size=3)
                return
            add_vertical_table(strings["tbl_grid"], grid_rows, accent)

        def render_tariff_table() -> None:
            if not site["tariffs"]:
                return

            tariff_headers = [
                strings["col_tariff_name"],
                strings["col_tariff_qty"],
                strings["col_tariff_unit"],
            ]
            if show_costs:
                tariff_headers.extend([strings["col_tariff_rate"], strings["col_tariff_cost"]])

            tariff_table = doc.add_table(rows=len(site["tariffs"]) + 1, cols=len(tariff_headers))
            _apply_table_borders(tariff_table)
            for col_idx, header_text in enumerate(tariff_headers):
                _set_cell_bg(tariff_table.cell(0, col_idx), accent)
                _p(tariff_table.cell(0, col_idx), header_text, bold=True, size=8.6, rgb=_WHITE)

            for row_idx, tariff in enumerate(site["tariffs"], start=1):
                if row_idx % 2 == 0:
                    for col_idx in range(len(tariff_headers)):
                        _set_cell_bg(tariff_table.cell(row_idx, col_idx), "F7F9FB")
                values = [
                    tariff["name"],
                    f"{float(tariff['quantity']):,.2f}",
                    tariff["unit"],
                ]
                if show_costs:
                    values.extend([
                        f"{symbol}{float(tariff['unit_cost']):.4f}",
                        f"{symbol}{float(tariff['cost']):,.2f}",
                    ])
                for col_idx, value in enumerate(values):
                    _p(tariff_table.cell(row_idx, col_idx), str(value), size=8.4, rgb=_DARK)
            _spacer(doc, 5)

        def render_total_box() -> None:
            if not show_costs:
                return

            total_table = doc.add_table(rows=1, cols=2)
            _apply_table_borders(total_table)
            _set_col_widths(total_table, 12.5, 5.5)
            _set_cell_bg(total_table.cell(0, 0), company["accent_soft"])
            _set_cell_bg(total_table.cell(0, 1), company["accent_soft"])
            _p(total_table.cell(0, 0), strings["charge_total"], bold=True, size=10.2, rgb=accent_rgb)
            _p(
                total_table.cell(0, 1),
                str(site['total_cost']) if isinstance(site['total_cost'], str) else f"{symbol}{float(site['total_cost']):,.2f}",
                bold=True,
                size=10.2,
                rgb=accent_rgb,
                align=WD_ALIGN_PARAGRAPH.RIGHT,
            )
            _spacer(doc, 5)

        def render_footer() -> None:
            footer = doc.add_table(rows=1, cols=2)
            _apply_table_borders(footer, none=True)
            _set_col_widths(footer, 13.0, 5.0)
            _p(footer.cell(0, 0), strings["doc_subtitle"], size=7.2, rgb=_GRAY)
            _p(footer.cell(0, 1), f"Page {page_no} / {total_pages}", size=7.2, rgb=_GRAY, align=WD_ALIGN_PARAGRAPH.RIGHT)

        renderers = {
            "addresses": render_addresses,
            "period_meta": render_period_meta,
            "meter_table": render_meter_table,
            "grid_table": render_grid_table,
            "tariff_table": render_tariff_table,
            "total_box": render_total_box,
            "footer": render_footer,
        }
        section_order = [
            section_name
            for section_name in (
                layout_plan.get("section_order")
                or ["addresses", "period_meta", "meter_table", "grid_table", "tariff_table", "total_box", "footer"]
            )
            if section_name in renderers
        ]
        for section_name in section_order:
            renderers[section_name]()

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_docx(
    config: dict,
    sections: list[dict],
    blank_fields: set[str] | None = None,
    category: str = "heat",
) -> bytes:
    if category == "electricity":
        return _generate_electricity_docx(config, sections)
    return _generate_heat_docx(config, sections, blank_fields=blank_fields)
