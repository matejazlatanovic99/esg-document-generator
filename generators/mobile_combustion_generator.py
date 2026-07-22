from __future__ import annotations

import csv
import json
import random
from datetime import date, datetime, time, timedelta
from decimal import Decimal
from io import BytesIO, StringIO
from typing import Any

import openpyxl
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

from generators.stationary_combustion_generator import (
    _augment_stationary_field_ids as _augment_field_ids,
    _currency_symbol,
    _draw_fitted_string,
    _draw_multiline,
    _financial_period,
    _fmt_date,
    _fmt_money,
    _fmt_num,
    _months_in_range,
    _multi_document_count,
    _ordered_stationary_field_ids as _ordered_field_ids,
    _parse_decimal,
    _q2,
    _set_docx_cell_text,
    _shade_docx_cell,
    _slugify_filename,
    _stationary_distractor_field_map as _distractor_field_map,
    _stationary_document_lines as _document_lines,
    _stationary_row_values as _row_values,
    _style_docx_document,
    _with_special_chars,
    _write_stationary_csv_preamble as _write_csv_preamble,
    _write_stationary_xlsx_preamble as _write_xlsx_preamble,
    _zip_documents,
)
from utils.bad_data import (
    FTYPE_CURRENCY_UNIT,
    FTYPE_DATE_TIME,
    FTYPE_IDENTIFIER,
    FTYPE_NUMERIC,
    FTYPE_TEXT,
    corrupt_records,
    get_bad_data_config,
)
from utils.distractor_fields import resolve_distractor_plan
from utils.layouts import resolve_layout_plan

PAGE_W, PAGE_H = A4

SCOPE_LABEL = "Scope 1"
CATEGORY_MOBILE = "mobile_combustion"
CATEGORY_STATIONARY = "stationary_combustion"

# ── field-type maps for invalid-data corruption ──────────────────────────────

_MOBILE_INVOICE_FIELD_TYPES: dict[str, str] = {
    "company": FTYPE_TEXT,
    "supplier": FTYPE_TEXT,
    "customer": FTYPE_TEXT,
    "site": FTYPE_TEXT,
    "vehicle_reg": FTYPE_IDENTIFIER,
    "driver": FTYPE_TEXT,
    "fuel": FTYPE_TEXT,
    "account_number": FTYPE_IDENTIFIER,
    "invoice_no": FTYPE_IDENTIFIER,
    "invoiced_date": FTYPE_DATE_TIME,
    "due_date": FTYPE_DATE_TIME,
    "quantity": FTYPE_NUMERIC,
    "unit_price": FTYPE_NUMERIC,
    "fuel_cost": FTYPE_NUMERIC,
    "subtotal": FTYPE_NUMERIC,
    "vat": FTYPE_NUMERIC,
    "total": FTYPE_NUMERIC,
    "currency": FTYPE_CURRENCY_UNIT,
    "unit": FTYPE_CURRENCY_UNIT,
}

_MOBILE_FUEL_CARD_FIELD_TYPES: dict[str, str] = {
    "vehicle_reg": FTYPE_IDENTIFIER,
    "driver": FTYPE_TEXT,
    "merchant": FTYPE_TEXT,
    "site": FTYPE_TEXT,
    "fuel": FTYPE_TEXT,
    "card_number": FTYPE_IDENTIFIER,
    "receipt_no": FTYPE_IDENTIFIER,
    "quantity": FTYPE_NUMERIC,
    "unit_price": FTYPE_NUMERIC,
    "total": FTYPE_NUMERIC,
    "odometer": FTYPE_NUMERIC,
    "unit": FTYPE_CURRENCY_UNIT,
}

_TELEMATICS_FUEL_FIELD_TYPES: dict[str, str] = {
    "vehicle_reg": FTYPE_IDENTIFIER,
    "vehicle_name": FTYPE_TEXT,
    "vehicle_type": FTYPE_TEXT,
    "fuel": FTYPE_TEXT,
    "period_start": FTYPE_DATE_TIME,
    "period_end": FTYPE_DATE_TIME,
    "distance": FTYPE_NUMERIC,
    "fuel_used": FTYPE_NUMERIC,
    "idle_fuel": FTYPE_NUMERIC,
    "engine_hours": FTYPE_NUMERIC,
    "avg_consumption": FTYPE_NUMERIC,
    "unit": FTYPE_CURRENCY_UNIT,
    "distance_unit": FTYPE_CURRENCY_UNIT,
}

_TELEMATICS_TRIP_FIELD_TYPES: dict[str, str] = {
    "trip_id": FTYPE_IDENTIFIER,
    "vehicle_reg": FTYPE_IDENTIFIER,
    "driver": FTYPE_TEXT,
    "trip_start": FTYPE_DATE_TIME,
    "trip_end": FTYPE_DATE_TIME,
    "start_location": FTYPE_TEXT,
    "end_location": FTYPE_TEXT,
    "distance": FTYPE_NUMERIC,
    "avg_speed": FTYPE_NUMERIC,
    "fuel": FTYPE_TEXT,
    "distance_unit": FTYPE_CURRENCY_UNIT,
}

_TELEMATICS_ODOMETER_FIELD_TYPES: dict[str, str] = {
    "vehicle_reg": FTYPE_IDENTIFIER,
    "vehicle_name": FTYPE_TEXT,
    "fuel": FTYPE_TEXT,
    "period_start": FTYPE_DATE_TIME,
    "period_end": FTYPE_DATE_TIME,
    "odometer_start": FTYPE_NUMERIC,
    "odometer_end": FTYPE_NUMERIC,
    "distance": FTYPE_NUMERIC,
    "distance_unit": FTYPE_CURRENCY_UNIT,
}

# ── translations ─────────────────────────────────────────────────────────────

_TRANSLATIONS: dict[str, dict[str, str]] = {
    "en": {
        "fuel_invoice_title": "Fuel Invoice",
        "bill_to": "Bill To",
        "invoice_details": "Invoice Details",
        "vehicle_details": "Vehicle",
        "invoice_no": "Invoice No",
        "account_no": "Account No",
        "invoiced_date": "Invoice Date",
        "due_date": "Due Date",
        "currency": "Currency",
        "registration": "Registration",
        "driver": "Driver",
        "site": "Site / Depot",
        "product": "Product",
        "description": "Description",
        "quantity": "Quantity",
        "unit": "Unit",
        "unit_price": "Unit Price",
        "amount": "Amount",
        "each": "Each",
        "subtotal": "Subtotal",
        "total": "Total",
        "vat": "VAT",
        "service_fee": "Service Fee",
        "delivery_charge": "Delivery Charge",
        "fuel_invoice_footer": "Generated for Scope 1 mobile combustion. Vehicle details are illustrative.",
        "fuel_card_title": "Fuel Card Statement",
        "account_name": "Account Name",
        "provider": "Provider",
        "statement_no": "Statement No",
        "statement_period": "Statement Period",
        "issue_date": "Issue Date",
        "card_no": "Card No",
        "date": "Date",
        "time": "Time",
        "merchant": "Merchant",
        "location": "Location",
        "receipt_no": "Receipt No",
        "odometer": "Odometer",
        "qty": "Qty",
        "net_amount": "Net Amount",
        "gross_amount": "Gross Amount",
        "account_details": "Account Details",
        "statement_details": "Statement Details",
        "fuel_card_footer": "Statement generated for mobile combustion fuel-card QA.",
        "telematics_fuel_title": "Telematics Fuel Usage Report",
        "telematics_trips_title": "Telematics Trip History Export",
        "telematics_odometer_title": "Fleet Mileage / Odometer Summary",
        "reporting_period": "Reporting Period",
        "period_start": "Period Start",
        "period_end": "Period End",
        "vehicle_reg": "Vehicle Registration",
        "vehicle_name": "Vehicle",
        "vehicle_type": "Vehicle Type",
        "fuel_type": "Fuel Type",
        "distance": "Distance",
        "distance_unit": "Distance Unit",
        "fuel_used": "Fuel Consumed",
        "idle_fuel": "Idle Fuel",
        "engine_hours": "Engine Hours",
        "avg_consumption": "Avg Consumption",
        "trip_id": "Trip ID",
        "trip_start": "Trip Start",
        "trip_end": "Trip End",
        "start_location": "Start Location",
        "end_location": "End Location",
        "duration": "Duration",
        "avg_speed": "Avg Speed",
        "odometer_start": "Odometer Start",
        "odometer_end": "Odometer End",
        "fleet": "Fleet",
        "generated": "Generated",
        "telematics_footer": "Telematics export generated for Scope 1 mobile combustion QA.",
        "equipment_ref": "Equipment",
    },
}


def _language(raw_config: dict) -> str:
    return str(raw_config.get("document", {}).get("language", "en"))


def _tr(raw_config: dict, key: str) -> str:
    table = _TRANSLATIONS.get(_language(raw_config), _TRANSLATIONS["en"])
    return table.get(key, _TRANSLATIONS["en"].get(key, key))


def _header_text(
    raw_config: dict,
    plan: dict,
    field_id: str,
    default_key: str,
    distractor_fields: dict[str, Any] | None = None,
) -> str:
    if distractor_fields and field_id in distractor_fields:
        return str(distractor_fields[field_id].label)
    return plan.get("header_aliases", {}).get(field_id) or _tr(raw_config, default_key)


# ── layout / distractor plan resolution ──────────────────────────────────────


def _layout_context(raw_config: dict) -> dict:
    return {
        "language": _language(raw_config),
        "period_label": raw_config.get("financial_period", {}).get("label", ""),
        "company_labels": [
            company.get("label", f"Company {index + 1}")
            for index, company in enumerate(raw_config.get("companies", []))
        ],
    }


def _document_type(raw_config: dict) -> str:
    return str(
        raw_config.get("document_type")
        or raw_config.get("document", {}).get("type")
        or "fuel_invoice"
    )


def _layout_plan(raw_config: dict, output_format: str) -> dict:
    return resolve_layout_plan(
        raw_config.get("document", {}),
        random_seed=int(raw_config.get("random_seed", 42)),
        category=CATEGORY_MOBILE,
        document_type=_document_type(raw_config),
        output_format=output_format,
        context=_layout_context(raw_config),
    )


def _distractor_plan(raw_config: dict, output_format: str):
    return resolve_distractor_plan(
        raw_config.get("document", {}),
        random_seed=int(raw_config.get("random_seed", 42)),
        category=CATEGORY_MOBILE,
        document_type=_document_type(raw_config),
        output_format=output_format,
        context=_layout_context(raw_config),
    )


# ── config helpers ───────────────────────────────────────────────────────────


def _document_option(raw_config: dict, key: str, default=None):
    return raw_config.get("document", {}).get(key, default)


def _distance_unit(raw_config: dict) -> str:
    return str(_document_option(raw_config, "distance_unit", "km") or "km")


def _cross_scope_enabled(raw_config: dict) -> bool:
    return bool(_document_option(raw_config, "cross_scope_items", False))


def _vat_rate(raw_config: dict, fallback: str = "20") -> Decimal:
    return _parse_decimal(_document_option(raw_config, "vat_rate", fallback), fallback)


def _iter_company_vehicles(raw_config: dict):
    for company_index, company in enumerate(raw_config.get("companies", []), start=1):
        for vehicle_index, vehicle in enumerate(company.get("vehicles", []), start=1):
            yield company_index, vehicle_index, company, vehicle


def _account_number(company: dict, seed: int, company_index: int) -> str:
    explicit = str(company.get("account_number", "") or "").strip()
    if explicit:
        return explicit
    rng = random.Random(f"{seed}:mobile_account:{company_index}")
    prefix = (company.get("customer_code") or "ACC").strip() or "ACC"
    return f"{prefix}-{rng.randint(100000, 999999)}"


def _receipt_number(rng: random.Random) -> str:
    return f"T-{rng.randint(10_000_000, 99_999_999)}"


_CROSS_SCOPE_STATIONARY_LINES = [
    {"fuel": "Gas Oil (Red Diesel)", "equipment": "Backup Generator"},
    {"fuel": "Heating Oil", "equipment": "Boiler Tank 1"},
    {"fuel": "Kerosene", "equipment": "Site Heater"},
]

_TRIP_LOCATIONS = [
    "London Depot",
    "Manchester DC",
    "Birmingham Hub",
    "Leeds Depot",
    "Bristol Yard",
    "Glasgow Terminal",
    "Sheffield Depot",
    "Reading Park",
]


# ── record builders ──────────────────────────────────────────────────────────


def _build_invoice_base_records(raw_config: dict) -> list[dict]:
    """One template record per configured company/vehicle; extra invoices
    beyond the configured vehicles reuse them with fresh randomized values."""
    seed = int(raw_config.get("random_seed", 42))
    records: list[dict] = []

    for company_index, vehicle_index, company, vehicle in _iter_company_vehicles(raw_config):
        vehicle_omit = vehicle.get("_omit", {})
        quantity = _q2(_parse_decimal(vehicle.get("quantity"), "0"))
        unit_price = _q2(_parse_decimal(vehicle.get("unit_price"), "0"))
        records.append({
            "company": _with_special_chars(raw_config, company.get("label", "")),
            "supplier": _with_special_chars(raw_config, company.get("supplier", "")),
            "supplier_code": company.get("supplier_code", "INV"),
            "supplier_address": [
                _with_special_chars(raw_config, line) for line in company.get("supplier_address", [])
            ],
            "customer": _with_special_chars(raw_config, company.get("customer", "")),
            "customer_code": company.get("customer_code", ""),
            "account_number": _account_number(company, seed, company_index),
            "site": "" if vehicle_omit.get("site", False) else _with_special_chars(raw_config, vehicle.get("site", "")),
            "vehicle_reg": _with_special_chars(raw_config, vehicle.get("registration", "")),
            "vehicle_name": _with_special_chars(raw_config, vehicle.get("make_model", "")),
            "driver": "" if vehicle_omit.get("driver", False) else _with_special_chars(raw_config, vehicle.get("driver", "")),
            "fuel": _with_special_chars(raw_config, vehicle.get("fuel", "Diesel")),
            "unit": vehicle.get("unit", "L"),
            "quantity": quantity,
            "unit_price": unit_price,
            "currency": company.get("currency", "GBP (£)"),
        })
    return records


def _invoice_non_fuel_lines(raw_config: dict, rng: random.Random, currency: str) -> list[dict]:
    """Non-fuel line items (must be ignored for emissions extraction)."""
    del currency
    lines: list[dict] = []
    if rng.random() < 0.7:
        fee = _q2(Decimal(str(round(rng.uniform(3.0, 18.0), 2))))
        label = rng.choice([_tr(raw_config, "service_fee"), _tr(raw_config, "delivery_charge")])
        lines.append({
            "kind": "fee",
            "description": label,
            "quantity": Decimal("1"),
            "unit": _tr(raw_config, "each"),
            "unit_price": fee,
            "amount": fee,
        })
    return lines


def _cross_scope_invoice_line(raw_config: dict, rng: random.Random) -> dict:
    template = rng.choice(_CROSS_SCOPE_STATIONARY_LINES)
    quantity = _q2(Decimal(rng.randrange(200, 1200, 25)))
    unit_price = _q2(Decimal(str(round(rng.uniform(0.75, 1.15), 2))))
    return {
        "kind": "fuel",
        "category": CATEGORY_STATIONARY,
        "description": f"{template['fuel']} ({template['equipment']})",
        "fuel": template["fuel"],
        "equipment": template["equipment"],
        "quantity": quantity,
        "unit": "L",
        "unit_price": unit_price,
        "amount": _q2(quantity * unit_price),
    }


def _build_invoice_records(raw_config: dict) -> list[dict]:
    fp = _financial_period(raw_config)
    seed = int(raw_config.get("random_seed", 42))
    base_records = _build_invoice_base_records(raw_config)
    if not base_records:
        return []

    count = _multi_document_count(raw_config)
    days_in_period = max((fp["end_date"] - fp["start_date"]).days, 0)
    vat_rate = _vat_rate(raw_config)
    cross_scope = _cross_scope_enabled(raw_config)
    records: list[dict] = []

    for document_index in range(count):
        template = base_records[document_index % len(base_records)]
        record = dict(template)
        rng = random.Random(f"{seed}:mobile_invoice_doc:{document_index}")

        if document_index >= len(base_records):
            record["quantity"] = _q2(Decimal(rng.randrange(30, 90)))
            record["unit_price"] = _q2(Decimal(str(round(rng.uniform(1.35, 1.75), 2))))

        fuel_amount = _q2(record["quantity"] * record["unit_price"])
        lines: list[dict] = [{
            "kind": "fuel",
            "category": CATEGORY_MOBILE,
            "description": record["fuel"],
            "fuel": record["fuel"],
            "vehicle_reg": record["vehicle_reg"],
            "quantity": record["quantity"],
            "unit": record["unit"],
            "unit_price": record["unit_price"],
            "amount": fuel_amount,
        }]
        lines.extend(_invoice_non_fuel_lines(raw_config, rng, record["currency"]))
        if cross_scope:
            lines.append(_cross_scope_invoice_line(raw_config, rng))

        subtotal = _q2(sum((line["amount"] for line in lines), Decimal("0")))
        vat = _q2(subtotal * vat_rate / Decimal("100"))
        invoiced_date = fp["start_date"] + timedelta(days=rng.randint(0, days_in_period))
        record.update({
            "lines": lines,
            "fuel_cost": fuel_amount,
            "subtotal": subtotal,
            "vat_rate": vat_rate,
            "vat": vat,
            "total": _q2(subtotal + vat),
            "invoiced_date": invoiced_date,
            "due_date": invoiced_date + timedelta(days=rng.randint(14, 30)),
            "invoice_no": f"{record.get('supplier_code', 'INV')}-{invoiced_date.strftime('%Y%m')}-{document_index + 1:04d}",
        })
        records.append(record)

    return records


def _build_fuel_card_statements(raw_config: dict) -> list[dict]:
    fp = _financial_period(raw_config)
    seed = int(raw_config.get("random_seed", 42))
    days_in_period = max((fp["end_date"] - fp["start_date"]).days, 0)
    line_count = _multi_document_count(raw_config)
    vat_rate = _vat_rate(raw_config)
    cross_scope = _cross_scope_enabled(raw_config)
    statements: list[dict] = []

    for company_index, company in enumerate(raw_config.get("companies", []), start=1):
        vehicles = company.get("vehicles", [])
        if not vehicles:
            continue

        odometers: dict[str, float] = {}
        transactions: list[dict] = []
        for line_index in range(line_count):
            rng = random.Random(f"{seed}:mobile_fuel_card:{company_index}:{line_index}")
            vehicle = vehicles[line_index % len(vehicles)]
            vehicle_omit = vehicle.get("_omit", {})
            registration = str(vehicle.get("registration", ""))
            is_cross_scope = cross_scope and line_index % 4 == 3
            transaction_date = fp["start_date"] + timedelta(days=rng.randint(0, days_in_period))
            transaction_time = time(hour=rng.randint(6, 21), minute=rng.randint(0, 59))

            if is_cross_scope:
                template = rng.choice(_CROSS_SCOPE_STATIONARY_LINES)
                quantity = _q2(Decimal(rng.randrange(60, 400, 10)))
                unit_price = _q2(Decimal(str(round(rng.uniform(0.75, 1.15), 2))))
                transactions.append({
                    "category": CATEGORY_STATIONARY,
                    "date": transaction_date,
                    "time": transaction_time,
                    "card_number": company.get("card_number") or vehicle.get("card_number", ""),
                    "vehicle_reg": "",
                    "equipment": template["equipment"],
                    "driver": "",
                    "merchant": _with_special_chars(raw_config, company.get("merchant") or "Depot Fuel Point"),
                    "site": _with_special_chars(raw_config, vehicle.get("site", "")),
                    "fuel": template["fuel"],
                    "quantity": quantity,
                    "unit": "L",
                    "unit_price": unit_price,
                    "total": _q2(quantity * unit_price),
                    "receipt_no": _receipt_number(rng),
                    "odometer": "",
                })
                continue

            base_quantity = _parse_decimal(vehicle.get("quantity"), "45")
            quantity = _q2(base_quantity * Decimal(str(round(rng.uniform(0.75, 1.2), 3))))
            unit_price = _q2(_parse_decimal(vehicle.get("unit_price"), "1.52") * Decimal(str(round(rng.uniform(0.96, 1.06), 3))))
            distance_step = rng.uniform(250.0, 900.0)
            odometers[registration] = odometers.get(
                registration,
                float(_parse_decimal(vehicle.get("odometer_start"), "42000")),
            ) + distance_step

            transactions.append({
                "category": CATEGORY_MOBILE,
                "date": transaction_date,
                "time": transaction_time,
                "card_number": vehicle.get("card_number") or company.get("card_number", ""),
                "vehicle_reg": _with_special_chars(raw_config, registration),
                "equipment": "",
                "driver": "" if vehicle_omit.get("driver", False) else _with_special_chars(raw_config, vehicle.get("driver", "")),
                "merchant": _with_special_chars(raw_config, rng.choice(company.get("merchants") or [company.get("merchant") or "Motorway Services"])),
                "site": _with_special_chars(raw_config, vehicle.get("site", "")),
                "fuel": _with_special_chars(raw_config, vehicle.get("fuel", "Diesel")),
                "quantity": quantity,
                "unit": vehicle.get("unit", "L"),
                "unit_price": unit_price,
                "total": _q2(quantity * unit_price),
                "receipt_no": _receipt_number(rng),
                "odometer": round(odometers[registration]),
            })

        transactions.sort(key=lambda row: (row["date"], row["time"], row["card_number"]))
        statement_rng = random.Random(f"{seed}:mobile_statement:{company_index}")
        net_total = _q2(sum((row["total"] for row in transactions), Decimal("0")))
        vat_amount = _q2(net_total * vat_rate / Decimal("100"))
        statements.append({
            "company": _with_special_chars(raw_config, company.get("label", "")),
            "account_name": _with_special_chars(raw_config, company.get("customer") or company.get("label", "")),
            "provider": _with_special_chars(raw_config, company.get("supplier", "")),
            "account_number": _account_number(company, seed, company_index),
            "statement_no": f"ST-{fp['end_date'].strftime('%Y%m')}-{statement_rng.randint(1000, 9999)}",
            "issue_date": fp["end_date"] + timedelta(days=statement_rng.randint(1, 6)),
            "currency": company.get("currency", "GBP (£)"),
            "period_label": fp["label"],
            "period_start": fp["start_date"],
            "period_end": fp["end_date"],
            "transactions": transactions,
            "net_total": net_total,
            "vat_rate": vat_rate,
            "vat_amount": vat_amount,
            "gross_total": _q2(net_total + vat_amount),
        })

    return statements


def _fuel_card_summary_rows(raw_config: dict, statement: dict) -> list[tuple[str, Any, bool]]:
    return [
        (_tr(raw_config, "net_amount"), statement["net_total"], False),
        (f"{_tr(raw_config, 'vat')} ({statement['vat_rate']}%)", statement["vat_amount"], False),
        (_tr(raw_config, "gross_amount"), statement["gross_total"], True),
    ]


def _monthly_distance(vehicle: dict, rng: random.Random) -> float:
    base = float(_parse_decimal(vehicle.get("monthly_distance_km"), "2400"))
    return base * rng.uniform(0.85, 1.15)


def _build_telematics_fuel_rows(raw_config: dict) -> list[dict]:
    fp = _financial_period(raw_config)
    seed = int(raw_config.get("random_seed", 42))
    months = _months_in_range(fp["start_date"], fp["end_date"])
    distance_unit = _distance_unit(raw_config)
    km_factor = 0.621371 if distance_unit == "mi" else 1.0
    rows: list[dict] = []

    for company_index, vehicle_index, company, vehicle in _iter_company_vehicles(raw_config):
        vehicle_omit = vehicle.get("_omit", {})
        efficiency = float(_parse_decimal(vehicle.get("efficiency_l_per_100km"), "9.5"))
        for year, month in months:
            rng = random.Random(f"{seed}:telematics_fuel:{company_index}:{vehicle_index}:{year}-{month:02d}")
            month_start = max(date(year, month, 1), fp["start_date"])
            next_month = date(year + 1, 1, 1) if month == 12 else date(year, month + 1, 1)
            month_end = min(next_month - timedelta(days=1), fp["end_date"])
            distance_km = _monthly_distance(vehicle, rng) * ((month_end - month_start).days + 1) / 30.0
            fuel_used = distance_km * efficiency / 100.0 * rng.uniform(0.95, 1.08)
            idle_fuel = fuel_used * rng.uniform(0.02, 0.09)
            engine_hours = distance_km / rng.uniform(38.0, 55.0)
            distance = distance_km * km_factor

            rows.append({
                "company": _with_special_chars(raw_config, company.get("label", "")),
                "period_start": month_start,
                "period_end": month_end,
                "vehicle_reg": _with_special_chars(raw_config, vehicle.get("registration", "")),
                "vehicle_name": _with_special_chars(raw_config, vehicle.get("make_model", "")),
                "vehicle_type": vehicle.get("vehicle_type", ""),
                "fuel": "" if vehicle_omit.get("fuel", False) else _with_special_chars(raw_config, vehicle.get("fuel", "Diesel")),
                "distance": round(distance, 1),
                "distance_unit": distance_unit,
                "fuel_used": round(fuel_used, 2),
                "unit": vehicle.get("unit", "L"),
                "idle_fuel": round(idle_fuel, 2),
                "engine_hours": round(engine_hours, 1),
                "avg_consumption": round(fuel_used / max(distance_km, 1.0) * 100.0, 2),
            })

    return rows


def _trips_per_vehicle(raw_config: dict) -> int:
    try:
        count = int(_document_option(raw_config, "trips_per_vehicle", 10) or 10)
    except (TypeError, ValueError):
        count = 10
    return max(count, 1)


def _build_trip_rows(raw_config: dict) -> list[dict]:
    fp = _financial_period(raw_config)
    seed = int(raw_config.get("random_seed", 42))
    days_in_period = max((fp["end_date"] - fp["start_date"]).days, 0)
    distance_unit = _distance_unit(raw_config)
    km_factor = 0.621371 if distance_unit == "mi" else 1.0
    trips_per_vehicle = _trips_per_vehicle(raw_config)
    rows: list[dict] = []

    for company_index, vehicle_index, company, vehicle in _iter_company_vehicles(raw_config):
        vehicle_omit = vehicle.get("_omit", {})
        for trip_index in range(trips_per_vehicle):
            rng = random.Random(f"{seed}:telematics_trip:{company_index}:{vehicle_index}:{trip_index}")
            trip_date = fp["start_date"] + timedelta(days=rng.randint(0, days_in_period))
            start_dt = datetime.combine(trip_date, time(hour=rng.randint(5, 18), minute=rng.randint(0, 59)))
            distance_km = rng.uniform(8.0, 320.0)
            avg_speed_kmh = rng.uniform(28.0, 88.0)
            duration_minutes = max(int(distance_km / avg_speed_kmh * 60), 4)
            end_dt = start_dt + timedelta(minutes=duration_minutes)
            start_location, end_location = rng.sample(_TRIP_LOCATIONS, 2)

            rows.append({
                "company": _with_special_chars(raw_config, company.get("label", "")),
                "trip_id": f"TRP-{trip_date.strftime('%Y%m%d')}-{rng.randint(1000, 9999)}",
                "vehicle_reg": _with_special_chars(raw_config, vehicle.get("registration", "")),
                "driver": "" if vehicle_omit.get("driver", False) else _with_special_chars(raw_config, vehicle.get("driver", "")),
                "trip_start": start_dt,
                "trip_end": end_dt,
                "start_location": start_location,
                "end_location": end_location,
                "distance": round(distance_km * km_factor, 1),
                "distance_unit": distance_unit,
                "duration": f"{duration_minutes // 60:02d}:{duration_minutes % 60:02d}",
                "avg_speed": round(avg_speed_kmh * km_factor, 1),
                "fuel": "" if vehicle_omit.get("fuel", False) else _with_special_chars(raw_config, vehicle.get("fuel", "Diesel")),
            })

    rows.sort(key=lambda row: (row["vehicle_reg"], row["trip_start"]))
    return rows


def _build_odometer_rows(raw_config: dict) -> list[dict]:
    fp = _financial_period(raw_config)
    seed = int(raw_config.get("random_seed", 42))
    days_in_period = max((fp["end_date"] - fp["start_date"]).days + 1, 1)
    distance_unit = _distance_unit(raw_config)
    km_factor = 0.621371 if distance_unit == "mi" else 1.0
    rows: list[dict] = []

    for company_index, vehicle_index, company, vehicle in _iter_company_vehicles(raw_config):
        vehicle_omit = vehicle.get("_omit", {})
        rng = random.Random(f"{seed}:telematics_odometer:{company_index}:{vehicle_index}")
        odometer_start_km = float(_parse_decimal(vehicle.get("odometer_start"), "42000"))
        distance_km = _monthly_distance(vehicle, rng) * days_in_period / 30.0
        odometer_end_km = odometer_start_km + distance_km

        rows.append({
            "company": _with_special_chars(raw_config, company.get("label", "")),
            "period_start": fp["start_date"],
            "period_end": fp["end_date"],
            "vehicle_reg": _with_special_chars(raw_config, vehicle.get("registration", "")),
            "vehicle_name": _with_special_chars(raw_config, vehicle.get("make_model", "")),
            "fuel": "" if vehicle_omit.get("fuel", False) else _with_special_chars(raw_config, vehicle.get("fuel", "Diesel")),
            "odometer_start": round(odometer_start_km * km_factor),
            "odometer_end": round(odometer_end_km * km_factor),
            "distance": round(distance_km * km_factor, 1),
            "distance_unit": distance_unit,
        })

    return rows


# ── invalid-data wrappers ────────────────────────────────────────────────────


def _corrupted_invoice_records(raw_config: dict) -> list[dict]:
    return corrupt_records(
        _build_invoice_records(raw_config),
        "mobile_fuel_invoice",
        _MOBILE_INVOICE_FIELD_TYPES,
        get_bad_data_config(raw_config),
    )


def _corrupted_fuel_card_statements(raw_config: dict) -> list[dict]:
    cfg = get_bad_data_config(raw_config)
    statements = []
    for statement_index, statement in enumerate(_build_fuel_card_statements(raw_config)):
        new_statement = dict(statement)
        new_statement["transactions"] = corrupt_records(
            statement["transactions"],
            f"mobile_fuel_card:s{statement_index}",
            _MOBILE_FUEL_CARD_FIELD_TYPES,
            cfg,
        )
        statements.append(new_statement)
    return statements


def _corrupted_telematics_fuel_rows(raw_config: dict) -> list[dict]:
    return corrupt_records(
        _build_telematics_fuel_rows(raw_config),
        "telematics_fuel",
        _TELEMATICS_FUEL_FIELD_TYPES,
        get_bad_data_config(raw_config),
    )


def _corrupted_trip_rows(raw_config: dict) -> list[dict]:
    return corrupt_records(
        _build_trip_rows(raw_config),
        "telematics_trips",
        _TELEMATICS_TRIP_FIELD_TYPES,
        get_bad_data_config(raw_config),
    )


def _corrupted_odometer_rows(raw_config: dict) -> list[dict]:
    return corrupt_records(
        _build_odometer_rows(raw_config),
        "telematics_odometer",
        _TELEMATICS_ODOMETER_FIELD_TYPES,
        get_bad_data_config(raw_config),
    )


# ── ground truth ─────────────────────────────────────────────────────────────


def _iso(value) -> str:
    return value.isoformat() if hasattr(value, "isoformat") else str(value)


def _number(value):
    if isinstance(value, str):
        return value
    if isinstance(value, Decimal):
        return float(value)
    return value


def _currency_code_value(value) -> str:
    if isinstance(value, str) and " " in value:
        return value.split()[0]
    return value


def _ground_truth_entries(raw_config: dict) -> list[dict]:
    document_type = _document_type(raw_config)

    if document_type == "fuel_invoice":
        entries = []
        for record in _corrupted_invoice_records(raw_config):
            for line in record.get("lines", []):
                if line.get("kind") != "fuel":
                    continue
                entry = {
                    "Scope": SCOPE_LABEL,
                    "Category": line.get("category", CATEGORY_MOBILE),
                    "Company": record["company"],
                    "Supplier": record["supplier"],
                    "Invoice No": record["invoice_no"],
                    "Account Number": record["account_number"],
                    "Invoiced date": _iso(record["invoiced_date"]),
                    "Due date": _iso(record["due_date"]),
                    "Fuel": line.get("fuel", ""),
                    "Quantity": _number(line.get("quantity")),
                    "Unit": line.get("unit", ""),
                    "Cost": _number(line.get("amount")),
                    "Currency": _currency_code_value(record["currency"]),
                }
                if line.get("category", CATEGORY_MOBILE) == CATEGORY_MOBILE:
                    entry["Vehicle"] = record["vehicle_reg"]
                    entry["Site"] = record["site"]
                else:
                    entry["Equipment"] = line.get("equipment", "")
                entries.append(entry)
        return entries

    if document_type == "fuel_card_statement":
        entries = []
        for statement in _corrupted_fuel_card_statements(raw_config):
            for transaction in statement["transactions"]:
                entries.append({
                    "Scope": SCOPE_LABEL,
                    "Category": transaction.get("category", CATEGORY_MOBILE),
                    "Company": statement["company"],
                    "Statement No": statement["statement_no"],
                    "Account Number": statement["account_number"],
                    "Card Number": transaction["card_number"],
                    "Vehicle": transaction["vehicle_reg"],
                    "Equipment": transaction.get("equipment", ""),
                    "Date": _iso(transaction["date"]),
                    "Time": transaction["time"].strftime("%H:%M") if hasattr(transaction["time"], "strftime") else str(transaction["time"]),
                    "Merchant": transaction["merchant"],
                    "Receipt No": transaction["receipt_no"],
                    "Fuel": transaction["fuel"],
                    "Quantity": _number(transaction["quantity"]),
                    "Unit": transaction["unit"],
                    "Cost": _number(transaction["total"]),
                    "Currency": _currency_code_value(statement["currency"]),
                    "Odometer": transaction["odometer"],
                })
        return entries

    if document_type == "telematics_fuel":
        return [
            {
                "Scope": SCOPE_LABEL,
                "Category": CATEGORY_MOBILE,
                "Company": row["company"],
                "Vehicle": row["vehicle_reg"],
                "Period start": _iso(row["period_start"]),
                "Period end": _iso(row["period_end"]),
                "Fuel": row["fuel"],
                "Quantity": _number(row["fuel_used"]),
                "Unit": row["unit"],
                "Distance": _number(row["distance"]),
                "Distance unit": row["distance_unit"],
                "Idle fuel": _number(row["idle_fuel"]),
            }
            for row in _corrupted_telematics_fuel_rows(raw_config)
        ]

    if document_type == "telematics_trips":
        return [
            {
                "Scope": SCOPE_LABEL,
                "Category": CATEGORY_MOBILE,
                "Company": row["company"],
                "Vehicle": row["vehicle_reg"],
                "Trip ID": row["trip_id"],
                "Trip start": _iso(row["trip_start"]),
                "Trip end": _iso(row["trip_end"]),
                "Start location": row["start_location"],
                "End location": row["end_location"],
                "Distance": _number(row["distance"]),
                "Distance unit": row["distance_unit"],
                "Fuel": row["fuel"],
            }
            for row in _corrupted_trip_rows(raw_config)
        ]

    return [
        {
            "Scope": SCOPE_LABEL,
            "Category": CATEGORY_MOBILE,
            "Company": row["company"],
            "Vehicle": row["vehicle_reg"],
            "Period start": _iso(row["period_start"]),
            "Period end": _iso(row["period_end"]),
            "Odometer start": _number(row["odometer_start"]),
            "Odometer end": _number(row["odometer_end"]),
            "Distance": _number(row["distance"]),
            "Distance unit": row["distance_unit"],
            "Fuel": row["fuel"],
        }
        for row in _corrupted_odometer_rows(raw_config)
    ]


def generate_ground_truth_json(raw_config: dict) -> bytes:
    return json.dumps(_ground_truth_entries(raw_config), indent=2).encode("utf-8")


# ── fuel invoice renderers ───────────────────────────────────────────────────


def _invoice_filename(record: dict, index: int, ext: str) -> str:
    return f"{_slugify_filename(record.get('invoice_no'), f'invoice_{index + 1:04d}')}.{ext}"


def _is_multi_document(raw_config: dict) -> bool:
    return _multi_document_count(raw_config) > 1


def generate_fuel_invoice_pdf(raw_config: dict) -> bytes:
    records = _corrupted_invoice_records(raw_config)
    if _is_multi_document(raw_config):
        return _zip_documents(raw_config, records, _render_invoice_pdf, _invoice_filename, "pdf")
    return _render_invoice_pdf(raw_config, records)


def _invoice_meta_lines(raw_config: dict, record: dict, distractor_plan) -> list[str]:
    return [
        f"{_tr(raw_config, 'invoice_no')}: {record['invoice_no']}",
        f"{_tr(raw_config, 'account_no')}: {record['account_number']}",
        f"{_tr(raw_config, 'invoiced_date')}: {_fmt_date(record['invoiced_date'])}",
        f"{_tr(raw_config, 'due_date')}: {_fmt_date(record['due_date'])}",
        f"{_tr(raw_config, 'currency')}: {record['currency']}",
        *_document_lines(distractor_plan, placement="meta"),
    ]


def _invoice_vehicle_lines(raw_config: dict, record: dict, distractor_plan) -> list[str]:
    lines = [
        f"{_tr(raw_config, 'registration')}: {record['vehicle_reg']}",
        record["vehicle_name"],
    ]
    if record.get("driver"):
        lines.append(f"{_tr(raw_config, 'driver')}: {record['driver']}")
    if record.get("site"):
        lines.append(f"{_tr(raw_config, 'site')}: {record['site']}")
    lines.extend(_document_lines(distractor_plan, placement="summary"))
    return [line for line in lines if line]


def _invoice_line_rows(raw_config: dict, record: dict) -> list[tuple[str, str, str, str, str]]:
    symbol = _currency_symbol(record["currency"])
    rows = []
    for line in record.get("lines", []):
        rows.append((
            str(line.get("description", "")),
            _fmt_num(line.get("quantity")),
            str(line.get("unit", "")),
            f"{symbol}{_fmt_money(line.get('unit_price'))}",
            f"{symbol}{_fmt_money(line.get('amount'))}",
        ))
    return rows


def _invoice_summary_lines(raw_config: dict, record: dict) -> list[tuple[str, Any, bool]]:
    return [
        (_tr(raw_config, "subtotal"), record["subtotal"], False),
        (f"{_tr(raw_config, 'vat')} ({record['vat_rate']}%)", record["vat"], False),
        (_tr(raw_config, "total"), record["total"], True),
    ]


def _render_invoice_pdf(raw_config: dict, records: list[dict]) -> bytes:
    distractor_plan = _distractor_plan(raw_config, "PDF")
    layout_plan = _layout_plan(raw_config, "PDF")

    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    c.setTitle(raw_config.get("document", {}).get("title", _tr(raw_config, "fuel_invoice_title")))
    c.setSubject(raw_config.get("document", {}).get("subject", "Scope 1 mobile combustion"))

    default_sections = ["addresses", "meta", "line_items", "totals"]
    section_order = [
        section
        for section in (layout_plan.get("section_order") or default_sections)
        if section in set(default_sections)
    ] if layout_plan.get("enabled") else default_sections

    for index, record in enumerate(records):
        if index > 0:
            c.showPage()

        accent = colors.HexColor("#245C4F")
        c.setFillColor(accent)
        c.rect(36, PAGE_H - 72, PAGE_W - 72, 28, fill=1, stroke=0)
        c.setFillColor(colors.white)
        c.setFont("Helvetica-Bold", 14)
        c.drawString(48, PAGE_H - 62, record["supplier"])

        current_top = PAGE_H - 108
        currency_symbol = _currency_symbol(record["currency"])

        for section_name in section_order:
            if section_name == "addresses":
                c.setFillColor(colors.black)
                c.setFont("Helvetica", 10)
                current_top = _draw_multiline(c, 48, current_top, record["supplier_address"]) - 10
                c.setFont("Helvetica-Bold", 11)
                c.drawString(48, current_top, _tr(raw_config, "bill_to"))
                c.setFont("Helvetica", 10)
                current_top = _draw_multiline(c, 48, current_top - 16, [record["customer"]]) - 18
            elif section_name == "meta":
                c.setFillColor(colors.black)
                c.setFont("Helvetica-Bold", 11)
                c.drawString(320, PAGE_H - 108, _tr(raw_config, "invoice_details"))
                c.setFont("Helvetica", 10)
                meta_bottom = _draw_multiline(c, 320, PAGE_H - 126, _invoice_meta_lines(raw_config, record, distractor_plan))
                c.setFont("Helvetica-Bold", 11)
                c.drawString(320, meta_bottom - 8, _tr(raw_config, "vehicle_details"))
                c.setFont("Helvetica", 10)
                vehicle_bottom = _draw_multiline(c, 320, meta_bottom - 24, _invoice_vehicle_lines(raw_config, record, distractor_plan))
                current_top = min(current_top, vehicle_bottom - 12)
            elif section_name == "line_items":
                table_top = min(current_top, PAGE_H - 330)
                table_x = 48
                table_widths = [210, 68, 58, 84, 84]
                headers = [
                    _tr(raw_config, "description"),
                    _tr(raw_config, "quantity"),
                    _tr(raw_config, "unit"),
                    _tr(raw_config, "unit_price"),
                    _tr(raw_config, "amount"),
                ]
                c.setFillColor(accent)
                c.rect(table_x, table_top, sum(table_widths), 22, fill=1, stroke=0)
                c.setFillColor(colors.white)
                x_cursor = table_x + 6
                for header, width in zip(headers, table_widths):
                    _draw_fitted_string(c, x_cursor, table_top + 7, header, width - 8, "Helvetica-Bold", 9)
                    x_cursor += width

                y_row = table_top - 24
                c.setFillColor(colors.black)
                for row in _invoice_line_rows(raw_config, record):
                    c.rect(table_x, y_row, sum(table_widths), 20, fill=0, stroke=1)
                    x_cursor = table_x + 6
                    for value, width in zip(row, table_widths):
                        _draw_fitted_string(c, x_cursor, y_row + 6, str(value), width - 8, "Helvetica", 9)
                        x_cursor += width
                    y_row -= 20
                current_top = y_row - 14
            elif section_name == "totals":
                summary_y = current_top
                for label, value, is_total in _invoice_summary_lines(raw_config, record):
                    c.setFont("Helvetica-Bold" if is_total else "Helvetica", 10)
                    c.drawRightString(PAGE_W - 180, summary_y, label)
                    c.drawRightString(PAGE_W - 48, summary_y, f"{currency_symbol}{_fmt_money(value)}")
                    summary_y -= 18
                current_top = summary_y - 10

        c.setFont("Helvetica", 8)
        c.setFillColor(colors.grey)
        c.drawString(48, 42, _tr(raw_config, "fuel_invoice_footer"))

    c.save()
    return buffer.getvalue()


def generate_fuel_invoice_docx(raw_config: dict) -> bytes:
    records = _corrupted_invoice_records(raw_config)
    if _is_multi_document(raw_config):
        return _zip_documents(raw_config, records, _render_invoice_docx, _invoice_filename, "docx")
    return _render_invoice_docx(raw_config, records)


def _render_invoice_docx(raw_config: dict, records: list[dict]) -> bytes:
    distractor_plan = _distractor_plan(raw_config, "DOCX")
    layout_plan = _layout_plan(raw_config, "DOCX")

    document = Document()
    _style_docx_document(document)
    document.core_properties.title = raw_config.get("document", {}).get("title", _tr(raw_config, "fuel_invoice_title"))
    document.core_properties.subject = raw_config.get("document", {}).get("subject", "Scope 1 mobile combustion")

    default_sections = ["addresses", "meta", "line_items", "totals", "footer"]
    section_order = (
        layout_plan.get("section_order") or default_sections
    ) if layout_plan.get("enabled") else default_sections

    for index, record in enumerate(records):
        banner = document.add_table(rows=1, cols=2)
        banner.style = "Table Grid"
        banner.autofit = False
        banner.columns[0].width = Inches(4.8)
        banner.columns[1].width = Inches(2.0)
        _shade_docx_cell(banner.cell(0, 0), "245C4F")
        _shade_docx_cell(banner.cell(0, 1), "245C4F")
        _set_docx_cell_text(banner.cell(0, 0), record["supplier"], bold=True, color="FFFFFF", size=13)
        _set_docx_cell_text(banner.cell(0, 1), _tr(raw_config, "fuel_invoice_title"), bold=True, color="FFFFFF", size=12)
        banner.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        currency_symbol = _currency_symbol(record["currency"])

        def render_addresses() -> None:
            addresses = document.add_table(rows=2, cols=2)
            addresses.style = "Table Grid"
            for cell, heading in zip(addresses.rows[0].cells, [record["supplier"], _tr(raw_config, "bill_to")]):
                _shade_docx_cell(cell, "E4EFEA")
                _set_docx_cell_text(cell, heading, bold=True)
            _set_docx_cell_text(addresses.cell(1, 0), "\n".join(record["supplier_address"]))
            _set_docx_cell_text(addresses.cell(1, 1), record["customer"])
            document.add_paragraph()

        def render_meta() -> None:
            meta = document.add_table(rows=2, cols=2)
            meta.style = "Table Grid"
            for cell, heading in zip(meta.rows[0].cells, [_tr(raw_config, "invoice_details"), _tr(raw_config, "vehicle_details")]):
                _shade_docx_cell(cell, "E4EFEA")
                _set_docx_cell_text(cell, heading, bold=True)
            _set_docx_cell_text(meta.cell(1, 0), "\n".join(_invoice_meta_lines(raw_config, record, distractor_plan)))
            _set_docx_cell_text(meta.cell(1, 1), "\n".join(_invoice_vehicle_lines(raw_config, record, distractor_plan)))
            document.add_paragraph()

        def render_line_items() -> None:
            table = document.add_table(rows=1, cols=5)
            table.style = "Table Grid"
            headers = [
                _tr(raw_config, "description"),
                _tr(raw_config, "quantity"),
                _tr(raw_config, "unit"),
                _tr(raw_config, "unit_price"),
                _tr(raw_config, "amount"),
            ]
            for cell, header in zip(table.rows[0].cells, headers):
                _shade_docx_cell(cell, "F5F8FB")
                _set_docx_cell_text(cell, header, bold=True)
            for row_values in _invoice_line_rows(raw_config, record):
                row = table.add_row().cells
                for cell, value in zip(row, row_values):
                    _set_docx_cell_text(cell, str(value))
            document.add_paragraph()

        def render_totals() -> None:
            totals = document.add_table(rows=3, cols=2)
            totals.style = "Table Grid"
            for row_idx, (label, value, is_total) in enumerate(_invoice_summary_lines(raw_config, record)):
                _shade_docx_cell(totals.cell(row_idx, 0), "F5F8FB")
                _set_docx_cell_text(totals.cell(row_idx, 0), label, bold=True)
                _set_docx_cell_text(totals.cell(row_idx, 1), f"{currency_symbol}{_fmt_money(value)}", bold=is_total)

        def render_footer() -> None:
            footer = document.add_paragraph()
            footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = footer.add_run(_tr(raw_config, "fuel_invoice_footer"))
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor.from_string("6E7A86")

        renderers = {
            "addresses": render_addresses,
            "meta": render_meta,
            "line_items": render_line_items,
            "totals": render_totals,
            "footer": render_footer,
        }
        for section_name in section_order:
            renderer = renderers.get(section_name)
            if renderer is not None:
                renderer()

        if index < len(records) - 1:
            document.add_page_break()

    output = BytesIO()
    document.save(output)
    return output.getvalue()


# ── fuel card statement renderers ────────────────────────────────────────────

_FUEL_CARD_HEADER_KEYS: dict[str, str] = {
    "date": "date",
    "time": "time",
    "card_no": "card_no",
    "vehicle_reg": "vehicle_reg",
    "driver": "driver",
    "merchant": "merchant",
    "site": "location",
    "receipt_no": "receipt_no",
    "odometer": "odometer",
    "product": "fuel_type",
    "qty": "qty",
    "unit": "unit",
    "unit_price": "unit_price",
    "total": "total",
    "currency": "currency",
}


def _fuel_card_row_map(statement: dict, transaction: dict) -> dict[str, Any]:
    reference = transaction["vehicle_reg"] or transaction.get("equipment", "")
    return {
        "date": transaction["date"].strftime("%d-%m-%y") if hasattr(transaction["date"], "strftime") else transaction["date"],
        "time": transaction["time"].strftime("%H:%M") if hasattr(transaction["time"], "strftime") else transaction["time"],
        "card_no": transaction["card_number"],
        "vehicle_reg": reference,
        "driver": transaction["driver"],
        "merchant": transaction["merchant"],
        "site": transaction["site"],
        "receipt_no": transaction["receipt_no"],
        "odometer": transaction["odometer"],
        "product": transaction["fuel"],
        "qty": _number(transaction["quantity"]),
        "unit": transaction["unit"],
        "unit_price": _number(transaction["unit_price"]),
        "total": _number(transaction["total"]),
        "currency": _currency_code_value(statement["currency"]),
    }


def _fuel_card_meta_lines(raw_config: dict, statement: dict, distractor_plan) -> list[str]:
    return [
        f"{_tr(raw_config, 'account_name')}: {statement['account_name']}",
        f"{_tr(raw_config, 'account_no')}: {statement['account_number']}",
        f"{_tr(raw_config, 'provider')}: {statement['provider']}",
        f"{_tr(raw_config, 'statement_no')}: {statement['statement_no']}",
        f"{_tr(raw_config, 'statement_period')}: {_fmt_date(statement['period_start'])} - {_fmt_date(statement['period_end'])}",
        f"{_tr(raw_config, 'issue_date')}: {_fmt_date(statement['issue_date'])}",
        f"{_tr(raw_config, 'currency')}: {statement['currency']}",
        *_document_lines(distractor_plan),
    ]


def generate_fuel_card_statement_pdf(raw_config: dict) -> bytes:
    statements = _corrupted_fuel_card_statements(raw_config)
    distractor_plan = _distractor_plan(raw_config, "PDF")
    layout_plan = _layout_plan(raw_config, "PDF")

    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    c.setTitle(raw_config.get("document", {}).get("title", _tr(raw_config, "fuel_card_title")))
    c.setSubject(raw_config.get("document", {}).get("subject", "Scope 1 mobile combustion fuel card transactions"))

    default_sections = ["summary", "transactions"]
    section_order = [
        section
        for section in (layout_plan.get("section_order") or default_sections)
        if section in set(default_sections)
    ] if layout_plan.get("enabled") else default_sections

    page_size = 16
    for statement_index, statement in enumerate(statements):
        transactions = statement["transactions"]
        currency_symbol = _currency_symbol(statement["currency"])
        for page_start in range(0, max(len(transactions), 1), page_size):
            if statement_index > 0 or page_start > 0:
                c.showPage()

            accent = colors.HexColor("#245C4F")
            c.setFillColor(accent)
            c.rect(36, PAGE_H - 74, PAGE_W - 72, 30, fill=1, stroke=0)
            c.setFillColor(colors.white)
            c.setFont("Helvetica-Bold", 15)
            c.drawString(48, PAGE_H - 63, _tr(raw_config, "fuel_card_title"))

            current_top = PAGE_H - 106
            for section_name in section_order:
                if section_name == "summary":
                    c.setFillColor(colors.black)
                    c.setFont("Helvetica", 9.5)
                    current_top = _draw_multiline(c, 48, current_top, _fuel_card_meta_lines(raw_config, statement, distractor_plan), leading=13) - 14
                elif section_name == "transactions":
                    table_x = 30
                    table_top = current_top - 6
                    column_widths = [40, 22, 78, 52, 76, 56, 38, 40, 28, 16, 38, 40]
                    header_ids = ["date", "time", "card_no", "vehicle_reg", "merchant", "receipt_no", "odometer", "product", "qty", "unit", "unit_price", "total"]
                    compact_headers = {"vehicle_reg": "Vehicle", "receipt_no": "Receipt", "product": "Fuel"}
                    c.setFillColor(accent)
                    c.rect(table_x, table_top, sum(column_widths), 22, fill=1, stroke=0)
                    c.setFillColor(colors.white)
                    cursor = table_x + 3
                    for field_id, width in zip(header_ids, column_widths):
                        header = compact_headers.get(field_id) or _tr(raw_config, _FUEL_CARD_HEADER_KEYS.get(field_id, field_id))
                        _draw_fitted_string(c, cursor, table_top + 7, header, width - 4, "Helvetica-Bold", 6.8)
                        cursor += width

                    row_y = table_top - 18
                    c.setFillColor(colors.black)
                    for transaction in transactions[page_start:page_start + page_size]:
                        row_map = _fuel_card_row_map(statement, transaction)
                        c.rect(table_x, row_y, sum(column_widths), 18, fill=0, stroke=1)
                        cursor = table_x + 3
                        for field_id, width in zip(header_ids, column_widths):
                            value = row_map[field_id]
                            if field_id in {"unit_price", "total"} and not isinstance(value, str):
                                value = f"{currency_symbol}{_fmt_money(value)}"
                            elif field_id == "qty" and not isinstance(value, str):
                                value = _fmt_num(value)
                            _draw_fitted_string(c, cursor, row_y + 5, str(value), width - 4, "Helvetica", 6.8)
                            cursor += width
                        row_y -= 18

                    if page_start + page_size >= len(transactions):
                        summary_y = row_y - 18
                        for label, amount, is_total in _fuel_card_summary_rows(raw_config, statement):
                            c.setFont("Helvetica-Bold" if is_total else "Helvetica", 10)
                            c.drawRightString(PAGE_W - 180, summary_y, label)
                            c.drawRightString(PAGE_W - 48, summary_y, f"{currency_symbol}{_fmt_money(amount)}")
                            summary_y -= 16
                        row_y = summary_y - 12
                    current_top = row_y - 10

            c.setFont("Helvetica", 8)
            c.setFillColor(colors.grey)
            c.drawString(48, 42, _tr(raw_config, "fuel_card_footer"))

    c.save()
    return buffer.getvalue()


def generate_fuel_card_statement_docx(raw_config: dict) -> bytes:
    statements = _corrupted_fuel_card_statements(raw_config)
    distractor_plan = _distractor_plan(raw_config, "DOCX")
    layout_plan = _layout_plan(raw_config, "DOCX")

    document = Document()
    _style_docx_document(document)
    document.core_properties.title = raw_config.get("document", {}).get("title", _tr(raw_config, "fuel_card_title"))
    document.core_properties.subject = raw_config.get("document", {}).get("subject", "Scope 1 mobile combustion fuel card transactions")

    default_sections = ["summary", "transactions", "footer"]
    section_order = (
        layout_plan.get("section_order") or default_sections
    ) if layout_plan.get("enabled") else default_sections

    header_ids = ["date", "time", "card_no", "vehicle_reg", "merchant", "receipt_no", "odometer", "product", "qty", "unit", "unit_price", "total"]

    for statement_index, statement in enumerate(statements):
        banner = document.add_table(rows=1, cols=2)
        banner.style = "Table Grid"
        banner.autofit = False
        banner.columns[0].width = Inches(4.8)
        banner.columns[1].width = Inches(2.0)
        _shade_docx_cell(banner.cell(0, 0), "245C4F")
        _shade_docx_cell(banner.cell(0, 1), "245C4F")
        _set_docx_cell_text(banner.cell(0, 0), statement["account_name"], bold=True, color="FFFFFF", size=13)
        _set_docx_cell_text(banner.cell(0, 1), _tr(raw_config, "fuel_card_title"), bold=True, color="FFFFFF", size=12)
        banner.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        currency_symbol = _currency_symbol(statement["currency"])

        def render_summary() -> None:
            meta = document.add_table(rows=2, cols=1)
            meta.style = "Table Grid"
            _shade_docx_cell(meta.cell(0, 0), "E4EFEA")
            _set_docx_cell_text(meta.cell(0, 0), _tr(raw_config, "statement_details"), bold=True)
            _set_docx_cell_text(meta.cell(1, 0), "\n".join(_fuel_card_meta_lines(raw_config, statement, distractor_plan)))
            document.add_paragraph()

        def render_transactions() -> None:
            table = document.add_table(rows=1, cols=len(header_ids))
            table.style = "Table Grid"
            for cell, field_id in zip(table.rows[0].cells, header_ids):
                _shade_docx_cell(cell, "F5F8FB")
                _set_docx_cell_text(cell, _tr(raw_config, _FUEL_CARD_HEADER_KEYS.get(field_id, field_id)), bold=True)
            for transaction in statement["transactions"]:
                row_map = _fuel_card_row_map(statement, transaction)
                row = table.add_row().cells
                for cell, field_id in zip(row, header_ids):
                    value = row_map[field_id]
                    if field_id in {"unit_price", "total"} and not isinstance(value, str):
                        value = f"{currency_symbol}{_fmt_money(value)}"
                    elif field_id == "qty" and not isinstance(value, str):
                        value = _fmt_num(value)
                    _set_docx_cell_text(cell, str(value))

            document.add_paragraph()
            totals = document.add_table(rows=3, cols=2)
            totals.style = "Table Grid"
            for row_idx, (label, amount, is_total) in enumerate(_fuel_card_summary_rows(raw_config, statement)):
                _shade_docx_cell(totals.cell(row_idx, 0), "F5F8FB")
                _set_docx_cell_text(totals.cell(row_idx, 0), label, bold=True)
                _set_docx_cell_text(totals.cell(row_idx, 1), f"{currency_symbol}{_fmt_money(amount)}", bold=is_total)

        def render_footer() -> None:
            footer = document.add_paragraph()
            footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = footer.add_run(_tr(raw_config, "fuel_card_footer"))
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor.from_string("6E7A86")

        renderers = {
            "summary": render_summary,
            "transactions": render_transactions,
            "footer": render_footer,
        }
        for section_name in section_order:
            renderer = renderers.get(section_name)
            if renderer is not None:
                renderer()

        if statement_index < len(statements) - 1:
            document.add_page_break()

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def generate_fuel_card_statement_xlsx(raw_config: dict) -> bytes:
    statements = _corrupted_fuel_card_statements(raw_config)
    layout_plan = _layout_plan(raw_config, "XLSX")
    distractor_plan = _distractor_plan(raw_config, "XLSX")
    distractor_fields = _distractor_field_map(distractor_plan)
    ordered_ids = _augment_field_ids(_ordered_field_ids(layout_plan, list(_FUEL_CARD_HEADER_KEYS)), distractor_plan)

    workbook = openpyxl.Workbook()
    workbook.remove(workbook.active)

    for index, statement in enumerate(statements, start=1):
        sheet = workbook.create_sheet(title=(statement["company"] or f"Account {index}")[:31])
        sheet["A1"] = _tr(raw_config, "fuel_card_title")
        sheet["A1"].font = Font(size=14, bold=True)
        sheet["A2"] = _tr(raw_config, "account_name")
        sheet["B2"] = statement["account_name"]
        sheet["A3"] = _tr(raw_config, "account_no")
        sheet["B3"] = statement["account_number"]
        sheet["A4"] = _tr(raw_config, "statement_no")
        sheet["B4"] = statement["statement_no"]
        sheet["A5"] = _tr(raw_config, "statement_period")
        sheet["B5"] = f"{statement['period_start'].isoformat()} to {statement['period_end'].isoformat()}"
        sheet["A6"] = _tr(raw_config, "currency")
        sheet["B6"] = statement["currency"]

        header_row = _write_xlsx_preamble(sheet, 7, layout_plan)
        header_fill = PatternFill(fill_type="solid", fgColor="245C4F")
        for column_index, field_id in enumerate(ordered_ids, start=1):
            cell = sheet.cell(
                row=header_row,
                column=column_index,
                value=_header_text(raw_config, layout_plan, field_id, _FUEL_CARD_HEADER_KEYS.get(field_id, field_id), distractor_fields),
            )
            cell.font = Font(color="FFFFFF", bold=True)
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center")

        row_index = header_row + 1
        for transaction in statement["transactions"]:
            row_map = _fuel_card_row_map(statement, transaction)
            values = _row_values(
                row_map,
                ordered_ids,
                distractor_plan,
                row_key=f"{_iso(transaction['date'])}:{transaction['card_number']}:{transaction['receipt_no']}",
                statement_key=f"{statement['account_number']}:{statement['statement_no']}",
            )
            for column_index, value in enumerate(values, start=1):
                sheet.cell(row=row_index, column=column_index, value=value)
            row_index += 1

        row_index += 1
        for label, amount, is_total in _fuel_card_summary_rows(raw_config, statement):
            sheet.cell(row=row_index, column=1, value=label).font = Font(bold=True)
            sheet.cell(row=row_index, column=2, value=_number(amount)).font = Font(bold=is_total)
            row_index += 1

        for column_index in range(1, len(ordered_ids) + 1):
            sheet.column_dimensions[get_column_letter(column_index)].width = 15

    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def generate_fuel_card_statement_csv(raw_config: dict) -> bytes:
    statements = _corrupted_fuel_card_statements(raw_config)
    layout_plan = _layout_plan(raw_config, "CSV")
    distractor_plan = _distractor_plan(raw_config, "CSV")
    distractor_fields = _distractor_field_map(distractor_plan)
    ordered_ids = _augment_field_ids(_ordered_field_ids(layout_plan, list(_FUEL_CARD_HEADER_KEYS)), distractor_plan)

    buffer = StringIO()
    writer = csv.writer(buffer)

    for statement_index, statement in enumerate(statements):
        if statement_index > 0:
            writer.writerow([])
        writer.writerow([_tr(raw_config, "fuel_card_title")])
        writer.writerow([_tr(raw_config, "account_name"), statement["account_name"]])
        writer.writerow([_tr(raw_config, "account_no"), statement["account_number"]])
        writer.writerow([_tr(raw_config, "statement_no"), statement["statement_no"]])
        writer.writerow([_tr(raw_config, "statement_period"), f"{statement['period_start'].isoformat()} to {statement['period_end'].isoformat()}"])
        writer.writerow([_tr(raw_config, "currency"), statement["currency"]])
        writer.writerow([])
        _write_csv_preamble(writer, layout_plan)
        writer.writerow([
            _header_text(raw_config, layout_plan, field_id, _FUEL_CARD_HEADER_KEYS.get(field_id, field_id), distractor_fields)
            for field_id in ordered_ids
        ])
        for transaction in statement["transactions"]:
            row_map = _fuel_card_row_map(statement, transaction)
            writer.writerow(
                _row_values(
                    row_map,
                    ordered_ids,
                    distractor_plan,
                    row_key=f"{_iso(transaction['date'])}:{transaction['card_number']}:{transaction['receipt_no']}",
                    statement_key=f"{statement['account_number']}:{statement['statement_no']}",
                )
            )

        writer.writerow([])
        for label, amount, _is_total in _fuel_card_summary_rows(raw_config, statement):
            writer.writerow([label, _fmt_num(amount, ".2f") if not isinstance(amount, str) else amount])

    return buffer.getvalue().encode("utf-8-sig")


# ── telematics tabular renderers ─────────────────────────────────────────────

_TELEMATICS_FUEL_HEADER_KEYS: dict[str, str] = {
    "period_start": "period_start",
    "period_end": "period_end",
    "vehicle_reg": "vehicle_reg",
    "vehicle_name": "vehicle_name",
    "vehicle_type": "vehicle_type",
    "fuel_type": "fuel_type",
    "distance": "distance",
    "distance_unit": "distance_unit",
    "fuel_used": "fuel_used",
    "unit": "unit",
    "idle_fuel": "idle_fuel",
    "engine_hours": "engine_hours",
    "avg_consumption": "avg_consumption",
}

_TELEMATICS_TRIP_HEADER_KEYS: dict[str, str] = {
    "trip_id": "trip_id",
    "vehicle_reg": "vehicle_reg",
    "driver": "driver",
    "trip_start": "trip_start",
    "trip_end": "trip_end",
    "start_location": "start_location",
    "end_location": "end_location",
    "distance": "distance",
    "distance_unit": "distance_unit",
    "duration": "duration",
    "avg_speed": "avg_speed",
    "fuel_type": "fuel_type",
}

_TELEMATICS_ODOMETER_HEADER_KEYS: dict[str, str] = {
    "period_start": "period_start",
    "period_end": "period_end",
    "vehicle_reg": "vehicle_reg",
    "vehicle_name": "vehicle_name",
    "fuel_type": "fuel_type",
    "odometer_start": "odometer_start",
    "odometer_end": "odometer_end",
    "distance": "distance",
    "distance_unit": "distance_unit",
}


def _telematics_fuel_row_map(row: dict) -> dict[str, Any]:
    return {
        "period_start": _iso(row["period_start"]),
        "period_end": _iso(row["period_end"]),
        "vehicle_reg": row["vehicle_reg"],
        "vehicle_name": row["vehicle_name"],
        "vehicle_type": row["vehicle_type"],
        "fuel_type": row["fuel"],
        "distance": _number(row["distance"]),
        "distance_unit": row["distance_unit"],
        "fuel_used": _number(row["fuel_used"]),
        "unit": row["unit"],
        "idle_fuel": _number(row["idle_fuel"]),
        "engine_hours": _number(row["engine_hours"]),
        "avg_consumption": _number(row["avg_consumption"]),
    }


def _telematics_trip_row_map(row: dict) -> dict[str, Any]:
    return {
        "trip_id": row["trip_id"],
        "vehicle_reg": row["vehicle_reg"],
        "driver": row["driver"],
        "trip_start": row["trip_start"].strftime("%Y-%m-%d %H:%M") if hasattr(row["trip_start"], "strftime") else row["trip_start"],
        "trip_end": row["trip_end"].strftime("%Y-%m-%d %H:%M") if hasattr(row["trip_end"], "strftime") else row["trip_end"],
        "start_location": row["start_location"],
        "end_location": row["end_location"],
        "distance": _number(row["distance"]),
        "distance_unit": row["distance_unit"],
        "duration": row["duration"],
        "avg_speed": _number(row["avg_speed"]),
        "fuel_type": row["fuel"],
    }


def _telematics_odometer_row_map(row: dict) -> dict[str, Any]:
    return {
        "period_start": _iso(row["period_start"]),
        "period_end": _iso(row["period_end"]),
        "vehicle_reg": row["vehicle_reg"],
        "vehicle_name": row["vehicle_name"],
        "fuel_type": row["fuel"],
        "odometer_start": _number(row["odometer_start"]),
        "odometer_end": _number(row["odometer_end"]),
        "distance": _number(row["distance"]),
        "distance_unit": row["distance_unit"],
    }


def _telematics_title_key(raw_config: dict) -> str:
    return {
        "telematics_fuel": "telematics_fuel_title",
        "telematics_trips": "telematics_trips_title",
        "telematics_odometer": "telematics_odometer_title",
    }.get(_document_type(raw_config), "telematics_fuel_title")


def _render_telematics_csv(
    raw_config: dict,
    rows: list[dict],
    header_keys: dict[str, str],
    row_map_builder,
    row_key_builder,
) -> bytes:
    fp = _financial_period(raw_config)
    layout_plan = _layout_plan(raw_config, "CSV")
    distractor_plan = _distractor_plan(raw_config, "CSV")
    distractor_fields = _distractor_field_map(distractor_plan)
    ordered_ids = _augment_field_ids(_ordered_field_ids(layout_plan, list(header_keys)), distractor_plan)

    buffer = StringIO()
    writer = csv.writer(buffer)
    writer.writerow([_tr(raw_config, _telematics_title_key(raw_config))])
    writer.writerow([_tr(raw_config, "reporting_period"), f"{fp['start_date'].isoformat()} to {fp['end_date'].isoformat()}"])
    writer.writerow([_tr(raw_config, "generated"), fp["end_date"].isoformat()])
    writer.writerow([])
    _write_csv_preamble(writer, layout_plan)
    writer.writerow([
        _header_text(raw_config, layout_plan, field_id, header_keys.get(field_id, field_id), distractor_fields)
        for field_id in ordered_ids
    ])
    for row in rows:
        writer.writerow(
            _row_values(
                row_map_builder(row),
                ordered_ids,
                distractor_plan,
                row_key=row_key_builder(row),
                statement_key=str(row.get("company", "")),
            )
        )

    return buffer.getvalue().encode("utf-8-sig")


def _render_telematics_xlsx(
    raw_config: dict,
    rows: list[dict],
    header_keys: dict[str, str],
    row_map_builder,
    row_key_builder,
    sheet_title: str,
) -> bytes:
    fp = _financial_period(raw_config)
    layout_plan = _layout_plan(raw_config, "XLSX")
    distractor_plan = _distractor_plan(raw_config, "XLSX")
    distractor_fields = _distractor_field_map(distractor_plan)
    ordered_ids = _augment_field_ids(_ordered_field_ids(layout_plan, list(header_keys)), distractor_plan)

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = sheet_title[:31]
    sheet["A1"] = _tr(raw_config, _telematics_title_key(raw_config))
    sheet["A1"].font = Font(size=14, bold=True)
    sheet["A2"] = _tr(raw_config, "reporting_period")
    sheet["B2"] = f"{fp['start_date'].isoformat()} to {fp['end_date'].isoformat()}"

    header_row = _write_xlsx_preamble(sheet, 4, layout_plan)
    header_fill = PatternFill(fill_type="solid", fgColor="245C4F")
    for column_index, field_id in enumerate(ordered_ids, start=1):
        cell = sheet.cell(
            row=header_row,
            column=column_index,
            value=_header_text(raw_config, layout_plan, field_id, header_keys.get(field_id, field_id), distractor_fields),
        )
        cell.font = Font(color="FFFFFF", bold=True)
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")

    row_index = header_row + 1
    for row in rows:
        values = _row_values(
            row_map_builder(row),
            ordered_ids,
            distractor_plan,
            row_key=row_key_builder(row),
            statement_key=str(row.get("company", "")),
        )
        for column_index, value in enumerate(values, start=1):
            sheet.cell(row=row_index, column=column_index, value=value)
        row_index += 1

    for column_index in range(1, len(ordered_ids) + 1):
        sheet.column_dimensions[get_column_letter(column_index)].width = 17

    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def _telematics_fuel_row_key(row: dict) -> str:
    return f"{row['vehicle_reg']}:{_iso(row['period_start'])}"


def _telematics_trip_row_key(row: dict) -> str:
    return f"{row['vehicle_reg']}:{row['trip_id']}"


def _telematics_odometer_row_key(row: dict) -> str:
    return f"{row['vehicle_reg']}:{_iso(row['period_start'])}"


def generate_telematics_fuel_csv(raw_config: dict) -> bytes:
    return _render_telematics_csv(
        raw_config,
        _corrupted_telematics_fuel_rows(raw_config),
        _TELEMATICS_FUEL_HEADER_KEYS,
        _telematics_fuel_row_map,
        _telematics_fuel_row_key,
    )


def generate_telematics_fuel_xlsx(raw_config: dict) -> bytes:
    return _render_telematics_xlsx(
        raw_config,
        _corrupted_telematics_fuel_rows(raw_config),
        _TELEMATICS_FUEL_HEADER_KEYS,
        _telematics_fuel_row_map,
        _telematics_fuel_row_key,
        "Fuel Usage",
    )


def generate_telematics_trips_csv(raw_config: dict) -> bytes:
    return _render_telematics_csv(
        raw_config,
        _corrupted_trip_rows(raw_config),
        _TELEMATICS_TRIP_HEADER_KEYS,
        _telematics_trip_row_map,
        _telematics_trip_row_key,
    )


def generate_telematics_trips_xlsx(raw_config: dict) -> bytes:
    return _render_telematics_xlsx(
        raw_config,
        _corrupted_trip_rows(raw_config),
        _TELEMATICS_TRIP_HEADER_KEYS,
        _telematics_trip_row_map,
        _telematics_trip_row_key,
        "Trip History",
    )


def generate_telematics_odometer_csv(raw_config: dict) -> bytes:
    return _render_telematics_csv(
        raw_config,
        _corrupted_odometer_rows(raw_config),
        _TELEMATICS_ODOMETER_HEADER_KEYS,
        _telematics_odometer_row_map,
        _telematics_odometer_row_key,
    )


def generate_telematics_odometer_xlsx(raw_config: dict) -> bytes:
    return _render_telematics_xlsx(
        raw_config,
        _corrupted_odometer_rows(raw_config),
        _TELEMATICS_ODOMETER_HEADER_KEYS,
        _telematics_odometer_row_map,
        _telematics_odometer_row_key,
        "Mileage Summary",
    )
